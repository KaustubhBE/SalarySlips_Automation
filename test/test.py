import json
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
import comtypes.client
import os
import time
import re
import pyautogui
import pandas as pd
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Load config
try:
    with open("config.json", "r") as f:
        config = json.load(f)
except FileNotFoundError:
    print("Error: Configuration file 'config.json' not found.")
    exit(1)

# Email settings
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 465
SENDER_EMAIL = "hrd@bajajearths.com"
SENDER_PASSWORD = "wkcj ajvh exxs qhko"

if not SENDER_EMAIL or not SENDER_PASSWORD:
    print("Error: Sender email and password are missing in the config file.")
    exit(1)

# Load Service Account Credentials
try:
    service_account_file = r"C:\Users\Kaustubh\OneDrive\Desktop\BE_SS_Automation\service_account_credentials.json"
    creds = Credentials.from_service_account_file(
        service_account_file, 
        scopes=[
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/drive'
        ]
    )
except Exception as e:
    print(f"Error loading service account credentials: {e}")
    exit(1)

# Authenticate PyDrive with OAuth2
try:
    oauth2_file = r"C:\Users\Kaustubh\OneDrive\Desktop\BE_SS_Automation\Oauth2.json"
    gauth = GoogleAuth()
    gauth.LoadCredentialsFile(oauth2_file)
    if gauth.credentials is None:
        gauth.LocalWebserverAuth()
        gauth.SaveCredentialsFile(oauth2_file)
    drive = GoogleDrive(gauth)
except Exception as e:
    print(f"Error authenticating Google Drive: {e}")
    exit(1)

# Fetch data from Google Sheets
def fetch_google_sheet_data(sheet_id, sheet_name):
    try:
        client = gspread.authorize(creds)
        sheet = client.open_by_key(sheet_id).worksheet(sheet_name)
        data = sheet.get_all_values()
        return data
    except Exception as e:
        print(f"Error fetching data from Google Sheets (ID: {sheet_id}, Sheet: {sheet_name}): {e}")
        return None

# Preprocess headers
def preprocess_headers(headers):
    return [header.replace("\n", " ").strip().strip('"') for header in headers]

# Convert DOCX to PDF
def convert_docx_to_pdf(input_path, output_path):
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=17)
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        return False

# Upload PDF to Google Drive
def upload_to_google_drive(output_pdf, folder_id, employee_name):
    try:
        # Define the file title
        file_title = f"Salary Slip_{employee_name}_{month}{year}.pdf"

        # Search for an existing file with the same title in the folder
        query = f"'{folder_id}' in parents and title = '{file_title}' and trashed = false"
        existing_files = drive.ListFile({'q': query}).GetList()

        # If a match is found, delete the existing file
        if existing_files:
            for file in existing_files:
                print(f"Found existing file {file['title']} in folder. Deleting it.")
                file.Delete()

        # Create and upload the new file
        file = drive.CreateFile({
            "title": file_title,
            "parents": [{"id": folder_id}]
        })
        file.SetContentFile(output_pdf)
        file.Upload()
        print(f"Uploaded {employee_name}'s salary slip to folder {folder_id}")
    except Exception as e:
        print(f"Error uploading {employee_name}'s file to Google Drive: {e}")

# Fetch employee email
def get_employee_email(employee_name, email_employees):
    for record in email_employees:
        if record.get("Name") == employee_name:
            return record.get("Email ID", "")
    return ""

# Send email with PDF attachment
def send_email_with_attachment(recipient_email, subject, body, attachment_path):
    try:
        msg = MIMEMultipart()
        msg['From'] = SENDER_EMAIL
        msg['To'] = recipient_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html'))

        with open(attachment_path, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment_path)}')
            msg.attach(part)

        with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as server:
            server.login(SENDER_EMAIL, SENDER_PASSWORD)
            server.send_message(msg)
            print(f"Email sent to {recipient_email}")
    except Exception as e:
        print(f"Error sending email to {recipient_email}: {e}")

def get_employee_contact(employee_name, contact_employees):
    if not isinstance(contact_employees, list):
        print("Error: contact_employees is not a list of dictionaries.")
        return ""
    for record in contact_employees:
        if isinstance(record, dict) and record.get("Name") == employee_name:
            return str(record.get("Contact No.", ""))
    return ""

# Function to open WhatsApp Web
def open_whatsapp():
    pyautogui.press('win')  # Open a new browser window
    time.sleep(2)
    pyautogui.typewrite("whatsapp")
    time.sleep(1)
    pyautogui.press('enter')
    time.sleep(5)  # Wait for WhatsApp Web to load

def send_whatsapp_message(contact_name, message, file_path, whatsapp_number, full_month, full_year, month, year):
    try:
        
        # Ensure whatsapp_number is a string
        str_whatsapp_number = pd.Series(whatsapp_number)
        whatsapp_number_str = str_whatsapp_number.to_string(index=False).strip()

        open_whatsapp()

        if not whatsapp_number:
            print(f"Phone number not found for {contact_name}.")
            return

        # Interact with WhatsApp Web using pyautogui
        pyautogui.hotkey('ctrl', 'n')  # Shortcut for new chat
        time.sleep(1)
        pyautogui.typewrite(whatsapp_number_str)  # Type the contact number
        time.sleep(1)
        pyautogui.press('tab')
        time.sleep(0.5)
        pyautogui.press('tab')
        time.sleep(0.5)
        pyautogui.press('enter')  # Open the chat
        time.sleep(1)

        for line in message:
            pyautogui.typewrite(line)
            time.sleep(1)
            pyautogui.hotkey('shift','enter')
            time.sleep(1)
            # pyautogui.hotkey('ctrl', 'v')  # To paste custom message
        pyautogui.press('enter')
        time.sleep(1)

        # Attach the file
        pyautogui.hotkey('shift', 'tab')  # Shortcut for attachments
        time.sleep(1)
        pyautogui.press('enter')
        time.sleep(0.5)
        pyautogui.press('down')
        time.sleep(0.5)
        pyautogui.press('down')  # Navigate to "Document"
        time.sleep(0.5)
        pyautogui.press('enter')
        time.sleep(2)

        pyautogui.typewrite(file_path)  # Type the file path
        time.sleep(2)
        pyautogui.press('enter')
        time.sleep(1)
        pyautogui.press('enter')
        time.sleep(1)
        
        print(f"Sent salary slip to {contact_name} ({whatsapp_number_str}) via WhatsApp.")
    except Exception as e:
        print(f"Error sending WhatsApp message to {contact_name}: {e}")

def format_file_path(file_path):
    if isinstance(file_path, str):
        return file_path.replace("\\\\", "\\")
    elif isinstance(file_path, list):
        return [f.replace("\\\\", "\\") for f in file_path]
    return file_path

# Generate and process salary slips
def process_salary_slip(template_path, output_dir, employee_data, headers, drive_data, email_employees, contact_employees, month, year, full_month, full_year):
    headers = preprocess_headers(headers)
    placeholders = dict(zip(headers, employee_data))

    # Add month and year placeholders to the dictionary
    placeholders["Month"] = full_month
    placeholders["Year"] = full_year
    
    # Calculate components of salary
    try:
        present_salary = float(re.sub(r'[^\d.]', '', placeholders.get("Present Salary", "")))
        placeholders["BS"] = str(round(present_salary * 0.40))
        placeholders["HRA"] = str(round(present_salary * 0.20))
        placeholders["SA"] = str(round(present_salary * 0.40))
    except ValueError:
        print(f"Invalid Present Salary for {placeholders.get('Name', 'Unknown')}. Skipping.")
        return

    # Load template and replace placeholders
    try:
        template = Document(template_path)
        for paragraph in template.paragraphs:
            for run in paragraph.runs:
                for placeholder, value in placeholders.items():
                    if f"{{{placeholder}}}" in run.text:
                        run.text = run.text.replace(f"{{{placeholder}}}", value)

        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for placeholder, value in placeholders.items():
                                if f"{{{placeholder}}}" in run.text:
                                    run.text = run.text.replace(f"{{{placeholder}}}", value)

        # Save output files
        employee_name = re.sub(r'[^\w\s]', '', placeholders.get("Name", "Employee"))
        output_docx = os.path.join(output_dir, f"Salary_Slip_{employee_name}_{month}{year}.docx")
        template.save(output_docx)
        output_pdf = os.path.join(output_dir, f"Salary_Slip_{employee_name}_{month}{year}.pdf")

        if convert_docx_to_pdf(output_docx, output_pdf):
            # Upload to Google Drive
            folder_id = next((item.get("Google Drive ID") for item in drive_data if item.get("Name") == placeholders.get("Name")), None)
            if folder_id:
                upload_to_google_drive(output_pdf, folder_id, employee_name)

            # Send email
            recipient_email = get_employee_email(placeholders.get("Name"), email_employees)
            if recipient_email:
                email_subject = f"Salary Slip for {full_month} {full_year} - Bajaj Earths Pvt. Ltd."
                email_body = f"""
                <html>
                <body>
                <p>Dear <b>{placeholders.get('Name')}</b>,</p>
                <p>Please find attached your <b>salary slip</b> for the month of <b>{full_month} {full_year}</b>.</p>
                <p>This document includes:</p>
                <ul>
                <li>Earnings Breakdown</li>
                <li>Deductions Summary</li>
                <li>Net Salary Details</li>
                </ul>
                <p>Kindly review the salary slip, and if you have any questions or concerns, please feel free to reach out to the HR department.</p>
                <p>Thanks & Regards,</p>
                </body>
                </html>
                """
                # send_email_with_attachment(recipient_email, email_subject, email_body, output_pdf)
            else:
                print(f"No email found for {placeholders.get('Name')}.")

            # Send salary slip via WhatsApp
            contact_name = placeholders.get("Name")
            whatsapp_number = get_employee_contact(contact_name, contact_employees)
            if whatsapp_number:
                message = [
                            f"Dear *{placeholders.get('Name')}*,",
                            "",
                            f"Please find attached your *salary slip* for the month of *{full_month} {full_year}*.",
                            "",
                            " This document includes:",
                            "   -  Earnings Breakdown",
                            "   -  Deductions Summary",
                            "   -  Net Salary Details",
                            "",
                            "Kindly review the salary slip, and if you have any questions or concerns, please feel free to reach out to the HR department.",
                            "",
                            "Thanks & Regards,",
                            "HR Department",
                            "Bajaj Earths Pvt. Ltd.",
                            "+91 - 86557 88172"
                        ]
                file_path = os.path.join(output_dir, f"Salary_Slip_{contact_name}_{month}{year}.pdf")
                #  send_whatsapp_message(contact_name, message, file_path, whatsapp_number, full_month, full_year, month, year)
    except Exception as e:
        print(f"Error processing salary slip for {placeholders.get('Name', 'Unknown')}: {e}")

# Main Script
if __name__ == "__main__":
    sheet_id_salary = "1HUUF8g3GJ3ZaPyUsRhRgCURu5m0prtZRy7kEIZoc10M"
    sheet_id_employees = "1DmV91n5ryeAJ7t4xM4jMzv99X5f5wbtf_1LsCOOBj8Q"
    template_path = r"C:\Users\Kaustubh\OneDrive\Desktop\BE_SS_Automation\ssformat.docx"
    output_dir = r"C:\Users\Kaustubh\OneDrive\Desktop\BE_SS_Automation\Salary_Slips"

    if not all([sheet_id_salary, sheet_id_employees, template_path, output_dir]):
        print("Error: Missing required configuration values.")
        exit(1)

    os.makedirs(output_dir, exist_ok=True)
    full_month = input("Enter sheet name for the salary slip (e.g., January): ").strip()
    full_year = input("Enter the year (e.g., 2024): ").strip()

    try:
        month = full_month[:3].capitalize()  # Ensure the abbreviation is capitalized (e.g., "Jan")
        year = full_year[-2:]               # Extract the last two digits of the year (e.g., "24")
        
        if not month.isalpha() or len(full_month) < 3:
            raise ValueError("Invalid month input. Please enter a valid full month name.")
        if not full_year.isdigit() or len(full_year) != 4:
            raise ValueError("Invalid year input. Please enter a valid 4-digit year.")
    except Exception as e:
        print(f"Error: {e}")

    salary_data = fetch_google_sheet_data(sheet_id_salary, month)
    drive_data = fetch_google_sheet_data(sheet_id_employees, "Official Details")
    email_data = fetch_google_sheet_data(sheet_id_employees, "Onboarding Details")
    contact_data = fetch_google_sheet_data(sheet_id_employees, "Onboarding Details")

    if salary_data and drive_data and email_data:
        salary_headers = salary_data[1]  # Assuming headers are in the second row (index 1)
        employees = salary_data[2:]      # Employee data starts from the third row (index 2)

        drive_headers = drive_data[1]
        drive_employees = [dict(zip(drive_headers, row)) for row in drive_data[2:]]

        email_headers = email_data[1]
        email_employees = [dict(zip(email_headers, row)) for row in email_data[2:]]

        contact_headers = contact_data[1]
        contact_employees = [dict(zip(contact_headers, row)) for row in contact_data[2:]]  # Ensure correct format

        for employee in employees:
            process_salary_slip(template_path, output_dir, employee, salary_headers, drive_employees, email_employees, contact_employees, month, year, full_month, full_year)
    else:
        print("Error: Failed to fetch data from Google Sheets.")
