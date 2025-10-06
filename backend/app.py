import os
import logging
import webbrowser
import sys
import threading
import time
from flask import Flask, request, jsonify, Response, g, session, redirect, url_for, make_response
from flask_cors import CORS
from logging.handlers import RotatingFileHandler
from Utils.fetch_data import fetch_google_sheet_data
from Utils.process_utils import *
from google.oauth2 import service_account
from googleapiclient.discovery import build
from Utils.config import CLIENT_SECRETS_FILE, drive, creds
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
from Utils.auth import auth_bp
from Utils.email_utils import send_email_smtp
from Utils.whatsapp_utils import (
    send_whatsapp_message,
    get_employee_contact,
    WhatsAppNodeClient,
    WHATSAPP_NODE_SERVICE_URL
)
from firebase_admin import firestore
from Utils.firebase_utils import (
    add_user as firebase_add_user,
    get_user_by_id,
    get_user_by_email,
    get_all_users as firebase_get_all_users,
    update_user_role as firebase_update_role,
    delete_user as firebase_delete_user,
    add_salary_slip,
    get_salary_slips_by_user,
    update_user_app_password,
    update_user,
    db,
    add_order,
    get_orders_by_factory,
    get_order_by_id,
    update_order_status,
    delete_order,
    get_all_orders,
    get_factory_initials,
    get_next_order_id,
    get_user_oauth_tokens,
    update_user_oauth_tokens
)
import json
from docx import Document
import re
import base64
import requests
from datetime import datetime, timedelta
import gspread
from dotenv import load_dotenv

# ============================================================================
# CENTRALIZED RBAC CONFIGURATION
# ============================================================================

# Centralized factory RBAC configuration (mirrors frontend config)
FACTORY_RBAC_CONFIG = {
    'gulbarga': {
        'name': 'Gulbarga',
        'document_name': 'GB',
        'departments': {
            'store': {
                'name': 'Store',
                'services': {
                    'gb_place_order': {'name': 'Place Order', 'permission': 'gb_place_order'},
                    'gb_material_list': {'name': 'Add Material', 'permission': 'gb_material_list'},
                    'gb_material_inward': {'name': 'Material Inward', 'permission': 'gb_material_inward'},
                    'gb_material_outward': {'name': 'Material Outward', 'permission': 'gb_material_outward'}
                }
            },
            'humanresource': {
                'name': 'Human Resource',
                'services': {
                    'gb_single_processing': {'name': 'Single Processing', 'permission': 'gb_single_processing'},
                    'gb_batch_processing': {'name': 'Batch Processing', 'permission': 'gb_batch_processing'}
                }
            },
            'operations': {
                'name': 'Operations',
                'services': {
                    'gb_general_reports': {'name': 'General Reports', 'permission': 'gb_general_reports'}
                }
            }
        }
    },
    'kerur': {
        'name': 'Kerur',
        'document_name': 'KR',
        'departments': {
            'store': {
                'name': 'Store',
                'services': {
                    'kr_place_order': {'name': 'Place Order', 'permission': 'kr_place_order'},
                    'kr_material_list': {'name': 'Material List', 'permission': 'kr_material_list'},
                    'kr_material_inward': {'name': 'Material Inward', 'permission': 'kr_material_inward'},
                    'kr_material_outward': {'name': 'Material Outward', 'permission': 'kr_material_outward'},
                    'kr_order_status': {'name': 'Order Status', 'permission': 'kr_order_status'}
                }
            },
            'humanresource': {
                'name': 'Human Resource',
                'services': {
                    'kr_single_processing': {'name': 'Single Processing', 'permission': 'kr_single_processing'},
                    'kr_batch_processing': {'name': 'Batch Processing', 'permission': 'kr_batch_processing'}
                }
            },
            'operations': {
                'name': 'Operations',
                'services': {
                    'kr_general_reports': {'name': 'General Reports', 'permission': 'kr_general_reports'},
                    'kr_reactor_reports': {'name': 'Reactor Reports', 'permission': 'kr_reactor_reports'}
                }
            }
        }
    },
    'humnabad': {
        'name': 'Humnabad',
        'document_name': 'HB',
        'departments': {
            'store': {
                'name': 'Store',
                'services': {
                    'hb_place_order': {'name': 'Place Order', 'permission': 'hb_place_order'},
                    'hb_material_list': {'name': 'Add Material', 'permission': 'hb_material_list'},
                    'hb_material_inward': {'name': 'Material Inward', 'permission': 'hb_material_inward'},
                    'hb_material_outward': {'name': 'Material Outward', 'permission': 'hb_material_outward'}
                }
            },
            'humanresource': {
                'name': 'Human Resource',
                'services': {
                    'hb_single_processing': {'name': 'Single Processing', 'permission': 'hb_single_processing'},
                    'hb_batch_processing': {'name': 'Batch Processing', 'permission': 'hb_batch_processing'}
                }
            },
            'operations': {
                'name': 'Operations',
                'services': {
                    'hb_general_reports': {'name': 'General Reports', 'permission': 'hb_general_reports'}
                }
            }
        }
    },
    'omkar': {
        'name': 'Omkar',
        'document_name': 'OM',
        'departments': {
            'store': {
                'name': 'Store',
                'services': {
                    'om_place_order': {'name': 'Place Order', 'permission': 'om_place_order'},
                    'om_material_list': {'name': 'Add Material', 'permission': 'om_material_list'},
                    'om_material_inward': {'name': 'Material Inward', 'permission': 'om_material_inward'},
                    'om_material_outward': {'name': 'Material Outward', 'permission': 'om_material_outward'}
                }
            },
            'humanresource': {
                'name': 'Human Resource',
                'services': {
                    'om_single_processing': {'name': 'Single Processing', 'permission': 'om_single_processing'},
                    'om_batch_processing': {'name': 'Batch Processing', 'permission': 'om_batch_processing'}
                }
            },
            'operations': {
                'name': 'Operations',
                'services': {
                    'om_general_reports': {'name': 'General Reports', 'permission': 'om_general_reports'}
                }
            }
        }
    },
    'padmavati': {
        'name': 'Padmavati',
        'document_name': 'PV',
        'departments': {
            'store': {
                'name': 'Store',
                'services': {
                    'pv_place_order': {'name': 'Place Order', 'permission': 'pv_place_order'},
                    'pv_material_list': {'name': 'Add Material', 'permission': 'pv_material_list'},
                    'pv_material_inward': {'name': 'Material Inward', 'permission': 'pv_material_inward'},
                    'pv_material_outward': {'name': 'Material Outward', 'permission': 'pv_material_outward'}
                }
            },
            'humanresource': {
                'name': 'Human Resource',
                'services': {
                    'pv_single_processing': {'name': 'Single Processing', 'permission': 'pv_single_processing'},
                    'pv_batch_processing': {'name': 'Batch Processing', 'permission': 'pv_batch_processing'}
                }
            },
            'operations': {
                'name': 'Operations',
                'services': {
                    'pv_general_reports': {'name': 'General Reports', 'permission': 'pv_general_reports'}
                }
            }
        }
    },
    'headoffice': {
        'name': 'Head Office',
        'document_name': 'HO',
        'departments': {
            'store': {
                'name': 'Store',
                'services': {
                    'ho_material_list': {'name': 'Material List', 'permission': 'ho_material_list'}
                }
            },
            'humanresource': {
                'name': 'Human Resource',
                'services': {
                    'ho_single_processing': {'name': 'Single Processing', 'permission': 'ho_single_processing'},
                    'ho_batch_processing': {'name': 'Batch Processing', 'permission': 'ho_batch_processing'}
                }
            },
            'accounts': {
                'name': 'Accounts',
                'services': {
                    # No specific services found in HO_Services folder
                }
            },
            'marketing': {
                'name': 'Marketing',
                'services': {
                    # No specific services found in HO_Services folder
                }
            },
            'operations': {
                'name': 'Operations',
                'services': {
                    'ho_general_reports': {'name': 'General Reports', 'permission': 'ho_general_reports'}
                }
            }
        }
    }
}

def generate_permission_metadata(role, user_permission_metadata=None):
    """
    Generate permission metadata based on role and user-specific permissions.
    Admin gets all permissions, regular users get specific permissions.
    """
    logger.info(f"generate_permission_metadata called with role: {role}, user_permission_metadata: {user_permission_metadata}")
    
    if role == 'admin':
        # Admin gets access to all factories, departments, and services
        all_factories = list(FACTORY_RBAC_CONFIG.keys())
        all_departments = {}
        all_services = {}
        
        for factory in all_factories:
            factory_config = FACTORY_RBAC_CONFIG[factory]
            factory_short_form = factory_config['document_name'].lower()  # e.g., 'gb', 'kr', 'hb'
            
            factory_departments = list(factory_config['departments'].keys())
            # Store departments with factory short form prefix (e.g., gb_store, kr_humanresource)
            prefixed_departments = [f"{factory_short_form}_{dept}" for dept in factory_departments]
            all_departments[factory] = prefixed_departments
            
            for department in factory_departments:
                service_key = f"{factory}.{department}"
                department_services = list(FACTORY_RBAC_CONFIG[factory]['departments'][department]['services'].keys())
                all_services[service_key] = department_services
        
        result = {
            'factories': all_factories,
            'departments': all_departments,
            'services': all_services
        }
        logger.info(f"Generated admin permissions: {result}")
        return result
    
    # For regular users, use the provided permission metadata or return empty
    result = user_permission_metadata or {
        'factories': [],
        'departments': {},
        'services': {}
    }
    logger.info(f"Generated user permissions: {result}")
    return result


# Load environment variables from .env file
load_dotenv()

# Removed complex thread pool executor - using direct client calls instead

# Initialize gspread client
try:
    
    client = gspread.authorize(creds)
except Exception as e:
    logger.error(f"Error initializing gspread client: {e}")
    client = None

# Old RBAC configuration removed - using centralized FACTORY_RBAC_CONFIG instead

# Configure logging first
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('app.log')
    ]
)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "your_default_secret_key")  # Set secret key for sessions

# Configure session settings
app.config.update(
    SESSION_COOKIE_SECURE=True,  # Only send cookie over HTTPS
    SESSION_COOKIE_HTTPONLY=True,  # Prevent JavaScript access
    SESSION_COOKIE_SAMESITE='Lax',  # Allow cross-site requests
    SESSION_COOKIE_DOMAIN='.bajajearths.com',  # Set domain to allow subdomain access
    PERMANENT_SESSION_LIFETIME=timedelta(hours=24),  # Session expires in 24 hours
    SESSION_COOKIE_PATH='/',  # Set cookie path
    SESSION_REFRESH_EACH_REQUEST=True  # Refresh session on each request
)

logger.info("Flask app initialized")

# Frontend URL
# FRONTEND_URL = "http://uatadmin.bajajearths.com"
# _frontend_opened = False

# Load configurations from environment variables
def get_base_dir():
    if getattr(sys, 'frozen', False):
        # If running as compiled executable
        return sys._MEIPASS
    # If running in development
    return os.path.dirname(os.path.abspath(__file__))

BASE_DIR = get_base_dir()
OUTPUT_DIR = os.getenv("OUTPUT_DIR", os.path.join(os.path.expanduser("~"), "Salary_Slips"))
TEMPLATE_PATH = os.getenv("TEMPLATE_PATH", os.path.join(BASE_DIR, "ssformat.docx"))
LOG_FILE_PATH = os.getenv("LOG_FILE_PATH", os.path.join(os.path.expanduser("~"), "app.log"))
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO")

# Print the actual template path for debugging
print("Template path: {}".format(TEMPLATE_PATH))
print("Base directory: {}".format(BASE_DIR))

# Ensure directories exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Configure logging
if not app.logger.handlers:
    handler = RotatingFileHandler(LOG_FILE_PATH, maxBytes=10*1024*1024, backupCount=5)
    handler.setLevel(logging.INFO)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    app.logger.addHandler(handler)
    app.logger.setLevel(logging.INFO)
    handler.propagate = False

from flask_cors import CORS

# CORS Configuration
app.config['CORS_HEADERS'] = 'Content-Type'
app.config['CORS_ORIGINS'] = [
    "https://uatadmin.bajajearths.com",
    "https://uatwhatsapp.bajajearths.com",
    "http://uatadmin.bajajearths.com",
]
app.config['CORS_METHODS'] = ["GET", "POST", "PUT", "DELETE", "OPTIONS"]
app.config['CORS_ALLOW_HEADERS'] = ["Content-Type", "Authorization", "X-User-Role", "X-User-Email", "Accept", "Origin", "X-Requested-With"]
app.config['CORS_EXPOSE_HEADERS'] = ["Content-Type", "Authorization", "X-User-Email"]
app.config['CORS_SUPPORTS_CREDENTIALS'] = True
app.config['CORS_MAX_AGE'] = 120

# Initialize CORS with proper configuration
CORS(app, 
     origins=app.config['CORS_ORIGINS'],
     methods=app.config['CORS_METHODS'],
     allow_headers=app.config['CORS_ALLOW_HEADERS'],
     expose_headers=app.config['CORS_EXPOSE_HEADERS'],
     supports_credentials=True,
     max_age=app.config['CORS_MAX_AGE'])

@app.before_request
def log_request_info():
    """Log only essential request information for debugging"""
    # Comment out or remove verbose logs
    # logger.info('Headers: %s', request.headers)
    # logger.info('Body: %s', request.get_data())
    # logger.info('URL: %s', request.url)
    # Only log method and URL if needed
    # logger.info('Method: %s URL: %s', request.method, request.url)
    pass

@app.after_request
def after_request(response):
    """Add CORS headers to all responses, but do not log response status or headers."""
    # Flask-CORS is already handling CORS headers, so we don't need to add them manually
    # logger.info('Response status: %s', response.status)
    # logger.info('Response headers: %s', response.headers)
    # origin = request.headers.get('Origin')
    # logger.info('Request origin: %s', origin)
    return response

# Manual OPTIONS handler for better CORS support
@app.route("/api/<path:endpoint>", methods=["OPTIONS"])
def handle_options(endpoint):
    """Handle preflight OPTIONS requests for CORS"""
    logger.info(f"Handling OPTIONS request for endpoint: {endpoint}")
    logger.info(f"Request origin: {request.headers.get('Origin')}")
    logger.info(f"Request headers: {dict(request.headers)}")
    
    response = make_response()
    response.headers.add("Access-Control-Allow-Origin", "https://uatadmin.bajajearths.com")
    response.headers.add("Access-Control-Allow-Headers", "Content-Type, Authorization, X-User-Role, X-User-Email, Accept, Origin, X-Requested-With")
    response.headers.add("Access-Control-Allow-Methods", "GET, POST, PUT, DELETE, OPTIONS")
    response.headers.add("Access-Control-Allow-Credentials", "true")
    response.headers.add("Access-Control-Max-Age", "120")
    
    logger.info(f"Response headers: {dict(response.headers)}")
    return response

@app.route("/api/preview-file", methods=["POST"])
def preview_file():
    logger.info('Handling preview file request')

    temp_path = None
    try:
        # Get user_id from session for authentication and temp directory
        user_id = session.get('user', {}).get('email')
        if not user_id:
            logger.error("No user_id found in session. User must be logged in to preview files.")
            return jsonify({"error": "User not authenticated"}), 401

        if 'file' not in request.files:
            logger.error('No file in request')
            return jsonify({"error": "No file provided"}), 400

        file = request.files['file']
        if file.filename == '':
            logger.error('Empty filename')
            return jsonify({"error": "No file selected"}), 400

        logger.info('Processing file: %s', file.filename)
        
        # Get user-specific temporary directory for file preview
        from Utils.temp_manager import get_user_temp_dir
        temp_dir = get_user_temp_dir(user_id, BASE_DIR)
        temp_path = os.path.join(temp_dir, file.filename)
        
        # Save the uploaded file
        file.save(temp_path)
        logger.info('File saved to: %s', temp_path)
        
        # Read the file content
        content = process_reports(temp_path)
        logger.info('File processed successfully')
        return jsonify({"content": content})
    except Exception as e:
        logger.error("Error previewing file: {}".format(e), exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        # Clean up only the specific file that was processed, not the entire directory
        # This prevents race conditions when multiple files are being previewed simultaneously
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                logger.info(f'Cleaned up preview file: {temp_path}')
            except Exception as cleanup_error:
                logger.warning(f'Failed to cleanup preview file {temp_path}: {cleanup_error}')

# Register blueprints
app.register_blueprint(auth_bp, url_prefix='/api/auth')

def get_drive_service():
    try:
        credentials = service_account.Credentials.from_service_account_file(CLIENT_SECRETS_FILE)
        service = build('drive', 'v3', credentials=credentials)
        return service
    except Exception as e:
        app.logger.error("Error initializing Google Drive service: {}".format(e))
        raise

def get_user_drive_service(user_id):
    """Get Google Drive service using user's OAuth tokens"""
    try:
        oauth_tokens = get_user_oauth_tokens(user_id)
        if not oauth_tokens or not oauth_tokens.get('access_token'):
            raise Exception("No OAuth tokens found for user")
        
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        
        credentials = Credentials(
            token=oauth_tokens['access_token'],
            refresh_token=oauth_tokens['refresh_token'],
            token_uri="https://oauth2.googleapis.com/token",
            client_id=os.environ.get('GOOGLE_CLIENT_ID'),
            client_secret=os.environ.get('GOOGLE_CLIENT_SECRET')
        )
        
        # Refresh token if needed
        if credentials.expired and credentials.refresh_token:
            credentials.refresh(Request())
            # Update the refreshed token in database
            update_user_oauth_tokens(
                user_id, 
                credentials.token, 
                credentials.refresh_token,
                oauth_tokens.get('granted_scopes', [])
            )
        
        service = build('drive', 'v3', credentials=credentials)
        return service
    except Exception as e:
        app.logger.error(f"Error initializing user Google Drive service: {e}")
        raise

def get_user_sheets_service(user_id):
    """Get Google Sheets service using user's OAuth tokens"""
    try:
        oauth_tokens = get_user_oauth_tokens(user_id)
        if not oauth_tokens or not oauth_tokens.get('access_token'):
            raise Exception("No OAuth tokens found for user")
        
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        
        credentials = Credentials(
            token=oauth_tokens['access_token'],
            refresh_token=oauth_tokens['refresh_token'],
            token_uri="https://oauth2.googleapis.com/token",
            client_id=os.environ.get('GOOGLE_CLIENT_ID'),
            client_secret=os.environ.get('GOOGLE_CLIENT_SECRET')
        )
        
        # Refresh token if needed
        if credentials.expired and credentials.refresh_token:
            credentials.refresh(Request())
            # Update the refreshed token in database
            update_user_oauth_tokens(
                user_id, 
                credentials.token, 
                credentials.refresh_token,
                oauth_tokens.get('granted_scopes', [])
            )
        
        service = build('sheets', 'v4', credentials=credentials)
        return service
    except Exception as e:
        app.logger.error(f"Error initializing user Google Sheets service: {e}")
        raise

@app.before_request
def load_user():
    g.user_role = request.headers.get('X-User-Role', 'user')

@app.route("/api/add_user", methods=["POST"])
def add_user():
    try:
        data = request.json
        username = data.get('username')
        email = data.get('email')
        role = data.get('role')
        password = generate_password_hash(data.get('password'))
        app_password = data.get('appPassword') or data.get('app_password')
        permission_metadata = data.get('permission_metadata', {})

        # Debug logging
        logger.info(f"Add user request - Role: {role}")
        logger.info(f"Add user request - Original permission_metadata: {permission_metadata}")

        if not all([username, email, role, password]):
            return jsonify({"error": "Missing required fields"}), 400

        # Generate permission metadata using centralized configuration
        permission_metadata = generate_permission_metadata(role, permission_metadata)
        
        # Debug logging
        logger.info(f"Add user request - Generated permission_metadata: {permission_metadata}")

        user_id = firebase_add_user(
            username=username, 
            email=email, 
            role=role, 
            password_hash=password, 
            app_password=app_password,
            permission_metadata=permission_metadata
        )
        return jsonify({"message": "User added successfully", "user_id": user_id}), 201

    except Exception as e:
        logger.error("Error adding user: {}".format(e))
        return jsonify({"error": str(e)}), 500

@app.route("/api/delete_user", methods=["POST"])
def delete_user():
    try:
        # Check if user is logged in
        current_user = session.get('user')
        if not current_user:
            return jsonify({"error": "User not authenticated"}), 401

        data = request.json
        user_id = data.get('user_id')

        if not user_id:
            return jsonify({"error": "User ID is required"}), 400

        # Get target user
        target_user = get_user_by_id(user_id)
        if not target_user:
            return jsonify({"error": "User not found"}), 404

        current_user_role = current_user.get('role')
        target_user_role = target_user.get('role')

        # Admin can delete any user
        if current_user_role == 'admin':
            firebase_delete_user(user_id)
            return jsonify({"message": "User deleted successfully"}), 200
        
        # Regular users cannot delete other users
        else:
            return jsonify({"error": "Insufficient permissions to delete users"}), 403

    except Exception as e:
        logger.error("Error deleting user: {}".format(e))
        return jsonify({"error": str(e)}), 500

@app.route("/api/update_role", methods=["POST"])
def update_role():
    try:
        # Check if user is logged in
        current_user = session.get('user')
        if not current_user:
            return jsonify({"error": "User not authenticated"}), 401

        data = request.json
        user_id = data.get('user_id')
        new_role = data.get('role')

        if not all([user_id, new_role]):
            return jsonify({"error": "User ID and role are required"}), 400

        # Get target user
        target_user = get_user_by_id(user_id)
        if not target_user:
            return jsonify({"error": "User not found"}), 404

        current_user_role = current_user.get('role')
        target_user_role = target_user.get('role')

        # Admin can update any user's role
        if current_user_role == 'admin':
            # Update role
            firebase_update_role(user_id, new_role)
            
            # Only update permissions if changing TO admin (admin gets all permissions)
            # or if changing FROM admin (preserve existing permissions for regular users)
            if new_role == 'admin':
                # Admin role gets all permissions automatically
                update_user_comprehensive_permissions(user_id, {
                    'permissions': {
                        'inventory': True,
                        'reports': True,
                        'single_processing': True,
                        'batch_processing': True,
                        'expense_management': True,
                        'marketing_campaigns': True,
                        'reactor_reports': True
                    }
                })
            # For other role changes, preserve existing permissions (don't overwrite)
            
            return jsonify({"message": "Role updated successfully"}), 200
        
        # Regular users cannot update roles
        else:
            return jsonify({"error": "Insufficient permissions to update user roles"}), 403

    except Exception as e:
        logger.error("Error updating role: {}".format(e))
        return jsonify({"error": str(e)}), 500
    


@app.route("/api/get_users", methods=["GET"])
def get_users():
    try:
        users = firebase_get_all_users()
        return jsonify(users), 200
        
    except Exception as e:
        logger.error("Error getting users: {}".format(e))
        return jsonify({"error": str(e)}), 500

@app.route("/api/user/<string:user_id>", methods=["GET"])
def get_user(user_id):
    try:
        user = get_user_by_id(user_id)
        if not user:
            return jsonify({"error": "User not found"}), 404
        return jsonify(user), 200

    except Exception as e:
        logger.error("Error getting user: {}".format(e))
        return jsonify({"error": str(e)}), 500

@app.route("/api/generate-salary-slip-single", methods=["POST"])
def generate_salary_slip_single():
    try:
        # Get user_id from session
        user_id = session.get('user', {}).get('email')
        
        user_inputs = request.json
        app.logger.info("Processing single salary slip request")

        # Check for required data structure
        if not user_inputs.get("months_data"):
            return jsonify({"error": "months_data is required"}), 400

        employee_identifier = user_inputs.get("employee_code")
        if not employee_identifier or employee_identifier.strip() == "":
            return jsonify({"error": "employee_code is required and cannot be empty"}), 400

        send_whatsapp = user_inputs.get("send_whatsapp", False)
        send_email = user_inputs.get("send_email", False)

        # Initialize list to collect PDFs
        collected_pdfs = []
        results = []
        employee = None
        drive_employees = None
        email_employees = None
        contact_employees = None
        
        # First pass: Generate all PDFs
        for month_data in user_inputs["months_data"]:
            try:
                sheet_id_salary = month_data.get("sheet_id_salary")
                sheet_id_drive = month_data.get("sheet_id_drive")
                full_month = month_data.get("month")
                full_year = month_data.get("year")

                if not all([sheet_id_salary, sheet_id_drive, full_month, full_year]):
                    results.append({
                        "month": full_month,
                        "year": full_year,
                        "status": "error",
                        "message": "Missing required data for this month"
                    })
                    continue

                sheet_name = full_month[:3]

                # Fetch data
                try:
                    salary_data = fetch_google_sheet_data(sheet_id_salary, sheet_name)
                    drive_data = fetch_google_sheet_data(sheet_id_salary, "Salary Details")
                    email_data = fetch_google_sheet_data(sheet_id_drive, "Onboarding Details")
                    contact_data = fetch_google_sheet_data(sheet_id_drive, "Onboarding Details")
                except Exception as e:
                    results.append({
                        "month": full_month,
                        "year": full_year,
                        "status": "error",
                        "message": "Error fetching data: {}".format(str(e))
                    })
                    continue

                if not all([salary_data, drive_data, email_data, contact_data]):
                    results.append({
                        "month": full_month,
                        "year": full_year,
                        "status": "error",
                        "message": "Failed to fetch required data"
                    })
                    continue

                # Process data
                salary_headers = salary_data[1]
                employees = salary_data[2:]
                
                # Convert salary data to dictionaries
                salary_employees = [dict(zip(salary_headers, row)) for row in employees]

                drive_headers = drive_data[1]
                drive_employees = [dict(zip(drive_headers, row)) for row in drive_data[2:]]

                email_headers = email_data[1]
                email_employees = [dict(zip(email_headers, row)) for row in email_data[2:]]

                contact_headers = contact_data[1]
                contact_employees = [dict(zip(contact_headers, row)) for row in contact_data[2:]]

                # Find employee
                employee = next((emp for emp in drive_employees 
                               if emp.get('Employee\nCode') == employee_identifier), None)
                if not employee:
                    results.append({
                        "month": full_month,
                        "year": full_year,
                        "status": "error",
                        "message": "Employee not found in drive data"
                    })
                    continue

                # Find employee in salary data
                salary_employee = next((emp for emp in salary_employees 
                                      if emp.get('Employee\nCode') == employee_identifier), None)
                if not salary_employee:
                    results.append({
                        "month": full_month,
                        "year": full_year,
                        "status": "error",
                        "message": "Employee salary data not found"
                    })
                    continue

                # Process salary slip
                try:
                    # Convert dictionary values back to list in the correct order
                    employee_data = [str(salary_employee.get(header, '')) for header in salary_headers]
                    pdf_path = process_salary_slip(
                        template_path=TEMPLATE_PATH,
                        output_dir=OUTPUT_DIR,
                        employee_identifier=employee_identifier,
                        employee_data=employee_data,
                        headers=salary_headers,
                        drive_data=drive_employees,
                        email_employees=email_employees,
                        contact_employees=contact_employees,
                        month=sheet_name,
                        year=str(full_year)[-2:],
                        full_month=full_month,
                        full_year=full_year,
                        send_whatsapp=send_whatsapp,
                        send_email=send_email,
                        is_special=len(user_inputs["months_data"]) > 1,
                        months_data=user_inputs["months_data"],
                        collected_pdfs=collected_pdfs
                    )
                    if pdf_path and pdf_path.get("success") and pdf_path.get("output_file"):
                        collected_pdfs.append(pdf_path["output_file"])
                        results.append({
                            "month": full_month,
                            "year": full_year,
                            "status": "success",
                            "message": "Salary slip generated successfully",
                            "pdf_path": pdf_path["output_file"]
                        })
                    else:
                        results.append({
                            "month": full_month,
                            "year": full_year,
                            "status": "error",
                            "message": "Failed to generate salary slip" + (f": {pdf_path.get('errors', [])}" if pdf_path and isinstance(pdf_path, dict) else "")
                        })
                except Exception as e:
                    results.append({
                        "month": full_month,
                        "year": full_year,
                        "status": "error",
                        "message": "Error processing salary slip: {}".format(str(e))
                    })

            except Exception as e:
                results.append({
                    "month": full_month if 'full_month' in locals() else "Unknown",
                    "year": full_year if 'full_year' in locals() else "Unknown",
                    "status": "error",
                    "message": "Error processing month: {}".format(str(e))
                })

        # Only proceed with notifications if we have successfully generated PDFs
        # Note: WhatsApp notifications are sent here to prevent duplicates from process_salary_slip function
        if collected_pdfs and employee and drive_employees and email_employees and contact_employees:
            # Send email if enabled
            if send_email:
                try:
                    recipient_email = get_employee_email(employee.get("Name"), email_employees)
                    if recipient_email:
                        is_special = len(user_inputs["months_data"]) > 1
                        if is_special:
                            email_subject = "Salary Slips - Bajaj Earths Pvt. Ltd."
                        else:
                            email_subject = "Salary Slip - Bajaj Earths Pvt. Ltd."
                        email_body = """
                                    <html>
                                    <body>
                                    <p>Dear <b>{}</b>,</p>
                                    <p>Please find attached your <b>salary slip</b> for the month of <b>{} {}</b>.</p>
                                    <p>This document includes:</p>
                                    <ul>
                                    <li>Earnings Breakdown</li>
                                    <li>Deductions Summary</li>
                                    <li>Net Salary Details</li>
                                    </ul>
                                    <p>Kindly review the salary slip, and if you have any questions or concerns, please feel free to reach out to the HR department.</p>
                                    <p>Thanks & Regards,</p>
                                    </body>
                                    </html>
                        """.format(employee.get("Name"), full_month, full_year)
                        logging.info("Sending email to {}".format(recipient_email))
                        user_email = session.get('user', {}).get('email')
                        success = send_email_smtp(
                            recipient_email=recipient_email,
                            subject=email_subject,
                            body=email_body,
                            attachment_paths=collected_pdfs,
                            user_email=user_email
                        )
                        if success == "TOKEN_EXPIRED":
                            return jsonify({"error": "TOKEN_EXPIRED"}), 401
                        elif success == "USER_NOT_LOGGED_IN":
                            return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                        elif success == "NO_SMTP_CREDENTIALS":
                            return jsonify({"error": "NO_SMTP_CREDENTIALS", "message": "Email credentials not found. Please check your settings."}), 400
                        elif success == "INVALID_RECIPIENT":
                            return jsonify({"error": "INVALID_RECIPIENT", "message": "Invalid recipient email address."}), 400
                        elif success == "NO_VALID_RECIPIENTS":
                            return jsonify({"error": "NO_VALID_RECIPIENTS", "message": "No valid recipient emails found."}), 400
                        elif success == "SMTP_AUTH_FAILED":
                            return jsonify({"error": "SMTP_AUTH_FAILED", "message": "Email authentication failed. Please check your credentials."}), 400
                        elif success == "SMTP_ERROR":
                            return jsonify({"error": "SMTP_ERROR", "message": "Email service error. Please try again later."}), 500
                        elif success == "EMAIL_SEND_ERROR":
                            return jsonify({"error": "EMAIL_SEND_ERROR", "message": "Failed to send email. Please try again."}), 500
                        elif not success:
                            logging.error("Failed to send email to {}".format(recipient_email))
                            return jsonify({"error": "EMAIL_SEND_FAILED", "message": "Failed to send email. Please try again."}), 500
                        else:
                            logging.info("Email sent to {}".format(recipient_email))
                    else:
                        logging.warning("No email found for {}".format(employee.get("Name")))
                except Exception as e:
                    logging.error("Error sending email: {}".format(str(e)))
                    return jsonify({"error": "EMAIL_ERROR", "message": f"Error sending email: {str(e)}"}), 500

            # Send WhatsApp messages if enabled
            if send_whatsapp:
                try:
                    contact_name = employee.get("Name")
                    whatsapp_number = get_employee_contact(contact_name, contact_employees)
                    if whatsapp_number and collected_pdfs:
                        logging.info(f"Sending WhatsApp notification to {contact_name} for {len(collected_pdfs)} PDF(s)")
                        logging.info(f"PDF files to send: {collected_pdfs}")
                        # Send all collected PDFs in one WhatsApp message
                        success = handle_whatsapp_notification(
                            contact_name=contact_name,
                            full_month=full_month,
                            full_year=full_year,
                            whatsapp_number=whatsapp_number,
                            file_path=collected_pdfs,
                            is_special=len(user_inputs["months_data"]) > 1,
                            months_data=user_inputs["months_data"]
                        )
                        
                        if success == "USER_NOT_LOGGED_IN":
                            return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                        elif success == "WHATSAPP_SERVICE_NOT_READY":
                            # Instead of returning 503, log a warning and continue
                            logging.warning("WhatsApp service is not ready. Salary slips generated but WhatsApp messages could not be sent.")
                            app.logger.warning("WhatsApp service is not ready. Please authenticate WhatsApp first.")
                            # Continue processing - don't return error
                            # return jsonify({"error": "WHATSAPP_SERVICE_NOT_READY", "message": "WhatsApp service is not ready. Please try again later."}), 503
                        elif success == "INVALID_FILE_PATH":
                            return jsonify({"error": "INVALID_FILE_PATH", "message": "Invalid file path for WhatsApp message."}), 400
                        elif success == "INVALID_FILE_PATH_TYPE":
                            return jsonify({"error": "INVALID_FILE_PATH_TYPE", "message": "Invalid file path type for WhatsApp message."}), 400
                        elif success == "NO_VALID_FILES":
                            return jsonify({"error": "NO_VALID_FILES", "message": "No valid files found for WhatsApp message."}), 400
                        elif success == "NO_FILES_FOR_UPLOAD":
                            return jsonify({"error": "NO_FILES_FOR_UPLOAD", "message": "No files available for WhatsApp upload."}), 400
                        elif success == "WHATSAPP_API_ERROR":
                            return jsonify({"error": "WHATSAPP_API_ERROR", "message": "WhatsApp API error. Please try again later."}), 500
                        elif success == "WHATSAPP_CONNECTION_ERROR":
                            return jsonify({"error": "WHATSAPP_CONNECTION_ERROR", "message": "WhatsApp connection error. Please try again later."}), 500
                        elif success == "WHATSAPP_TIMEOUT_ERROR":
                            return jsonify({"error": "WHATSAPP_TIMEOUT_ERROR", "message": "WhatsApp timeout error. Please try again later."}), 500
                        elif success == "WHATSAPP_SEND_ERROR":
                            return jsonify({"error": "WHATSAPP_SEND_ERROR", "message": "Failed to send WhatsApp message. Please try again."}), 500
                        elif not success:
                            logging.error("Failed to send WhatsApp message to {}".format(contact_name))
                            return jsonify({"error": "WHATSAPP_SEND_FAILED", "message": "Failed to send WhatsApp message. Please try again."}), 500
                        else:
                            logging.info("WhatsApp message sent successfully to {}".format(contact_name))
                    else:
                        logging.warning("No WhatsApp number found for {}".format(contact_name))
                except Exception as e:
                    logging.error("Error sending WhatsApp message: {}".format(str(e)))
                    return jsonify({"error": "WHATSAPP_ERROR", "message": f"Error sending WhatsApp message: {str(e)}"}), 500

        # Return results for all processed months
        return jsonify({
            "message": "Processing completed",
            "results": results
        }), 200 if all(r["status"] == "success" for r in results) else 207

    except Exception as e:
        app.logger.error("Error: {}".format(e))
        return jsonify({"error": str(e)}), 500
    finally:
        # Optional cleanup or logging here if needed
        pass

@app.route("/api/generate-salary-slips-batch", methods=["POST"])
def generate_salary_slips_batch():
    try:
        user_inputs = request.json
        app.logger.info("Processing batch salary slips request")

        required_keys = ["sheet_id_salary", "sheet_id_drive", "full_month", "full_year"]
        missing_keys = [key for key in required_keys if not user_inputs.get(key)]
        if missing_keys:
            return jsonify({"error": "Missing parameters: {}".format(', '.join(missing_keys))}), 400

        sheet_id_salary = user_inputs["sheet_id_salary"]
        sheet_id_drive = user_inputs["sheet_id_drive"]
        full_month = user_inputs["full_month"]
        full_year = user_inputs["full_year"]
        sheet_name = full_month[:3]
        send_whatsapp = user_inputs.get("send_whatsapp", False)
        send_email = user_inputs.get("send_email", False)

        # Fetch data
        try:
            salary_data = fetch_google_sheet_data(sheet_id_salary, sheet_name)
            drive_data = fetch_google_sheet_data(sheet_id_salary, "Salary Details")
            email_data = fetch_google_sheet_data(sheet_id_drive, "Onboarding Details")
            contact_data = fetch_google_sheet_data(sheet_id_drive, "Onboarding Details")
        except Exception as e:
            return jsonify({"error": "Error fetching data: {}".format(e)}), 500

        if not all([salary_data, drive_data, email_data, contact_data]):
            return jsonify({"error": "Failed to fetch required data"}), 500

        # Process data
        salary_headers = (salary_data[1])
        employees = salary_data[2:]

        drive_headers = (drive_data[1])
        drive_employees = [dict(zip(drive_headers, row)) for row in drive_data[2:]]

        email_headers = (email_data[1])
        email_employees = [dict(zip(email_headers, row)) for row in email_data[2:]]

        contact_headers = (contact_data[1])
        contact_employees = [dict(zip(contact_headers, row)) for row in contact_data[2:]]

        
       # Generate salary slips for each employee sequentially
        for employee in employees:
            employee_name = employee[4]  # Assuming the employee name is at index 4
            app.logger.info("Processing salary slip for employee: {}".format(employee_name))
            try:
                employee_data = [str(item) if item is not None else '' for item in employee]
                # Get employee code from the employee data
                employee_code_index = next((i for i, header in enumerate(salary_headers) if 'Employee' in header and 'Code' in header), 0)
                employee_identifier = employee[employee_code_index] if employee_code_index < len(employee) else ''
                
                process_salary_slips(
                    template_path=TEMPLATE_PATH,
                    output_dir=OUTPUT_DIR,
                    employee_identifier=employee_identifier,
                    employee_data=employee_data,
                    headers=salary_headers,
                    drive_data=drive_employees,
                    email_employees=email_employees,
                    contact_employees=contact_employees,
                    month=sheet_name,
                    year=str(full_year)[-2:],  # Last two digits of the year
                    full_month=full_month,
                    full_year=full_year,
                    send_whatsapp=send_whatsapp,
                    send_email=send_email
                )
                if send_email:
                    app.logger.info("Sending email to {}".format(employee[5]))  # Assuming email is at index 5
                    user_email = session.get('user', {}).get('email')
                    if not user_email:
                        return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                    
                    # Get recipient email
                    recipient_email = get_employee_email(employee[4], email_employees)  # Assuming name is at index 4
                    if recipient_email:
                        email_subject = "Salary Slip - Bajaj Earths Pvt. Ltd."
                        email_body = f"""
                        <html>
                        <body>
                        <p>Dear <b>{employee[4]}</b>,</p>
                        <p>Please find attached your <b>salary slip</b> for the month of <b>{full_month} {full_year}</b>.</p>
                        <p>This document includes:</p>
                        <ul>
                        <li>Earnings Breakdown</li>
                        <li>Deductions Summary</li>
                        <li>Net Salary Details</li>
                        </ul>
                        <p>Kindly review the salary slip, and if you have any questions or concerns, please feel free to reach out to the HR department.</p>
                        <p>Thanks & Regards,</p>
                        </body>
                        </html>
                        """
                        
                        success = send_email_smtp(
                            recipient_email=recipient_email,
                            subject=email_subject,
                            body=email_body,
                            attachment_paths=[output_pdf] if 'output_pdf' in locals() else [],
                            user_email=user_email
                        )
                        
                        if success == "TOKEN_EXPIRED":
                            return jsonify({"error": "TOKEN_EXPIRED"}), 401
                        elif success == "USER_NOT_LOGGED_IN":
                            return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                        elif success == "NO_SMTP_CREDENTIALS":
                            return jsonify({"error": "NO_SMTP_CREDENTIALS", "message": "Email credentials not found. Please check your settings."}), 400
                        elif success == "INVALID_RECIPIENT":
                            return jsonify({"error": "INVALID_RECIPIENT", "message": "Invalid recipient email address."}), 400
                        elif success == "NO_VALID_RECIPIENTS":
                            return jsonify({"error": "NO_VALID_RECIPIENTS", "message": "No valid recipient emails found."}), 400
                        elif success == "SMTP_AUTH_FAILED":
                            return jsonify({"error": "SMTP_AUTH_FAILED", "message": "Email authentication failed. Please check your credentials."}), 400
                        elif success == "SMTP_ERROR":
                            return jsonify({"error": "SMTP_ERROR", "message": "Email service error. Please try again later."}), 500
                        elif success == "EMAIL_SEND_ERROR":
                            return jsonify({"error": "EMAIL_SEND_ERROR", "message": "Failed to send email. Please try again."}), 500
                        elif not success:
                            app.logger.error("Failed to send email to {}".format(recipient_email))
                            return jsonify({"error": "EMAIL_SEND_FAILED", "message": "Failed to send email. Please try again."}), 500
                    else:
                        app.logger.warning("No email found for {}".format(employee[4]))
                        
                if send_whatsapp:
                    app.logger.info("Sending WhatsApp message to {}".format(employee[6]))  # Assuming phone number is at index 6
                    user_email = session.get('user', {}).get('email')
                    if not user_email:
                        return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                    
                    contact_name = employee[4]  # Assuming name is at index 4
                    whatsapp_number = get_employee_contact(contact_name, contact_employees)
                    if whatsapp_number:
                        success = handle_whatsapp_notification(
                            contact_name=contact_name,
                            full_month=full_month,
                            full_year=full_year,
                            whatsapp_number=whatsapp_number,
                            file_path=[output_pdf] if 'output_pdf' in locals() else [],
                            is_special=False
                        )
                        
                        if success == "USER_NOT_LOGGED_IN":
                            return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                        elif success == "WHATSAPP_SERVICE_NOT_READY":
                            # Instead of returning 503, log a warning and continue
                            logging.warning("WhatsApp service is not ready. Salary slips generated but WhatsApp messages could not be sent.")
                            app.logger.warning("WhatsApp service is not ready. Please authenticate WhatsApp first.")
                            # Continue processing - don't return error
                            # return jsonify({"error": "WHATSAPP_SERVICE_NOT_READY", "message": "WhatsApp service is not ready. Please try again later."}), 503
                        elif success == "INVALID_FILE_PATH":
                            return jsonify({"error": "INVALID_FILE_PATH", "message": "Invalid file path for WhatsApp message."}), 400
                        elif success == "INVALID_FILE_PATH_TYPE":
                            return jsonify({"error": "INVALID_FILE_PATH_TYPE", "message": "Invalid file path type for WhatsApp message."}), 400
                        elif success == "NO_VALID_FILES":
                            return jsonify({"error": "NO_VALID_FILES", "message": "No valid files found for WhatsApp message."}), 400
                        elif success == "NO_FILES_FOR_UPLOAD":
                            return jsonify({"error": "NO_FILES_FOR_UPLOAD", "message": "No files available for WhatsApp upload."}), 400
                        elif success == "WHATSAPP_API_ERROR":
                            return jsonify({"error": "WHATSAPP_API_ERROR", "message": "WhatsApp API error. Please try again later."}), 500
                        elif success == "WHATSAPP_CONNECTION_ERROR":
                            return jsonify({"error": "WHATSAPP_CONNECTION_ERROR", "message": "WhatsApp connection error. Please try again later."}), 500
                        elif success == "WHATSAPP_TIMEOUT_ERROR":
                            return jsonify({"error": "WHATSAPP_TIMEOUT_ERROR", "message": "WhatsApp timeout error. Please try again later."}), 500
                        elif success == "WHATSAPP_SEND_ERROR":
                            return jsonify({"error": "WHATSAPP_SEND_ERROR", "message": "Failed to send WhatsApp message. Please try again."}), 500
                        elif not success:
                            app.logger.error("Failed to send WhatsApp message to {}".format(contact_name))
                            return jsonify({"error": "WHATSAPP_SEND_FAILED", "message": "Failed to send WhatsApp message. Please try again."}), 500
                        else:
                            app.logger.info("WhatsApp message sent successfully to {}".format(contact_name))
            except Exception as e:
                error_msg = "Error processing salary slip for employee {}: {}".format(employee_name, e)
                app.logger.error(error_msg)
                return jsonify({"error": error_msg}), 500

        return jsonify({"message": "Batch salary slips generated successfully"}), 200

    except Exception as e:
        app.logger.error("Error: {}".format(e))
        return jsonify({"error": str(e)}), 500
    finally:
        # Optional cleanup or logging here if needed
        pass

@app.route("/api/get-logs", methods=["GET"])
def get_logs():
    try:
        with open(LOG_FILE_PATH, "r") as log_file:
            logs = log_file.read()
        return logs, 200, {'Content-Type': 'text/plain'}
    except Exception as e:
        app.logger.error("Error reading logs: {}".format(e))
        return jsonify({"error": "Failed to read logs"}), 500

@app.route("/api/", methods=["GET"])
def home():
    return jsonify({"message": "Welcome to the Salary Slip Automation API!"}), 200

@app.route("/api/user/oauth-status", methods=["GET"])
def get_user_oauth_status():
    """Get user's OAuth token status and permissions"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        user_id = session.get('user', {}).get('id')
        if not user_id:
            return jsonify({"error": "User ID not found"}), 400
        
        oauth_tokens = get_user_oauth_tokens(user_id)
        if not oauth_tokens:
            return jsonify({
                "success": True,
                "has_oauth": False,
                "message": "No OAuth tokens found"
            }), 200
        
        return jsonify({
            "success": True,
            "has_oauth": True,
            "has_sheets_access": oauth_tokens.get('has_sheets_access', False),
            "has_drive_access": oauth_tokens.get('has_drive_access', False),
            "has_gmail_access": oauth_tokens.get('has_gmail_access', False),
            "granted_scopes": oauth_tokens.get('granted_scopes', [])
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting user OAuth status: {e}")
        return jsonify({"error": "Failed to get OAuth status"}), 500

@app.route("/api/update_permissions", methods=["POST"])
def update_permissions():
    try:
        # Check if user is logged in
        current_user = session.get('user')
        if not current_user:
            return jsonify({"error": "User not authenticated"}), 401

        data = request.json
        user_id = data.get('user_id')
        permission_metadata = data.get('permission_metadata')

        if not user_id:
            return jsonify({"error": "User ID is required"}), 400

        if not permission_metadata:
            return jsonify({"error": "Permission metadata is required"}), 400

        # Get target user
        target_user = get_user_by_id(user_id)
        if not target_user:
            return jsonify({"error": "User not found"}), 404

        current_user_role = current_user.get('role')

        # Admin can update permissions for any user
        if current_user_role == 'admin':
            # Update user with permission_metadata only
            from Utils.firebase_utils import update_user_complete_rbac
            
            # Use the complete RBAC update function
            update_user_complete_rbac(user_id, permission_metadata)
            
            return jsonify({"message": "Permissions updated successfully"}), 200
        
        # Regular users cannot update permissions
        else:
            return jsonify({"error": "Insufficient permissions to update user permissions"}), 403

    except Exception as e:
        logger.error("Error updating permissions: {}".format(e))
        return jsonify({"error": str(e)}), 500

@app.route("/api/send-reports", methods=["POST"])
def generate_report():
    try:
        user_id = session.get('user', {}).get('email')
        if not user_id:
            logger.error("No user_id found in session. User must be logged in to send reports.")
            return jsonify({"error": "User not authenticated"}), 401
        
        # Get form data
        template_files = request.files.getlist('template_files')
        attachment_files = request.files.getlist('attachment_files')
        file_sequence = json.loads(request.form.get('file_sequence', '{}'))
        sheet_id = request.form.get('sheet_id')
        sheet_name = request.form.get('sheet_name')
        send_whatsapp = request.form.get('send_whatsapp') == 'true'
        send_email = request.form.get('send_email') == 'true'
        mail_subject = request.form.get('mail_subject')
        use_template_as_caption = request.form.get('use_template_as_caption') == 'true'

        # Import the new function from process_utils
        from Utils.process_utils import process_general_reports

        # Call the new function that handles all the heavy processing
        result = process_general_reports(
            template_files=template_files,
            attachment_files=attachment_files,
            file_sequence=file_sequence,
            sheet_id=sheet_id,
            sheet_name=sheet_name,
            send_whatsapp=send_whatsapp,
            send_email=send_email,
            mail_subject=mail_subject,
            use_template_as_caption=use_template_as_caption,
            user_id=user_id,
            output_dir=OUTPUT_DIR,
            logger=logger,
            send_email_smtp=send_email_smtp,
            send_whatsapp_message=send_whatsapp_message,
            validate_sheet_id_func=validate_sheet_id,
            prepare_file_paths_func=prepare_file_paths,
            fetch_google_sheet_data_func=fetch_google_sheet_data,
            process_template_func=process_template,
            send_log_report_to_user_func=send_log_report_to_user
        )

        # Check if there were any errors
        if not result["success"]:
            if result["errors"]:
                return jsonify({"error": result["errors"][0]}), 400
            else:
                return jsonify({"error": "Unknown error occurred"}), 500

        # Return success response
        return jsonify({
            "message": result["message"],
            "generated_files": result["generated_files"],
            "notifications_sent": result["notifications_sent"],
            "delivery_stats": result["delivery_stats"]
        }), 200

    except Exception as e:
        logger.error("Error generating reports: {}".format(e))
        return jsonify({"error": str(e)}), 500
    finally:
        # Optional cleanup or logging here if needed
        pass

@app.route("/api/retry-reports", methods=["POST"])
def retry_reports():
    """Retry sending reports after token refresh"""
    try:
        user_id = session.get('user', {}).get('email')
        if not user_id:
            logger.error("No user_id found in session. User must be logged in to send reports.")
            return jsonify({"error": "User not authenticated"}), 401
        
        # Get the original request data from the request
        data = request.json
        original_request_data = data.get('original_request_data')
        
        if not original_request_data:
            return jsonify({"error": "Original request data is required"}), 400
        
        # Extract the data needed for report generation
        template_files_data = original_request_data.get('template_files_data', [])
        attachment_files_data = original_request_data.get('attachment_files_data', [])
        file_sequence = original_request_data.get('file_sequence', {})
        sheet_id = original_request_data.get('sheet_id')
        sheet_name = original_request_data.get('sheet_name')
        send_whatsapp = original_request_data.get('send_whatsapp', False)
        send_email = original_request_data.get('send_email', True)
        mail_subject = original_request_data.get('mail_subject', '')
        
        if not sheet_id or not sheet_name:
            return jsonify({"error": "Sheet ID and name are required"}), 400
        
        # Validate sheet ID format
        if not validate_sheet_id(sheet_id):
            return jsonify({"error": "Invalid Google Sheet ID format"}), 400
        
        # Get user-specific temporary directory for attachments
        from Utils.temp_manager import get_user_temp_dir
        temp_dir = get_user_temp_dir(user_id, OUTPUT_DIR)
        
        # Reconstruct attachment files from base64 data
        attachment_paths = []
        for attachment_data in attachment_files_data:
            try:
                file_name = attachment_data.get('name')
                file_content = base64.b64decode(attachment_data.get('content'))
                file_path = os.path.join(temp_dir, file_name)
                
                with open(file_path, 'wb') as f:
                    f.write(file_content)
                attachment_paths.append(file_path)
            except Exception as e:
                logger.error(f"Error reconstructing attachment {file_name}: {e}")
                continue
        
        try:
            # Fetch data from Google Sheet
            sheet_data = fetch_google_sheet_data(sheet_id, sheet_name)
            if not sheet_data or len(sheet_data) < 2:
                return jsonify({"error": "No data found in the Google Sheet"}), 400
        except Exception as e:
            logger.error("Error fetching Google Sheet data: {}".format(e))
            return jsonify({"error": "Failed to fetch data from Google Sheet"}), 500
        
        # Process headers and data
        headers = sheet_data[0]
        data_rows = sheet_data[1:]
        
        # Create output directory if it doesn't exist
        output_dir = os.path.join(OUTPUT_DIR, "reports")
        os.makedirs(output_dir, exist_ok=True)
        
        # Process each template file
        generated_files = []
        for template_data in template_files_data:
            try:
                file_name = template_data.get('name')
                file_content = base64.b64decode(template_data.get('content'))
                temp_template_path = os.path.join(output_dir, f"temp_{file_name}")
                
                with open(temp_template_path, 'wb') as f:
                    f.write(file_content)
                
                # Read template content for messages
                try:
                    doc = Document(temp_template_path)
                    template_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                except Exception as e:
                    logger.error("Error reading template content: {}".format(e))
                    continue
                
                if send_whatsapp:
                    pass
                
                # Process each row of data
                for row in data_rows:
                    try:
                        # Create data dictionary from headers and row
                        data_dict = dict(zip(headers, row))
                        
                        recipient_name = data_dict.get('Name', 'unknown')
                        
                        # Generate report for this row
                        output_filename = "report_{}.docx".format(recipient_name)
                        output_path = os.path.join(output_dir, output_filename)
                        
                        # Process the template with data
                        process_template(temp_template_path, output_path, data_dict)
                        generated_files.append(output_path)
                        
                        # Process template content for messages
                        message_content = template_content
                        email_content = template_content
                        
                        # Replace placeholders in message content
                        for key, value in data_dict.items():
                            placeholder = "{{{}}}".format(key)
                            message_content = message_content.replace(placeholder, str(value))
                            email_content = email_content.replace(placeholder, str(value))
                            mail_subject = mail_subject.replace(placeholder, str(value))
                        
                        # Get contact details from Google Sheet data
                        recipient_email = data_dict.get('Email ID - To')
                        cc_email = data_dict.get('Email ID - CC', '')
                        bcc_email = data_dict.get('Email ID - BCC', '')
                        
                        # Helper to split emails by comma or newline and join as comma-separated string
                        def clean_emails(email_str):
                            if not email_str:
                                return None
                            emails = [e.strip() for e in re.split(r'[\n,]+', email_str) if e.strip()]
                            return ','.join(emails) if emails else None
                        
                        recipient_email = clean_emails(recipient_email)
                        cc_email = clean_emails(cc_email)
                        bcc_email = clean_emails(bcc_email)
                        
                        country_code = data_dict.get('Country Code', '').strip()
                        phone_no = data_dict.get('Contact No.', '').strip()
                        # Format phone number properly: remove spaces and combine country code + number
                        recipient_phone = f"{country_code}{phone_no}".replace(' ', '')
                        
                        if send_whatsapp:
                            # Validate phone number components
                            if not country_code or not country_code.strip():
                                logger.warning("Skipping WhatsApp message for {}: Missing Country Code".format(recipient_name))
                                continue
                            
                            if not phone_no or not phone_no.strip():
                                logger.warning("Skipping WhatsApp message for {}: Missing Contact No.".format(recipient_name))
                                continue
                            
                            if not recipient_phone or len(recipient_phone.replace(' ', '')) < 4:
                                logger.warning("Skipping WhatsApp message for {}: Invalid phone number format (Country Code: {}, Contact No.: {})".format(
                                    recipient_name, country_code, phone_no))
                                continue
                            
                            logger.info("Valid phone number for {}: Country Code '{}', Contact No. '{}' -> Formatted: '{}'".format(
                                recipient_name, country_code, phone_no, recipient_phone))
                            
                            try:
                                # Send WhatsApp message with attachments
                                success = send_whatsapp_message(
                                    contact_name=recipient_name,
                                    message=message_content,  # Use processed template content
                                    file_paths=attachment_paths,
                                    file_sequence=file_sequence,
                                    whatsapp_number=recipient_phone,
                                    process_name="report"
                                )
                                
                                if success == "USER_NOT_LOGGED_IN":
                                    return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                                elif success == "WHATSAPP_SERVICE_NOT_READY":
                                    return jsonify({"error": "WHATSAPP_SERVICE_NOT_READY", "message": "WhatsApp service is not ready. Please try again later."}), 503
                                elif success == "INVALID_FILE_PATH":
                                    return jsonify({"error": "INVALID_FILE_PATH", "message": "Invalid file path for WhatsApp message."}), 400
                                elif success == "INVALID_FILE_PATH_TYPE":
                                    return jsonify({"error": "INVALID_FILE_PATH_TYPE", "message": "Invalid file path type for WhatsApp message."}), 400
                                elif success == "NO_VALID_FILES":
                                    return jsonify({"error": "NO_VALID_FILES", "message": "No valid files found for WhatsApp message."}), 400
                                elif success == "NO_FILES_FOR_UPLOAD":
                                    return jsonify({"error": "NO_FILES_FOR_UPLOAD", "message": "No files available for WhatsApp upload."}), 400
                                elif success == "WHATSAPP_API_ERROR":
                                    return jsonify({"error": "WHATSAPP_API_ERROR", "message": "WhatsApp API error. Please try again later."}), 500
                                elif success == "WHATSAPP_CONNECTION_ERROR":
                                    return jsonify({"error": "WHATSAPP_CONNECTION_ERROR", "message": "WhatsApp connection error. Please try again later."}), 500
                                elif success == "WHATSAPP_TIMEOUT_ERROR":
                                    return jsonify({"error": "WHATSAPP_TIMEOUT_ERROR", "message": "WhatsApp timeout error. Please try again later."}), 500
                                elif success == "WHATSAPP_SEND_ERROR":
                                    return jsonify({"error": "WHATSAPP_SEND_ERROR", "message": "Failed to send WhatsApp message. Please try again."}), 500
                                elif not success:
                                    logger.error("Failed to send WhatsApp message to {}".format(recipient_phone))
                                    return jsonify({"error": "WHATSAPP_SEND_FAILED", "message": "Failed to send WhatsApp message. Please try again."}), 500

                            except Exception as e:
                                logger.error("Error sending WhatsApp message to {}: {}".format(recipient_phone, e))
                        
                        # Handle email notifications
                        if send_email:
                            if not recipient_email:
                                logger.warning("No email found for recipient: {}".format(recipient_name))
                                continue
                            try:
                                email_subject = mail_subject
                                email_body = """
                                <html>
                                <body>
                                {} 
                                </body>
                                </html>
                                """.format(email_content.replace('\n', '<br>'))
                                user_email = session.get('user', {}).get('email')
                                success = send_email_smtp(
                                    recipient_email=recipient_email,
                                    subject=email_subject,
                                    body=email_body,
                                    attachment_paths=attachment_paths,
                                    user_email=user_email,
                                    cc=cc_email,
                                    bcc=bcc_email
                                )
                                if success == "TOKEN_EXPIRED":
                                    return jsonify({"error": "TOKEN_EXPIRED"}), 401
                                elif success == "USER_NOT_LOGGED_IN":
                                    return jsonify({"error": "USER_NOT_LOGGED_IN", "message": "User session expired. Please log in again."}), 401
                                elif success == "NO_SMTP_CREDENTIALS":
                                    return jsonify({"error": "NO_SMTP_CREDENTIALS", "message": "Email credentials not found. Please check your settings."}), 400
                                elif success == "INVALID_RECIPIENT":
                                    return jsonify({"error": "INVALID_RECIPIENT", "message": "Invalid recipient email address."}), 400
                                elif success == "NO_VALID_RECIPIENTS":
                                    return jsonify({"error": "NO_VALID_RECIPIENTS", "message": "No valid recipient emails found."}), 400
                                elif success == "SMTP_AUTH_FAILED":
                                    return jsonify({"error": "SMTP_AUTH_FAILED", "message": "Email authentication failed. Please check your credentials."}), 400
                                elif success == "SMTP_ERROR":
                                    return jsonify({"error": "SMTP_ERROR", "message": "Email service error. Please try again later."}), 500
                                elif success == "EMAIL_SEND_ERROR":
                                    return jsonify({"error": "EMAIL_SEND_ERROR", "message": "Failed to send email. Please try again."}), 500
                                elif not success:
                                    logger.error("Failed to send email to {}".format(recipient_email))
                                    return jsonify({"error": "EMAIL_SEND_FAILED", "message": "Failed to send email. Please try again."}), 500
                            except Exception as e:
                                logger.error("Error sending email: {}".format(str(e)))
                                
                    except Exception as e:
                        logger.error("Error processing row: {}".format(e))
                        continue
                
                # Clean up temporary template
                os.remove(temp_template_path)
                
            except Exception as e:
                logger.error(f"Error processing template {file_name}: {e}")
                continue
        
        # Clean up user-specific temporary directory
        from Utils.temp_manager import cleanup_user_temp_dir
        cleanup_user_temp_dir(user_id, OUTPUT_DIR)
        
        return jsonify({
            "success": True,
            "message": "Reports generated and sent successfully!",
            "generated_files": len(generated_files)
        }), 200
        
    except Exception as e:
        logger.error("Error retrying reports: {}".format(e))
        return jsonify({"error": str(e)}), 500
    
@app.route("/api/reactor_reports", methods=["POST"])
def reactor_report():
    try:
        user_id = session.get('user', {}).get('email')
        if not user_id:
            logger.error("No user_id found in session. User must be logged in to send reports.")
            return jsonify({"error": "User not authenticated"}), 401
        # Get form data
        send_email = request.form.get('send_email') == 'true'
        send_whatsapp = request.form.get('send_whatsapp') == 'true'
        date = request.form.get('date')
        if not date:
            return jsonify({"error": "Date is required"}), 400
        # Fetch sheet IDs from Google Sheet (mapping sheet)
        try:
            reactor_reports_sheet_id = "1XOLQvy6j7syAlOKpQ3J2o6DcgiSZsSO1xxWlWih_QOY"
            sheet_id_mapping_data = fetch_google_sheet_data(reactor_reports_sheet_id, 'Sheet_ID_Reactor')
            sheet_recipients_data = fetch_google_sheet_data(reactor_reports_sheet_id, 'Recipients')
            table_range_data = fetch_google_sheet_data(reactor_reports_sheet_id, 'Table Ranges')
        except Exception as e:
            logger.error(f"Error fetching sheet IDs from Google Sheet: {e}")
            return jsonify({"error": "Failed to fetch sheet IDs from Google Sheet"}), 500
        # Get user-specific temporary directory for processing
        from Utils.temp_manager import get_user_temp_dir
        temp_dir = get_user_temp_dir(user_id, OUTPUT_DIR)
        template_path = os.path.join(os.path.dirname(__file__), "reactorreportformat.docx")
        if not os.path.exists(template_path):
            return jsonify({"error": "Reactor report template not found"}), 500
        gspread_client = gspread.authorize(creds)
        # Call the new utility function
        result = process_reactor_reports(
            sheet_id_mapping_data=sheet_id_mapping_data,
            sheet_recipients_data=sheet_recipients_data,
            table_range_data=table_range_data,
            input_date=date,
            user_id=user_id,
            send_email=send_email,
            send_whatsapp=send_whatsapp,
            template_path=template_path,
            output_dir=temp_dir,
            gspread_client=gspread_client,
            logger=logger,
            send_email_smtp=send_email_smtp
        )
        # Clean up user-specific temporary directory
        from Utils.temp_manager import cleanup_user_temp_dir
        cleanup_user_temp_dir(user_id, OUTPUT_DIR)
        if 'error' in result:
            return jsonify({"error": result['error']}), 400
        return jsonify(result), 200
    except Exception as e:
        logger.error(f"Error generating reactor reports: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        pass

@app.route("/api/kr_reactor_reports", methods=["POST"])
def kr_reactor_report():
    """New endpoint specifically for KR_ReactorReports.jsx with OAuth email support"""
    try:
        user_id = session.get('user', {}).get('email')
        if not user_id:
            logger.error("No user_id found in session. User must be logged in to send reports.")
            return jsonify({"error": "User not authenticated"}), 401
        
        # Get form data
        send_email = request.form.get('send_email') == 'true'
        send_whatsapp = request.form.get('send_whatsapp') == 'true'
        date = request.form.get('date')
        process_name = request.form.get('process_name', 'reactor-report')  # Default to 'reactor-report'
        
        # Get Google tokens if provided by frontend
        google_access_token = request.form.get('google_access_token')
        google_refresh_token = request.form.get('google_refresh_token')
        
        if not date:
            return jsonify({"error": "Date is required"}), 400
        
        # Fetch sheet IDs from Google Sheet (mapping sheet)
        try:
            reactor_reports_sheet_id = "1XOLQvy6j7syAlOKpQ3J2o6DcgiSZsSO1xxWlWih_QOY"
            sheet_id_mapping_data = fetch_google_sheet_data(reactor_reports_sheet_id, 'Sheet_ID_Reactor')
            sheet_recipients_data = fetch_google_sheet_data(reactor_reports_sheet_id, 'Recipients')
            table_range_data = fetch_google_sheet_data(reactor_reports_sheet_id, 'Table Ranges')
        except Exception as e:
            logger.error(f"Error fetching sheet IDs from Google Sheet: {e}")
            return jsonify({"error": "Failed to fetch sheet IDs from Google Sheet"}), 500
        
        # Get user-specific temporary directory for processing
        from Utils.temp_manager import get_user_temp_dir
        temp_dir = get_user_temp_dir(user_id, OUTPUT_DIR)
        template_path = os.path.join(os.path.dirname(__file__), "reactorreportformat.docx")
        
        if not os.path.exists(template_path):
            return jsonify({"error": "Reactor report template not found"}), 500
        
        gspread_client = gspread.authorize(creds)
        
        # Call the new utility function with OAuth email support
        result = process_reactor_reports(
            sheet_id_mapping_data=sheet_id_mapping_data,
            sheet_recipients_data=sheet_recipients_data,
            table_range_data=table_range_data,
            input_date=date,
            user_id=user_id,
            send_email=send_email,
            send_whatsapp=send_whatsapp,
            template_path=template_path,
            output_dir=temp_dir,
            gspread_client=gspread_client,
            logger=logger,
            send_email_smtp=send_email_smtp,  # This will use OAuth email function
            process_name=process_name,  # Pass the process name for OAuth email selection
            google_access_token=google_access_token,  # Pass Google tokens if available
            google_refresh_token=google_refresh_token
        )
        
        # Clean up user-specific temporary directory
        from Utils.temp_manager import cleanup_user_temp_dir
        cleanup_user_temp_dir(user_id, OUTPUT_DIR)
        
        if 'error' in result:
            return jsonify({"error": result['error']}), 400
        
        return jsonify(result), 200
        
    except Exception as e:
        logger.error(f"Error generating KR reactor reports: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        pass

@app.route("/api/update_app_password", methods=["POST"])
def update_app_password():
    try:
        # Check if user is logged in
        current_user = session.get('user')
        if not current_user:
            return jsonify({"error": "User not authenticated"}), 401

        data = request.json
        user_id = data.get('user_id')
        app_password = data.get('appPassword') or data.get('app_password')
        
        if not user_id or not app_password:
            return jsonify({"error": "User ID and app password are required"}), 400

        # Get target user
        target_user = get_user_by_id(user_id)
        if not target_user:
            return jsonify({"error": "User not found"}), 404

        current_user_role = current_user.get('role')
        target_user_role = target_user.get('role')

        # Admin can update any user's app password
        if current_user_role == 'admin':
            update_user_app_password(user_id, app_password)
            return jsonify({"message": "App password updated successfully"}), 200
        
        # Regular users cannot update other users' app passwords
        else:
            return jsonify({"error": "Insufficient permissions to update app passwords"}), 403

    except Exception as e:
        logger.error("Error updating app password: {}".format(e))
        return jsonify({"error": str(e)}), 500

@app.route("/api/update_user", methods=["POST"])
def update_user_endpoint():
    try:
        # Check if user is logged in
        current_user = session.get('user')
        if not current_user:
            return jsonify({"error": "User not authenticated"}), 401

        data = request.json
        user_id = data.get('user_id')
        permissions = data.get('permissions')
        
        if not user_id:
            return jsonify({"error": "User ID is required"}), 400

        # Get target user
        target_user = get_user_by_id(user_id)
        if not target_user:
            return jsonify({"error": "User not found"}), 404

        current_user_role = current_user.get('role')
        target_user_role = target_user.get('role')

        # Admin can update any user
        if current_user_role == 'admin':
                # Prepare update data
                update_data = {}
                if permissions is not None:
                    update_data['permissions'] = permissions
                    
                if not update_data:
                    return jsonify({"error": "No valid update data provided"}), 400
                
                # Update user permissions
                from Utils.firebase_utils import update_user_comprehensive_permissions
                
                if 'permissions' in update_data:
                    update_user_comprehensive_permissions(user_id, {
                        'permissions': update_data['permissions']
                    })
                
                return jsonify({"message": "User updated successfully"}), 200
        
        # Regular users cannot update other users
        else:
            return jsonify({"error": "Insufficient permissions to update users"}), 403
        
    except Exception as e:
        logger.error("Error updating user: {}".format(e))
        return jsonify({"error": str(e)}), 500

def validate_sheet_id(sheet_id):
    """Validate Google Sheet ID format"""
    return bool(sheet_id and len(sheet_id) == 44)

def process_template(template_path, output_path, data_dict):
    """Process the template file with the provided data"""
    try:
        # Load the template
        doc = Document(template_path)
        
        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            for key, value in data_dict.items():
                placeholder = "{{{}}}".format(key)
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Save the processed document
        doc.save(output_path)
        return True
    except Exception as e:
        logger.error("Error processing template: {}".format(e))
        return False

def get_user_emails(user_id):
    """Get recipient emails for reactor reports from user settings"""
    try:
        # Get user document from Firebase
        user_doc = db.collection('users').document(user_id).get()
        if user_doc.exists:
            user_data = user_doc.to_dict()
            # Check if user has configured reactor report recipients
            reactor_recipients = user_data.get('reactor_report_recipients', [])
            if reactor_recipients:
                return reactor_recipients
        # Fallback to user's own email
        return [user_id]
    except Exception as e:
        logger.error(f"Error getting user emails: {e}")
        return [user_id]

def get_user_phones(user_id):
    """Get recipient phone numbers for reactor reports from user settings"""
    try:
        # Get user document from Firebase
        user_doc = db.collection('users').document(user_id).get()
        if user_doc.exists:
            user_data = user_doc.to_dict()
            # Check if user has configured reactor report phone recipients
            reactor_phones = user_data.get('reactor_report_phones', [])
            if reactor_phones:
                return reactor_phones
        # Return empty list if no phones configured
        return []
    except Exception as e:
        logger.error(f"Error getting user phones: {e}")
        return []

# Error handlers
@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "API endpoint not found"}), 404

@app.errorhandler(Exception)
def handle_exception(e):
    logger.error("Unhandled exception: {}".format(str(e)), exc_info=True)
    return jsonify({"error": "Internal server error", "details": str(e)}), 500

# Health check endpoint for Render
@app.route('/healthz')
def health_check():
    try:
        # Check Firebase connection by getting users
        users = firebase_get_all_users()
        return jsonify({"status": "healthy", "database": "connected"}), 200
    except Exception as e:
        logger.error("Health check failed: {}".format(e))
        return jsonify({"status": "unhealthy", "error": str(e)}), 500

def ensure_directories():
    """Ensure all required directories exist"""
    directories = [
        OUTPUT_DIR,
        os.path.dirname(LOG_FILE_PATH)
    ]
    for directory in directories:
        os.makedirs(directory, exist_ok=True)

@app.route("/api/auth/logout", methods=["POST"])
def logout():
    try:
        # Clear session
        session.clear()
        
        return jsonify({"message": "Logged out successfully"}), 200
    except Exception as e:
        logger.error("Error during logout: {}".format(e))
        return jsonify({"error": str(e)}), 500

@app.route("/api/whatsapp-login", methods=["POST"])
def whatsapp_login():
    """Simple WhatsApp login endpoint - no complex threading"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        user_email = session.get('user', {}).get('email')
        if not user_email:
            return jsonify({"error": "No user email found"}), 400
        
        logging.info(f"Starting WhatsApp login for user: {user_email}")
        
        # Create WhatsApp client directly
        client = WhatsAppNodeClient(WHATSAPP_NODE_SERVICE_URL, user_email)
        
        # Trigger login and get QR code
        result = client.trigger_login()
        logging.info(f"WhatsApp login result: {result}")
        
        # Return clean response
        return jsonify({
            "qr": result.get("qr", ""),
            "authenticated": result.get("authenticated", False),
            "message": result.get("message", ""),
            "userInfo": result.get("userInfo")
        })
        
    except Exception as e:
        logging.error(f"Error in WhatsApp login: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/whatsapp-phone-login", methods=["POST"])
def whatsapp_phone_login():
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        user_email = session.get('user', {}).get('email')
        data = request.get_json()
        phone_number = data.get('phoneNumber')
        code = data.get('code')
        
        if not phone_number:
            return jsonify({"error": "Phone number is required"}), 400
        
        # Use thread pool for WhatsApp phone authentication
        future = whatsapp_executor.submit(whatsapp_phone_login_worker, user_email, phone_number, code)
        
        try:
            result = future.result(timeout=60)  # 60 second timeout for phone login
            return jsonify(result)
        except concurrent.futures.TimeoutError:
            return jsonify({"error": "WhatsApp phone authentication timeout"}), 408
        except Exception as e:
            return jsonify({"error": str(e)}), 500
            
    except Exception as e:
        logging.error(f"Error in WhatsApp phone login: {e}")
        return jsonify({"error": str(e)}), 500

# Removed complex worker functions - using direct client calls instead

@app.route("/api/whatsapp-status", methods=["GET"])
def whatsapp_status():
    """Simple WhatsApp status endpoint - no complex threading"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        user_email = session.get('user', {}).get('email')
        if not user_email:
            return jsonify({"error": "No user email found"}), 400
        
        logging.info(f"Checking WhatsApp status for user: {user_email}")
        
        # Create WhatsApp client directly
        client = WhatsAppNodeClient(WHATSAPP_NODE_SERVICE_URL, user_email)
        
        # Get status
        result = client.get_status()
        logging.info(f"WhatsApp status result: {result}")
        
        # Return clean response
        return jsonify({
            "isReady": result.get("isReady", False),
            "status": result.get("status", "unavailable"),
            "authenticated": result.get("authenticated", False),
            "userInfo": result.get("userInfo")
        })
        
    except Exception as e:
        logging.error(f"Error in WhatsApp status: {e}")
        return jsonify({"error": str(e)}), 500

# Removed complex worker functions - using direct client calls instead

@app.route("/api/whatsapp-force-new-session", methods=["POST"])
def whatsapp_force_new_session():
    """Simple WhatsApp force new session endpoint - no complex threading"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        user_email = session.get('user', {}).get('email')
        if not user_email:
            return jsonify({"error": "No user email found"}), 400
        
        logging.info(f"Forcing new WhatsApp session for user: {user_email}")
        
        # Create WhatsApp client directly
        client = WhatsAppNodeClient(WHATSAPP_NODE_SERVICE_URL, user_email)
        
        # Force new session
        success = client.force_new_session()
        logging.info(f"WhatsApp force new session result: {success}")
        
        if success:
            return jsonify({"success": True, "message": "Session cleared successfully"})
        else:
            return jsonify({"success": False, "message": "Failed to clear session"})
        
    except Exception as e:
        logging.error(f"Error in WhatsApp force new session: {e}")
        return jsonify({"error": str(e)}), 500

# Removed complex worker functions - using direct client calls instead

@app.route("/api/whatsapp-logout", methods=["POST"])
def whatsapp_logout():
    """Simple WhatsApp logout endpoint - no complex threading"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        user_email = session.get('user', {}).get('email')
        if not user_email:
            return jsonify({"error": "No user email found"}), 400
        
        logging.info(f"Logging out WhatsApp for user: {user_email}")
        
        # Create WhatsApp client directly
        client = WhatsAppNodeClient(WHATSAPP_NODE_SERVICE_URL, user_email)
        
        # Logout
        success = client.logout()
        logging.info(f"WhatsApp logout result: {success}")
        
        return jsonify({"success": bool(success)}), 200 if success else 500
        
    except Exception as e:
        logging.error(f"Error logging out WhatsApp: {str(e)}")
        return jsonify({"success": False}), 500

@app.route("/api/test-session", methods=["GET"])
def test_session():
    """Test endpoint to check if session is working properly"""
    try:
        logger.info(f"Test session endpoint called")
        logger.info(f"Request cookies: {request.cookies}")
        logger.info(f"Request headers: {dict(request.headers)}")
        logger.info(f"Session contents: {session}")
        logger.info(f"User in session: {session.get('user')}")
        
        if 'user' in session:
            return jsonify({
                "success": True,
                "session_working": True,
                "user": session['user'],
                "message": "Session is working properly"
            }), 200
        else:
            return jsonify({
                "success": False,
                "session_working": False,
                "message": "No user found in session"
            }), 401
    except Exception as e:
        logger.error(f"Error in test session endpoint: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/get_material_data", methods=["GET"])
def get_material_data():
    """Get material data from Firebase for dropdown population"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        # Get factory parameter (optional, defaults to KR for now)
        factory = request.args.get('factory', 'KR')
        
        from Utils.firebase_utils import get_material_data_by_factory
        
        # Fetch material data for the specified factory
        material_data = get_material_data_by_factory(factory)
        
        if not material_data:
            return jsonify({
                "success": True,
                "data": {},
                "message": f"No material data found for factory {factory}"
            }), 200
        
        return jsonify({
            "success": True,
            "data": material_data,
            "factory": factory
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching material data: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/get_authority_list", methods=["GET"])
def get_authority_list():
    """Get authority list from Google Sheets for dropdown population"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        # Get parameters
        factory = request.args.get('factory', 'KR')
        sheet_name = request.args.get('sheet_name', 'Authority List')
        sheet_id = request.args.get('sheet_id')
        
        if not sheet_id:
            return jsonify({
                "success": False,
                "error": "Sheet ID is required"
            }), 400
        
        # Fetch authority list from Google Sheets
        authority_data = fetch_google_sheet_data(sheet_id, sheet_name)
        
        if not authority_data or len(authority_data) < 2:
            return jsonify({
                "success": True,
                "data": [],
                "message": f"No authority data found in sheet {sheet_name}"
            }), 200
        
        # Extract authority names from the data (assuming first column contains names)
        # Skip header row
        authority_names = []
        for row in authority_data[1:]:
            if row and row[0].strip():  # Check if first column has data
                authority_names.append(row[0].strip())
        
        return jsonify({
            "success": True,
            "data": authority_names,
            "factory": factory,
            "sheet_name": sheet_name
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching authority list: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/get_recipients_list", methods=["GET"])
def get_recipients_list():
    """Get recipients list from Google Sheets for notification population"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        # Get parameters
        factory = request.args.get('factory', 'KR')
        sheet_name = request.args.get('sheet_name', 'Recipents List')
        sheet_id = request.args.get('sheet_id')
        
        if not sheet_id:
            return jsonify({
                "success": False,
                "error": "Sheet ID is required"
            }), 400
        
        # Fetch recipients list from Google Sheets
        recipients_data = fetch_google_sheet_data(sheet_id, sheet_name)
        
        if not recipients_data or len(recipients_data) < 3:
            return jsonify({
                "success": True,
                "data": [],
                "message": f"No recipients data found in sheet {sheet_name}. Need at least 3 rows (headers at row 2, data from row 3)"
            }), 200
        
        # Headers are in Row 2 (index 1), data starts from Row 3 (index 2)
        headers = [h.strip() for h in recipients_data[1]]  # Row 2 contains headers
        
        # Expected headers: "Name, Country Code, Contact No., Email ID - To"
        expected_headers = ['Name', 'Country Code', 'Contact No.', 'Email ID - To']
        logger.info(f"Found headers in recipients sheet: {headers}")
        
        # Build recipients list
        recipients = []
        for row in recipients_data[2:]:  # Process data rows starting from Row 3 (index 2)
            if row and len(row) > 0:
                recipient = {}
                for i, header in enumerate(headers):
                    if i < len(row):
                        recipient[header] = row[i].strip()
                
                # Only add if has name and at least one contact method
                if recipient.get('Name') and (recipient.get('Email ID - To') or recipient.get('Contact No.')):
                    recipients.append(recipient)
        
        return jsonify({
            "success": True,
            "data": recipients,
            "factory": factory,
            "sheet_name": sheet_name
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching recipients list: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/get_party_place_data", methods=["GET"])
def get_party_place_data():
    """Get party and place data from Google Sheets for dropdown population"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        # Get parameters
        factory = request.args.get('factory', 'KR')
        sheet_name = request.args.get('sheet_name', 'Party List')
        sheet_id = request.args.get('sheet_id')
        
        if not sheet_id:
            return jsonify({
                "success": False,
                "error": "Sheet ID is required"
            }), 400
        
        # Fetch party and place data from Google Sheets
        party_data = fetch_google_sheet_data(sheet_id, sheet_name)
        
        if not party_data or len(party_data) < 2:
            return jsonify({
                "success": True,
                "data": {
                    "party_names": [],
                    "places": [],
                    "party_place_mapping": {}
                },
                "message": f"No party data found in sheet {sheet_name}"
            }), 200
        
        # Extract party names and places from the data
        # Assuming structure: Party Name | Place
        party_names = []
        places = []
        party_place_mapping = {}
        
        for row in party_data[1:]:  # Skip header row
            if len(row) >= 2 and row[0].strip() and row[1].strip():
                party_name = row[0].strip()
                place = row[1].strip()
                
                party_names.append(party_name)
                places.append(place)
                party_place_mapping[party_name] = place
        
        # Remove duplicates
        party_names = list(set(party_names))
        places = list(set(places))
        
        return jsonify({
            "success": True,
            "data": {
                "party_names": party_names,
                "places": places,
                "party_place_mapping": party_place_mapping
            },
            "factory": factory,
            "sheet_name": sheet_name
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching party place data: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/material_inward", methods=["POST"])
def material_inward():
    """Record material inward transaction"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['category', 'materialName', 'uom', 'quantity', 'partyName', 'place']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({
                    "success": False,
                    "error": f"Missing required field: {field}"
                }), 400
        
        # Create material inward record
        inward_record = {
            'category': data['category'],
            'subCategory': data.get('subCategory', ''),
            'particulars': data.get('particulars', ''),
            'materialName': data['materialName'],
            'uom': data['uom'],
            'quantity': data['quantity'],
            'partyName': data['partyName'],
            'place': data['place'],
            'timestamp': data.get('timestamp', datetime.now().isoformat()),
            'department': data.get('department', 'KR'),
            'type': 'inward',
            'recordedBy': session.get('user', {}).get('email', 'unknown'),
            'recordedAt': datetime.now().isoformat()
        }
        
        # Here you would typically save to database
        # For now, we'll just log it and return success
        logger.info(f"Material inward recorded: {inward_record}")
        
        return jsonify({
            "success": True,
            "message": "Material inward recorded successfully",
            "data": inward_record
        }), 200
        
    except Exception as e:
        logger.error(f"Error recording material inward: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/material_outward", methods=["POST"])
def material_outward():
    """Record material outward transaction"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['category', 'materialName', 'uom', 'quantity', 'givenTo', 'description']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({
                    "success": False,
                    "error": f"Missing required field: {field}"
                }), 400
        
        # Create material outward record
        outward_record = {
            'category': data['category'],
            'subCategory': data.get('subCategory', ''),
            'particulars': data.get('particulars', ''),
            'materialName': data['materialName'],
            'uom': data['uom'],
            'quantity': data['quantity'],
            'givenTo': data['givenTo'],
            'description': data['description'],
            'timestamp': data.get('timestamp', datetime.now().isoformat()),
            'department': data.get('department', 'KR'),
            'type': 'outward',
            'recordedBy': session.get('user', {}).get('email', 'unknown'),
            'recordedAt': datetime.now().isoformat()
        }
        
        # Here you would typically save to database
        # For now, we'll just log it and return success
        logger.info(f"Material outward recorded: {outward_record}")
        
        return jsonify({
            "success": True,
            "message": "Material outward recorded successfully",
            "data": outward_record
        }), 200
        
    except Exception as e:
        logger.error(f"Error recording material outward: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/add_material", methods=["POST"])
def add_material():
    """Add a new material to Firebase"""
    try:
        logger.info("Add material endpoint called")
        logger.info(f"Request origin: {request.headers.get('Origin')}")
        logger.info(f"Request headers: {dict(request.headers)}")
        
        if 'user' not in session:
            logger.warning("User not in session for add_material request")
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        logger.info(f"Received data: {data}")
        
        # Validate required fields
        required_fields = ['category', 'materialName', 'uom']
        for field in required_fields:
            if not data.get(field):
                return jsonify({
                    "success": False,
                    "message": f"Missing required field: {field}"
                }), 400
        
        # Get factory from data or default to KR
        factory = data.get('department', 'KR')
        
        from Utils.firebase_utils import db
        from firebase_admin import firestore
        
        # Prepare material data (simplified - no extra fields)
        material_data = {
            'category': data.get('category'),
            'subCategory': data.get('subCategory', ''),
            'particulars': data.get('particulars', ''),
            'materialName': data.get('materialName'),
            'uom': data.get('uom')
        }
        
        # Get the factory document
        factory_ref = db.collection('MATERIAL').document(factory)
        factory_doc = factory_ref.get()
        
        if factory_doc.exists:
            # Get existing materials array
            existing_data = factory_doc.to_dict()
            materials = existing_data.get('materials', [])
        else:
            # Create new factory document with empty materials array
            materials = []
        
        # Add new material to the array
        materials.append(material_data)
        
        # Update the factory document with the new materials array
        factory_ref.set({
            'materials': materials
        }, merge=True)
        
        logger.info(f"Material added successfully: {material_data['materialName']} to factory {factory}")
        
        return jsonify({
            "success": True,
            "message": "Material added successfully"
        }), 200
        
    except Exception as e:
        logger.error(f"Error adding material: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error adding material: {str(e)}"
        }), 500


@app.route("/api/submit_order", methods=["POST"])
def submit_order():
    """Submit an order to Firebase ORDERS collection using factory initials as document name"""
    try:
        logger.info("Submit order endpoint called")
        logger.info(f"Request origin: {request.headers.get('Origin')}")
        
        if 'user' not in session:
            logger.warning("User not in session for submit_order request")
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        logger.info(f"Received order data: {data}")
        
        # Validate required fields
        required_fields = ['orderId', 'orderItems', 'givenBy', 'description', 'importance', 'factory']
        for field in required_fields:
            if not data.get(field):
                return jsonify({
                    "success": False,
                    "message": f"Missing required field: {field}"
                }), 400
        
        # Validate order items
        if not data.get('orderItems') or len(data.get('orderItems', [])) == 0:
            return jsonify({
                "success": False,
                "message": "Order must contain at least one item"
            }), 400
        
        # Prepare order data (without timestamps - they'll be added in add_order function)
        factory = data.get('factory', 'KR')
        order_data = {
            'orderId': data.get('orderId'),
            'orderItems': data.get('orderItems'),
            'givenBy': data.get('givenBy'),
            'description': data.get('description'),
            'importance': data.get('importance'),
            'factory': factory,
            'createdBy': session.get('user', {}).get('username', 'Unknown'),
            'status': 'Pending'
        }
        
        # Use the new add_order function with factory initials
        success = add_order(factory, order_data)
        
        if success:
            logger.info(f"Order submitted successfully: {order_data['orderId']} to factory {factory}")
            return jsonify({
                "success": True,
                "message": "Order submitted successfully",
                "orderId": order_data['orderId'],
                "factory": factory,
                "factoryDocument": get_factory_initials(factory)
            }), 200
        else:
            return jsonify({
                "success": False,
                "message": "Failed to submit order"
            }), 500
        
    except Exception as e:
        logger.error(f"Error submitting order: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error submitting order: {str(e)}"
        }), 500

@app.route("/api/get_orders", methods=["GET"])
def get_orders():
    """Get all orders for a specific factory"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        factory = request.args.get('factory', 'KR')
        orders = get_orders_by_factory(factory)
        
        return jsonify({
            "success": True,
            "data": orders,
            "factory": factory,
            "count": len(orders)
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching orders: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error fetching orders: {str(e)}"
        }), 500

@app.route("/api/get_order/<string:factory>/<string:order_id>", methods=["GET"])
def get_order(factory, order_id):
    """Get a specific order by ID from a factory"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        order = get_order_by_id(factory, order_id)
        
        if order:
            return jsonify({
                "success": True,
                "data": order
            }), 200
        else:
            return jsonify({
                "success": False,
                "message": "Order not found"
            }), 404
        
    except Exception as e:
        logger.error(f"Error fetching order: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error fetching order: {str(e)}"
        }), 500

@app.route("/api/update_order_status", methods=["POST"])
def update_order_status_endpoint():
    """Update the status of a specific order"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        factory = data.get('factory')
        order_id = data.get('orderId')
        new_status = data.get('status')
        
        if not all([factory, order_id, new_status]):
            return jsonify({
                "success": False,
                "message": "Missing required fields: factory, orderId, status"
            }), 400
        
        # Validate status values
        valid_statuses = ['Pending', 'In Progress', 'Completed', 'Cancelled', 'On Hold']
        if new_status not in valid_statuses:
            return jsonify({
                "success": False,
                "message": f"Invalid status. Must be one of: {', '.join(valid_statuses)}"
            }), 400
        
        updated_by = session.get('user', {}).get('username', 'Unknown')
        success = update_order_status(factory, order_id, new_status, updated_by)
        
        if success:
            return jsonify({
                "success": True,
                "message": f"Order status updated to {new_status}"
            }), 200
        else:
            return jsonify({
                "success": False,
                "message": "Order not found or update failed"
            }), 404
        
    except Exception as e:
        logger.error(f"Error updating order status: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error updating order status: {str(e)}"
        }), 500

@app.route("/api/delete_order", methods=["POST"])
def delete_order_endpoint():
    """Delete a specific order"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        factory = data.get('factory')
        order_id = data.get('orderId')
        
        if not all([factory, order_id]):
            return jsonify({
                "success": False,
                "message": "Missing required fields: factory, orderId"
            }), 400
        
        success = delete_order(factory, order_id)
        
        if success:
            return jsonify({
                "success": True,
                "message": "Order deleted successfully"
            }), 200
        else:
            return jsonify({
                "success": False,
                "message": "Order not found or delete failed"
            }), 404
        
    except Exception as e:
        logger.error(f"Error deleting order: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error deleting order: {str(e)}"
        }), 500

@app.route("/api/get_all_orders", methods=["GET"])
def get_all_orders_endpoint():
    """Get all orders from all factories (admin only)"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        # Check if user is admin
        current_user = session.get('user')
        if current_user.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403
        
        orders = get_all_orders()
        
        return jsonify({
            "success": True,
            "data": orders,
            "count": len(orders)
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching all orders: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error fetching all orders: {str(e)}"
        }), 500

@app.route("/api/get_plant_material_data", methods=["POST"])
def get_plant_material_data():
    """Get material data for a specific plant from Google Sheets"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        plant_id = data.get('plant_id')
        plant_data = data.get('plant_data')  # Plant data from frontend
        
        if not plant_id or not plant_data:
            return jsonify({"error": "Plant ID and plant data required"}), 400
        
        from Utils.process_utils import get_plant_material_data_from_sheets
        
        # Get plant name from the provided plant data
        plant_name = 'Unknown Plant'
        for plant in plant_data:
            if plant.get('material_sheet_id') == plant_id:
                plant_name = plant.get('name', 'Unknown Plant')
                break
        
        # Get user email from session for logging
        user_email = session.get('user', {}).get('email', 'Unknown User')
        logging.info(f"User {user_email} requested material data for plant {plant_name} (ID: {plant_id})")
        
        sheet_result = get_plant_material_data_from_sheets(plant_id, plant_data)
        
        # Handle both old and new return formats for backward compatibility
        if isinstance(sheet_result, dict) and 'material_data' in sheet_result:
            # New format with skipped data
            sheet_data = sheet_result['material_data']
        else:
            # Old format (fallback)
            sheet_data = sheet_result
        
        return jsonify({
            "success": True,
            "data": sheet_data,
            "plant_name": plant_name
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching plant material data: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/sync_plant_material_data", methods=["POST"])
def sync_plant_material_data():
    """Sync material data from Google Sheets to Firebase"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        data = request.get_json()
        plant_id = data.get('plant_id')
        plant_name = data.get('plant_name')
        plant_data = data.get('plant_data')  # Plant data from frontend
        sync_description = data.get('sync_description')
        sync_timestamp = data.get('sync_timestamp')
        
        if not all([plant_id, plant_name, sync_description, plant_data]):
            return jsonify({
                "success": False,
                "message": "Missing required fields: plant_id, plant_name, sync_description, plant_data"
            }), 400
        
        # Get user email from session for security
        user_email = session.get('user', {}).get('email', 'Unknown User')
        
        from Utils.process_utils import sync_plant_material_to_firebase
        
        # Sync material data to Firebase
        result = sync_plant_material_to_firebase(
            plant_id=plant_id,
            plant_name=plant_name,
            plant_data=plant_data,
            sync_description=sync_description,
            sync_timestamp=sync_timestamp,
            synced_by=user_email
        )
        
        if result['success']:
            return jsonify({
                "success": True,
                "message": f"Material data synced successfully for {plant_name}",
                "total_processed": result.get('total_processed', 0),
                "total_synced": result.get('total_synced', 0),
                "skipped_count": result.get('skipped_count', 0),
                "skipped_rows": result.get('skipped_rows', []),
                "skipped_reasons": result.get('skipped_reasons', {}),
                "data": result.get('data', {})
            }), 200
        else:
            return jsonify({
                "success": False,
                "message": result.get('message', 'Failed to sync material data')
            }), 500
        
    except Exception as e:
        logger.error(f"Error syncing plant material data: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error syncing material data: {str(e)}"
        }), 500


@app.route("/api/reset_order_counter", methods=["POST"])
def reset_order_counter_endpoint():
    """Reset the order counter for a factory (for testing purposes)"""
    try:
        data = request.get_json()
        factory = data.get('factory', 'KR')
        
        # Reset the counter
        success = reset_order_counter(factory)
        
        if success:
            return jsonify({
                "success": True, 
                "message": f"Order counter reset for {factory}",
                "factory": factory
            })
        else:
            return jsonify({
                "success": False, 
                "message": f"Failed to reset order counter for {factory}"
            }), 500
        
    except Exception as e:
        logging.error(f"Error in reset_order_counter_endpoint: {str(e)}", exc_info=True)
        return jsonify({
            "success": False,
            "message": f"Error resetting order counter: {str(e)}"
        }), 500


@app.route("/api/get_next_order_id", methods=["POST"])
def get_next_order_id_endpoint():
    """Get the next available order ID for a factory"""
    try:
        data = request.get_json()
        factory = data.get('factory', 'KR')
        
        # Generate next order ID using global counter
        order_id = get_next_order_id(factory)
        
        return jsonify({
            "success": True, 
            "orderId": order_id,
            "factory": factory
        })
        
    except Exception as e:
        logger.error(f"Error generating order ID: {str(e)}")
        return jsonify({"success": False, "message": f"Error generating order ID: {str(e)}"}), 500

@app.route("/api/send_order_notification", methods=["POST"])
def send_order_notification():
    """Send order notification via email and/or WhatsApp"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Not logged in"}), 401
        
        user_email = session.get('user', {}).get('email')
        if not user_email:
            return jsonify({"error": "No user email found"}), 400
        
        data = request.get_json()
        logger.info(f"Received order notification request: {data}")
        
        # Validate required fields
        required_fields = ['orderId', 'orderData', 'method', 'factory']
        for field in required_fields:
            if field not in data:
                return jsonify({
                    "success": False,
                    "message": f"Missing required field: {field}"
                }), 400
        
        order_id = data.get('orderId')
        order_data = data.get('orderData')
        recipients = data.get('recipients', [])
        method = data.get('method')  # 'email', 'whatsapp', or 'both'
        factory = data.get('factory', 'KR')
        auto_send = data.get('autoSend', False)
        sheet_id = data.get('sheetId')
        sheet_name = data.get('sheetName', 'Recipents List')
        
        # If autoSend is True and no recipients provided, fetch from Google Sheets
        if auto_send and (not recipients or len(recipients) == 0):
            if not sheet_id:
                return jsonify({
                    "success": False,
                    "message": "Sheet ID is required for auto-send"
                }), 400
            
            try:
                logger.info(f"Auto-send enabled, fetching recipients from Google Sheets")
                logger.info(f"Sheet: '{sheet_name}', ID: {sheet_id}, Factory: {factory}")
                
                # Fetch recipients from Google Sheets using provided sheet ID and name
                recipients_data = fetch_google_sheet_data(sheet_id, sheet_name)
                
                if not recipients_data or len(recipients_data) < 3:
                    return jsonify({
                        "success": False,
                        "message": f"No recipients found in sheet '{sheet_name}'. Need at least 3 rows (headers at row 2, data from row 3)"
                    }), 400
                
                # Headers are in Row 2 (index 1), data starts from Row 3 (index 2)
                headers = [h.strip() for h in recipients_data[1]]  # Row 2 contains headers
                
                # Expected headers: "Name, Country Code, Contact No., Email ID - To"
                expected_headers = ['Name', 'Country Code', 'Contact No.', 'Email ID - To']
                logger.info(f"Found headers in sheet: {headers}")
                
                # Process data rows starting from Row 3 (index 2)
                recipients = []
                for row in recipients_data[2:]:
                    if row and len(row) > 0:
                        recipient = {}
                        for i, header in enumerate(headers):
                            if i < len(row):
                                recipient[header] = row[i].strip()
                        
                        # Only add if has name and at least one contact method
                        if recipient.get('Name') and (recipient.get('Email ID - To') or recipient.get('Contact No.')):
                            recipients.append(recipient)
                
                logger.info(f"Successfully fetched {len(recipients)} recipients from Google Sheets")
                
            except Exception as e:
                logger.error(f"Error fetching recipients from Google Sheets: {str(e)}")
                return jsonify({
                    "success": False,
                    "message": f"Error fetching recipients: {str(e)}"
                }), 500
        
        if not recipients or len(recipients) == 0:
            return jsonify({
                "success": False,
                "message": "No recipients available for notification"
            }), 400
        
        # Get user-specific temporary directory
        from Utils.temp_manager import get_user_temp_dir
        temp_dir = get_user_temp_dir(user_email, OUTPUT_DIR)
        
        # Get template path
        template_path = os.path.join(os.path.dirname(__file__), "reactorreportformat.docx")
        
        if not os.path.exists(template_path):
            logger.warning("Reactor report template not found, will create basic document")
            template_path = None
        
        # Import the order notification processing function
        from Utils.process_utils import process_order_notification
        
        # Process and send order notification
        result = process_order_notification(
            order_id=order_id,
            order_data=order_data,
            recipients=recipients,
            method=method,
            factory=factory,
            template_path=template_path,
            output_dir=temp_dir,
            user_email=user_email,
            logger=logger,
            send_email_smtp=send_email_smtp,
            send_whatsapp_message=send_whatsapp_message
        )
        
        # Clean up user-specific temporary directory
        from Utils.temp_manager import cleanup_user_temp_dir
        cleanup_user_temp_dir(user_email, OUTPUT_DIR)
        
        if result.get('success'):
            return jsonify({
                "success": True,
                "message": result.get('message', 'Notifications sent successfully'),
                "delivery_stats": result.get('delivery_stats', {})
            }), 200
        else:
            return jsonify({
                "success": False,
                "message": result.get('message', 'Failed to send notifications'),
                "errors": result.get('errors', [])
            }), 500
        
    except Exception as e:
        logger.error(f"Error sending order notification: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error sending order notification: {str(e)}"
        }), 500


def send_log_report_to_user(user_email, delivery_stats, send_email_enabled, send_whatsapp_enabled, logger):
    """
    Send log report to the user who generated the reports
    """
    try:
        # Import the PDF generation function
        from Utils.process_utils import generate_log_report_pdf
        
        # Generate PDF log report
        pdf_path = generate_log_report_pdf(delivery_stats, OUTPUT_DIR, logger)
        
        if not pdf_path:
            logger.error("Failed to generate PDF log report, falling back to text format")
            # Fallback to text format if PDF generation fails
            send_log_report_text_fallback(user_email, delivery_stats, send_email_enabled, send_whatsapp_enabled, logger)
            return
        
        # Format the log report text for WhatsApp (simplified version)
        total_recipients = delivery_stats.get("total_recipients", 0)
        successful_deliveries = delivery_stats.get("successful_deliveries", 0)
        failed_deliveries = delivery_stats.get("failed_deliveries", 0)
        
        log_report_text = "📊 DELIVERY LOG REPORT\n\n"
        log_report_text += f"📋 Total messages to be delivered: {total_recipients}\n"
        log_report_text += f"✅ Messages delivered successfully: {successful_deliveries}\n"
        log_report_text += f"❌ Failed messages: {failed_deliveries}\n\n"
        log_report_text += "📄 Detailed PDF report has been attached to your email.\n"
        log_report_text += "📱 This WhatsApp message contains the summary only."

        # Send email log report if email notifications were enabled
        if send_email_enabled:
            try:
                email_subject = "📊 Report Generation Log - Delivery Summary"
                email_body = f"""
                <html>
                <body>
                <h2>Report Generation Log</h2>
                <p>Here is the delivery summary for the reports you generated:</p>
                <p><strong>📋 Total messages to be delivered:</strong> {total_recipients}</p>
                <p><strong>✅ Messages delivered successfully:</strong> {successful_deliveries}</p>
                <p><strong>❌ Failed messages:</strong> {failed_deliveries}</p>
                <p><strong>📊 Success Rate:</strong> {(successful_deliveries/total_recipients*100):.1f}%</p>
                
                <p>A detailed PDF report with complete delivery statistics and failed contacts table is attached to this email.</p>
                <p><strong>Note:</strong> If you have WhatsApp authenticated in your account, you will also receive a summary via WhatsApp automatically.</p>
                <br>
                <p>Best regards,<br>Bajaj Earths Automation System</p>
                </body>
                </html>
                """
                
                # Send email with PDF attachment
                send_email_smtp(
                    recipient_email=user_email,
                    subject=email_subject,
                    body=email_body,
                    attachment_paths=[pdf_path],
                    user_email=user_email
                )
                logger.info(f"Log report email with PDF attachment sent successfully to {user_email}")
                
            except Exception as e:
                logger.error(f"Failed to send log report email to {user_email}: {e}")

        # Send WhatsApp log report if WhatsApp notifications were enabled
        if send_whatsapp_enabled:
            try:
                # Get user phone number from WhatsApp service
                user_phone = get_user_phone_from_whatsapp_service(user_email)
                if user_phone:
                    # Send WhatsApp message with PDF attachment using caption
                    options = {
                        'use_template_as_caption': True  # Always use caption for log reports
                    }
                    success = send_whatsapp_message(
                        contact_name="Report Generator",
                        message=log_report_text,
                        file_paths=[pdf_path],
                        whatsapp_number=user_phone,
                        process_name="log_report",
                        options=options
                    )
                    if success:
                        logger.info(f"Log report WhatsApp with PDF attachment sent successfully to {user_email}")
                    else:
                        logger.warning(f"Failed to send log report WhatsApp to {user_email}")
                else:
                    logger.info(f"No phone number found for user {user_email} - skipping WhatsApp log report")
                
            except Exception as e:
                logger.error(f"Failed to send log report WhatsApp to user: {e}")
        
        # Clean up the PDF file after sending
        try:
            import os
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
                logger.info(f"Cleaned up temporary PDF log report: {pdf_path}")
        except Exception as e:
            logger.warning(f"Failed to clean up PDF log report {pdf_path}: {e}")

    except Exception as e:
        logger.error(f"Error in send_log_report_to_user: {e}")
        # Fallback to text format if there's an error
        try:
            send_log_report_text_fallback(user_email, delivery_stats, send_email_enabled, send_whatsapp_enabled, logger)
        except Exception as fallback_error:
            logger.error(f"Fallback text log report also failed: {fallback_error}")

def send_log_report_text_fallback(user_email, delivery_stats, send_email_enabled, send_whatsapp_enabled, logger):
    """
    Fallback function to send text-based log report if PDF generation fails
    """
    try:
        # Format the log report as text
        total_recipients = delivery_stats.get("total_recipients", 0)
        successful_deliveries = delivery_stats.get("successful_deliveries", 0)
        failed_deliveries = delivery_stats.get("failed_deliveries", 0)
        failed_contacts = delivery_stats.get("failed_contacts", [])

        # Create the log report text
        log_report_text = "📊 DELIVERY LOG REPORT\n\n"
        log_report_text += f"📋 Total messages to be delivered: {total_recipients}\n"
        log_report_text += f"✅ Messages delivered successfully: {successful_deliveries}\n"
        log_report_text += f"❌ Failed messages: {failed_deliveries}\n\n"

        if failed_contacts:
            log_report_text += "📋 FAILED CONTACTS TABLE:\n"
            log_report_text += "┌─────────────────┬──────────────────┬─────────────────────────────────────┐\n"
            log_report_text += "│ Name            │ Contact Number   │ Reason                              │\n"
            log_report_text += "├─────────────────┼──────────────────┼─────────────────────────────────────┤\n"
            
            for contact in failed_contacts:
                name = contact.get("name", "Unknown")
                contact_num = contact.get("contact", "N/A")
                reason = contact.get("reason", "Unknown error")
                
                # Truncate long strings to fit in table
                truncated_name = name[:15] if len(name) <= 15 else name[:12] + "..."
                truncated_contact = contact_num[:16] if len(contact_num) <= 16 else contact_num[:13] + "..."
                truncated_reason = reason[:35] if len(reason) <= 35 else reason[:32] + "..."
                
                log_report_text += f"│ {truncated_name:<15} │ {truncated_contact:<16} │ {truncated_reason:<35} │\n"
            
            log_report_text += "└─────────────────┴──────────────────┴─────────────────────────────────────┘\n"

        # Send email log report if email notifications were enabled
        if send_email_enabled:
            try:
                email_subject = "📊 Report Generation Log - Delivery Summary"
                email_body = f"""
                <html>
                <body>
                <h2>Report Generation Log</h2>
                <p>Here is the delivery summary for the reports you generated:</p>
                <pre style="font-family: monospace; background-color: #f5f5f5; padding: 10px; border-radius: 5px;">
                {log_report_text.replace('\n', '<br>')}
                </pre>
                <p>This log report shows the delivery status of all recipients in your Google Sheet.</p>
                <p><strong>Note:</strong> If you have WhatsApp authenticated in your account, you will also receive this log report via WhatsApp automatically.</p>
                <br>
                <p>Best regards,<br>Bajaj Earths Automation System</p>
                </body>
                </html>
                """
                
                send_email_smtp(
                    recipient_email=user_email,
                    subject=email_subject,
                    body=email_body,
                    user_email=user_email
                )
                logger.info(f"Log report email sent successfully to {user_email}")
                
            except Exception as e:
                logger.error(f"Failed to send log report email to {user_email}: {e}")

        # Send WhatsApp log report if WhatsApp notifications were enabled
        if send_whatsapp_enabled:
            try:
                # Get user phone number from WhatsApp service
                user_phone = get_user_phone_from_whatsapp_service(user_email)
                if user_phone:
                    # Send WhatsApp message with log report
                    success = send_whatsapp_message(
                        contact_name="Report Generator",
                        message=log_report_text,
                        whatsapp_number=user_phone,
                        process_name="log_report"
                    )
                    if success:
                        logger.info(f"Log report WhatsApp sent successfully to {user_email}")
                    else:
                        logger.warning(f"Failed to send log report WhatsApp to {user_email}")
                else:
                    logger.info(f"No phone number found for user {user_email} - skipping WhatsApp log report")
                
            except Exception as e:
                logger.error(f"Failed to send log report WhatsApp to user: {e}")

    except Exception as e:
        logger.error(f"Error in send_log_report_text_fallback: {e}")

def get_user_phone_from_whatsapp_service(user_email):
    """
    Get user phone number from WhatsApp service
    """
    try:
        # Create WhatsApp client to get user info
        client = WhatsAppNodeClient(WHATSAPP_NODE_SERVICE_URL, user_email)
        
        # Get WhatsApp status which includes user info
        status_data = client.get_status()
        
        if status_data and status_data.get('authenticated') and status_data.get('userInfo'):
            user_info = status_data.get('userInfo')
            phone_number = user_info.get('phoneNumber')
            
            if phone_number:
                logger.info(f"Found WhatsApp phone number for user {user_email}: {phone_number}")
                return phone_number
        
        logger.info(f"No WhatsApp phone number found for user {user_email}")
        return None
        
    except Exception as e:
        logger.error(f"Error getting WhatsApp phone number for user {user_email}: {e}")
        return None


if __name__ == "__main__":
    try:
        logger.info("Starting SS Automation backend server...")
        ensure_directories()
        logger.info("Directories ensured")
        
        # Start Flask app
        app.run(host="0.0.0.0", port=7082, debug=True, use_reloader=False)
    except Exception as e:
        logger.error("Failed to start server: {}".format(str(e)), exc_info=True)
        sys.exit(1)