import os
import re
import json
import logging
from docx import Document
from flask import session
from Utils.email_utils import *
from Utils.whatsapp_utils import (
    get_employee_contact,
    send_whatsapp_message,
)
from Utils.drive_utils import upload_to_google_drive, upload_reactor_report_to_drive
import shutil
import subprocess
import platform
# import pythoncom
# from comtypes.client import CreateObject
import os
from docx import Document
from docx.shared import Inches
import requests
from PIL import Image
import io
import base64
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import gspread
from Utils.config import creds
from Utils.firebase_utils import db
from datetime import datetime, timedelta
from Utils.whatsapp_utils import handle_reactor_report_notification, handle_reactor_report_notification_with_stats
import traceback

# Configure logging
logging.basicConfig(level=logging.INFO)

# ============================================================================
# SHARED TABLE FORMATTING FUNCTIONS
# ============================================================================

def set_cell_background_color(cell, color):
    """
    Set cell background color using XML styling
    """
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
    except Exception as e:
        logging.warning(f"Could not set cell background color: {e}")

def set_table_borders_professional(table, logger=None):
    """
    Set professional table borders with custom styling
    """
    try:
        from docx.oxml.shared import OxmlElement, qn
        
        # Get the table element
        tbl = table._tbl
        
        # Create table properties if they don't exist
        tblPr = tbl.xpath('w:tblPr')
        if not tblPr:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        else:
            tblPr = tblPr[0]
        
        # Remove any existing borders
        for border in tblPr.xpath('w:tblBorders'):
            tblPr.remove(border)
        
        # Create new borders element
        borders = OxmlElement('w:tblBorders')
        
        # Outer borders: thick orange (level 2 thickness, color e69138)
        outer_border_props = {
            'w:val': 'single',
            'w:sz': '16',  # Level 2 thickness (2pt)
            'w:color': 'e69138'
        }
        
        # Internal borders: thin black
        inner_border_props = {
            'w:val': 'single',
            'w:sz': '4',   # Thin internal borders (0.5pt)
            'w:color': '000000'  # Black color
        }
        
        # Add outer borders (thick orange)
        outer_border_types = ['w:top', 'w:bottom', 'w:left', 'w:right']
        for border_type in outer_border_types:
            border = OxmlElement(border_type)
            for prop, value in outer_border_props.items():
                border.set(qn(prop), value)
            borders.append(border)
        
        # Add internal borders (thin black)
        inner_border_types = ['w:insideH', 'w:insideV']
        for border_type in inner_border_types:
            border = OxmlElement(border_type)
            for prop, value in inner_border_props.items():
                border.set(qn(prop), value)
            borders.append(border)
        
        # Add borders to table properties
        tblPr.append(borders)
        
    except Exception as e:
        if logger:
            logger.warning(f"Could not set table borders: {e}")
        else:
            logging.warning(f"Could not set table borders: {e}")

def set_table_borders_simple(table, logger=None):
    """
    Set simple table borders with black lines
    """
    try:
        from docx.oxml.shared import OxmlElement, qn
        
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        
        # Define border styles - simple black borders
        border_elements = [
            ('top', 'single', '000000', '4'),
            ('left', 'single', '000000', '4'),
            ('bottom', 'single', '000000', '4'),
            ('right', 'single', '000000', '4'),
            ('insideH', 'single', '000000', '4'),
            ('insideV', 'single', '000000', '4')
        ]
        
        for border_name, border_style, border_color, border_size in border_elements:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), border_size)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)
        
        tbl.tblPr.append(tblBorders)
        
    except Exception as e:
        if logger:
            logger.warning(f"Could not set simple table borders: {e}")
        else:
            logging.warning(f"Could not set simple table borders: {e}")

def set_table_no_page_break(table, logger=None):
    """
    Set table properties to prevent page breaks
    """
    try:
        from docx.oxml.shared import OxmlElement, qn
        
        tbl = table._tbl
        tblPr = tbl.get_or_add_tblPr()
        
        # Remove any existing table layout
        for layout in tblPr.xpath('w:tblLayout'):
            tblPr.remove(layout)
        
        # Add table layout to prevent page breaks
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        
        # Add table properties to prevent page breaks
        tblLook = OxmlElement('w:tblLook')
        tblLook.set(qn('w:firstRow'), '1')
        tblLook.set(qn('w:lastRow'), '0')
        tblLook.set(qn('w:firstCol'), '0')
        tblLook.set(qn('w:lastCol'), '0')
        tblLook.set(qn('w:noHBand'), '0')
        tblLook.set(qn('w:noVBand'), '1')
        tblPr.append(tblLook)
        
        # Set table to not break across pages
        tblCellSpacing = OxmlElement('w:tblCellSpacing')
        tblCellSpacing.set(qn('w:w'), '0')
        tblCellSpacing.set(qn('w:type'), 'dxa')
        tblPr.append(tblCellSpacing)
        
        # Add table properties to prevent page breaks (critical for table integrity)
        tblPr.set(qn('w:tblStyle'), 'TableGrid')
        
        # Set table to not break across pages using tblPr properties
        tblPr.set(qn('w:tblW'), '0')
        tblPr.set(qn('w:tblInd'), '0')
        
        # Remove any existing table break properties that might force page breaks
        for tblBreak in tblPr.xpath('w:tblBreak'):
            tblPr.remove(tblBreak)
        
        # Add cantSplit property to prevent table from breaking across pages
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')  # 1 = true, prevents splitting
        tblPr.append(cantSplit)
        
        if logger:
            logger.info("Successfully set table properties to prevent page breaks")
        
    except Exception as e:
        if logger:
            logger.warning(f"Could not set table page break properties: {e}")
        else:
            logging.warning(f"Could not set table page break properties: {e}")

def format_table_professional(table, is_header=False, logger=None):
    """
    Apply professional styling to a table with reactor report style formatting
    """
    try:
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Set table alignment to center
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = True
        
        # Set table to not break across pages - CRITICAL FOR TABLE INTEGRITY
        set_table_no_page_break(table, logger)
        
        # Apply professional borders
        set_table_borders_professional(table, logger)
        
        # Format each cell
        for i, row in enumerate(table.rows):
            try:
                for j, cell in enumerate(row.cells):
                    try:
                        # Set vertical alignment to center
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # Set paragraph alignment
                        for paragraph in cell.paragraphs:
                            try:
                                if i == 0:  # Header row
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:  # Data rows
                                    if j == 0:  # First column
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    else:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # Format text
                                for run in paragraph.runs:
                                    try:
                                        if i == 0:  # Header row
                                            run.bold = True
                                            run.font.size = Pt(11)
                                        else:
                                            run.font.size = Pt(10)
                                    except Exception as e:
                                        if logger:
                                            logger.warning(f"Error formatting run in cell [{i},{j}]: {e}")
                                        continue
                            except Exception as e:
                                if logger:
                                    logger.warning(f"Error formatting paragraph in cell [{i},{j}]: {e}")
                                continue
                        
                        # Set background colors
                        try:
                            if i == 0:  # Header row
                                set_cell_background_color(cell, "f9cb9c")
                            else:
                                # Alternate row colors
                                if i % 2 == 0:
                                    set_cell_background_color(cell, "FFFFFF")
                                else:
                                    set_cell_background_color(cell, "e8f0fe")
                        except Exception as e:
                            if logger:
                                logger.warning(f"Error setting background color for cell [{i},{j}]: {e}")
                            continue
                    except Exception as e:
                        if logger:
                            logger.warning(f"Error processing cell [{i},{j}]: {e}")
                        continue
            except Exception as e:
                if logger:
                    logger.warning(f"Error processing row {i}: {e}")
                continue
        
    except Exception as e:
        if logger:
            logger.warning(f"Error in format_table_professional: {e}")
        else:
            logging.warning(f"Error in format_table_professional: {e}")

def format_table_simple(table, is_header=False, logger=None):
    """
    Apply simple styling to a table with basic formatting
    """
    try:
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Set table alignment
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Apply simple borders
        set_table_borders_simple(table, logger)
        
        # Format each cell
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Set cell margins
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Format paragraphs in cell
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Format runs in paragraph
                    for run in paragraph.runs:
                        if row_idx == 0:  # Header row
                            run.bold = True
                            run.font.size = Pt(10)
                        else:
                            run.font.size = Pt(9)
        
    except Exception as e:
        if logger:
            logger.warning(f"Error in format_table_simple: {e}")
        else:
            logging.warning(f"Error in format_table_simple: {e}")

def should_start_table_on_new_page(doc, table_rows, table_name_length=0, logger=None):
    """
    Rough estimation to determine if table should start on new page.
    """
    try:
        # Estimate space needed for table
        # Each row approximately 0.3 inches (including padding)
        # Table name approximately 0.2 inches
        # Header spacing approximately 0.1 inches
        estimated_table_height = (table_rows * 0.3) + 0.2 + 0.1
        
        # Standard page height is approximately 11 inches
        # Leave 1 inch margin at bottom
        available_page_height = 10.0
        
        # If table height exceeds available space, suggest new page
        if estimated_table_height > available_page_height:
            return True
            
        return False
    except Exception as e:
        if logger:
            logger.warning(f"Error estimating table space: {e}")
        else:
            logging.warning(f"Error estimating table space: {e}")
        return False  # Default to current page if estimation fails

def prepare_file_paths(file_paths, user_email=None, base_output_dir=None, is_upload=False):
    """
    Prepare file paths for processing, handling both uploaded files and existing file paths.
    This function validates file paths and saves uploaded files to user-specific temporary directory.
    
    Args:
        file_paths: List of file paths or uploaded file objects
        user_email: User's email address for user-specific temp directory
        base_output_dir: Base output directory path
        is_upload: Boolean indicating if files are uploaded (need to be saved)
        
    Returns:
        list: List of valid file paths ready for processing
    """
    try:
        # Import temp_manager here to avoid circular imports
        from .temp_manager import get_user_temp_dir
        
        # Initialize empty list if no file paths provided
        if not file_paths:
            return []
            
        # Convert single path to list if it's not already
        if not isinstance(file_paths, list):
            file_paths = [file_paths]
        
        # Get user-specific temp directory
        temp_dir = None
        if is_upload and user_email and base_output_dir:
            temp_dir = get_user_temp_dir(user_email, base_output_dir)
        elif is_upload and not user_email:
            # Fallback to base temp directory if no user email provided
            temp_dir = os.path.join(base_output_dir, "temp") if base_output_dir else None
            if temp_dir:
                os.makedirs(temp_dir, exist_ok=True)
            
        # Validate each path and collect valid ones
        valid_paths = []
        seen_filenames = set()  # Keep track of filenames we've already processed
        errors = []
        
        for path in file_paths:
            try:
                if is_upload:
                    # Handle uploaded file
                    if hasattr(path, 'filename') and path.filename:
                        if temp_dir:
                            temp_path = os.path.join(temp_dir, path.filename)
                            path.save(temp_path)
                            valid_paths.append(temp_path)
                            seen_filenames.add(path.filename)
                            logging.info(f"Saved attachment file to user temp dir: {temp_path}")
                        else:
                            errors.append("No temp directory available for uploaded file")
                            logging.error("No temp directory available for uploaded file")
                else:
                    # Handle existing file path
                    if os.path.exists(path) and os.path.isfile(path):
                        filename = os.path.basename(path)
                        if filename not in seen_filenames:
                            valid_paths.append(path)
                            seen_filenames.add(filename)
                            logging.info(f"Added file: {path}")
                        else:
                            logging.warning(f"Duplicate file found: {path}. Skipping.")
                    else:
                        errors.append(f"Invalid or non-existent file path: {path}")
                        logging.warning(f"Invalid or non-existent file path: {path}")
            except Exception as e:
                errors.append(f"Error processing file {path}: {str(e)}")
                logging.error(f"Error processing file {path}: {str(e)}")
                
        logging.info(f"Prepared {len(valid_paths)} valid file paths for user: {user_email}")
        if errors:
            logging.warning(f"File preparation completed with {len(errors)} errors")
        
        return valid_paths
    except Exception as e:
        logging.error(f"Error preparing file paths: {str(e)}")
        return []

# Load message templates
def load_message_templates():
    try:
        with open('backend/Utils/message.json', 'r') as f:
            return json.load(f)
    except Exception as e:
        logging.error("Error loading message templates: {}".format(e))
        return None

# Preprocess headers
def preprocess_headers(headers):
    return [header.replace("\n", " ").strip().strip('"') for header in headers]

def convert_docx_to_pdf(input_path, output_path):
    try:
        # Check if input file exists
        if not os.path.exists(input_path):
            logging.error(f"Input file does not exist: {input_path}")
            return False
            
        logging.info(f"Converting DOCX to PDF: {input_path} -> {output_path}")
        
        # Check if LibreOffice is available
        try:
            subprocess.run(['libreoffice', '--version'], check=True, capture_output=True)
            logging.info("LibreOffice is available")
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            logging.error(f"LibreOffice is not available: {e}")
            return False
        
        process = subprocess.Popen([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path),
            input_path
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        stdout, stderr = process.communicate()
        return_code = process.returncode
        
        logging.info(f"LibreOffice conversion return code: {return_code}")
        if stdout:
            logging.info(f"LibreOffice stdout: {stdout.decode()}")
        if stderr:
            logging.warning(f"LibreOffice stderr: {stderr.decode()}")
        
        # Check if the output file was actually created
        if os.path.exists(output_path):
            logging.info(f"PDF conversion successful: {output_path}")
            return True
        else:
            logging.error(f"PDF conversion failed - output file not created: {output_path}")
            return False
            
    except Exception as e:
        logging.error(f"Error converting DOCX to PDF: {e}")
        return False

# Format file path
def format_file_path(file_path):
    if isinstance(file_path, str):
        return file_path.replace("\\\\", "\\")
    elif isinstance(file_path, list):
        return [f.replace("\\\\", "\\") for f in file_path]
    return file_path

# Clear the Salary_Slips folder
def clear_salary_slips_folder(output_dir):
    try:
        for filename in os.listdir(output_dir):
            file_path = os.path.join(output_dir, filename)
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        logging.info("Cleared the contents of the folder: {}".format(output_dir))
    except Exception as e:
        logging.error("Error clearing the folder {}: {}".format(output_dir, e))

def delete_generated_files(file_paths, drive_upload_success=None, logger=None):
    """
    Generic function to delete generated files after processing.
    Only deletes files if Google Drive upload was successful.
    
    Args:
        file_paths: Can be:
            - Single file path (str): "path/to/file.pdf"
            - List of file paths (list): ["path1.pdf", "path2.docx"]
            - Result dictionary with file path keys: {"pdf_path": "...", "docx_path": "...", "output_file": "..."}
        drive_upload_success: bool or None
            - True: Delete files (upload succeeded)
            - False: Skip deletion, keep files (upload failed)
            - None: Skip deletion, keep files (upload not attempted or unknown)
        logger: Optional logger instance (uses logging if not provided)
        
    Returns:
        dict: {
            'success': bool,
            'deleted_files': list of successfully deleted file paths,
            'failed_files': list of failed deletions with reasons,
            'total_deleted': int,
            'skipped_reason': str or None
        }
    """
    log = logger if logger else logging
    
    result = {
        'success': False,
        'deleted_files': [],
        'failed_files': [],
        'total_deleted': 0,
        'skipped_reason': None
    }
    
    # Check if we should delete files
    if drive_upload_success is not True:
        if drive_upload_success is False:
            result['skipped_reason'] = "Google Drive upload failed - keeping files for retry"
            log.info("Skipping file deletion: Google Drive upload failed. Files will be kept for next upload attempt.")
        else:
            result['skipped_reason'] = "Google Drive upload not attempted or status unknown - keeping files"
            log.info("Skipping file deletion: Google Drive upload was not attempted or status is unknown. Files will be kept.")
        return result
    
    # Extract file paths from different input formats
    files_to_delete = []
    
    if isinstance(file_paths, str):
        # Single file path
        files_to_delete = [file_paths]
    elif isinstance(file_paths, list):
        # List of file paths
        files_to_delete = file_paths
    elif isinstance(file_paths, dict):
        # Result dictionary - extract file paths
        # Check for common keys: pdf_path, docx_path, output_file
        if 'pdf_path' in file_paths and file_paths['pdf_path']:
            files_to_delete.append(file_paths['pdf_path'])
        if 'docx_path' in file_paths and file_paths['docx_path']:
            files_to_delete.append(file_paths['docx_path'])
        if 'output_file' in file_paths and file_paths['output_file']:
            # Only add if not already in list (output_file might be same as pdf_path)
            if file_paths['output_file'] not in files_to_delete:
                files_to_delete.append(file_paths['output_file'])
    else:
        error_msg = f"Invalid file_paths type: {type(file_paths)}"
        log.error(error_msg)
        result['failed_files'].append({'path': str(file_paths), 'reason': error_msg})
        return result
    
    # Remove duplicates and None/empty values
    files_to_delete = [f for f in files_to_delete if f and isinstance(f, str)]
    files_to_delete = list(set(files_to_delete))  # Remove duplicates
    
    if not files_to_delete:
        log.info("No files to delete (no valid file paths provided)")
        result['skipped_reason'] = "No valid file paths provided"
        return result
    
    log.info(f"Attempting to delete {len(files_to_delete)} file(s) after successful Google Drive upload")
    
    # Delete each file
    for file_path in files_to_delete:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                result['deleted_files'].append(file_path)
                result['total_deleted'] += 1
                log.info(f"Successfully deleted file: {file_path}")
            else:
                log.info(f"File does not exist, skipping deletion: {file_path}")
                result['deleted_files'].append(file_path)  # Count as "deleted" since it's already gone
                result['total_deleted'] += 1
        except Exception as e:
            error_msg = f"Failed to delete file {file_path}: {str(e)}"
            log.warning(error_msg)
            result['failed_files'].append({'path': file_path, 'reason': str(e)})
    
    # Set overall success status
    if result['total_deleted'] > 0 and len(result['failed_files']) == 0:
        result['success'] = True
        log.info(f"File cleanup completed: Successfully deleted {result['total_deleted']} file(s)")
    elif result['total_deleted'] > 0:
        result['success'] = True  # Partial success
        log.warning(f"File cleanup completed with {len(result['failed_files'])} failure(s): Deleted {result['total_deleted']} file(s)")
    else:
        log.warning("File cleanup failed: No files were deleted")
    
    return result

def format_months_list(months_data):
    """Format the list of months for display in messages."""
    if not months_data:
        return ""
    
    if isinstance(months_data, list):
        # Format for WhatsApp message
        return "\n".join(["   -  {} {}".format(month['month'], month['year']) for month in months_data])
    else:
        # Format for email HTML
        return "".join(["<li>{} {}</li>".format(month['month'], month['year']) for month in months_data])

def handle_whatsapp_notification(contact_name, full_month, full_year, whatsapp_number, file_path, is_special=False, months_data=None):
    """Delegate to Node-backed WhatsApp notification handler"""
    try:
        # Call the actual WhatsApp sending function with file paths
        # Message content will be handled by the WhatsApp service using templates
        return send_whatsapp_message(
            contact_name=contact_name,
            message="",  # Empty message - will use template
            file_paths=file_path,
            whatsapp_number=whatsapp_number,
            process_name="salary_slip",
            options={"isMultiple": is_special, "variables": {"full_month": full_month, "full_year": full_year, "months_data": months_data}}
        )
    except Exception as e:
        logging.error(f"Error delegating WhatsApp notification: {e}")
        return False

# Generate and process salary slips for a single employee
def process_salary_slip(template_path, output_dir, employee_identifier, employee_data, headers, drive_data, email_employees, contact_employees, month, year, full_month, full_year, send_whatsapp, send_email, is_special=False, months_data=None, collected_pdfs=None):
    logging.info("Starting process_salary_slip function")
    headers = preprocess_headers(headers)
    errors = []
    warnings = []
    
    try:
        # Find indices of ESIC headers
        esic_indices = [i for i, header in enumerate(headers) if header == "ESIC"]
        
        # Create placeholders dictionary, using the second ESIC value if it exists
        placeholders = {}
        for i, header in enumerate(headers):
            if header == "ESIC" and i == esic_indices[1] and len(esic_indices) > 1:
                # Skip the first ESIC if there's a second one
                continue
            placeholders[header] = employee_data[i]

        # Add month and year placeholders to the dictionary
        placeholders["Month"] = full_month
        placeholders["Year"] = full_year

        # Merge data from "Official Details" sheet
        logging.info("Looking for employee code: {}".format(employee_identifier))
        logging.info("First few items in drive_data: {}".format(drive_data[:2]))
        official_details = next((item for item in drive_data if item.get("Employee\nCode") == employee_identifier), {})
        logging.info("Found official details: {}".format(official_details))
        placeholders.update(official_details)
        logging.info("Updated placeholders: {}".format(placeholders))

        # Calculate components of salary with error handling
        try:
            present_salary_str = placeholders.get("Present Salary", "")
            present_salary = float(re.sub(r'[^\d.]', '', present_salary_str))
            if present_salary <= 0:
                errors.append("Present Salary must be greater than zero")
                # Set default values to continue processing
                placeholders["BS"] = "0"
                placeholders["HRA"] = "0"
                placeholders["SA"] = "0"
            else:
                placeholders["BS"] = str(round(present_salary * 0.40))
                placeholders["HRA"] = str(round(present_salary * 0.20))
                placeholders["SA"] = str(round(present_salary * 0.40))

            # Calculate OT by combining OT (Days) and OT (EH)
            try:
                ot_days = float(re.sub(r'[^\d.]', '', placeholders.get("OT\n(Days)", "0")))
                ot_eh = float(re.sub(r'[^\d.]', '', placeholders.get("OT\n(EH)", "0")))
                total_ot = ot_days + ot_eh
                placeholders["OT"] = str(round(total_ot))
            except ValueError as e:
                warnings.append(f"Invalid OT values, setting to 0: {e}")
                placeholders["OT"] = "0"
                
        except ValueError as e:
            errors.append(f"Invalid Present Salary for {placeholders.get('Name', 'Unknown')}: {e}")
            # Set default values to continue processing
            placeholders["BS"] = "0"
            placeholders["HRA"] = "0"
            placeholders["SA"] = "0"
            placeholders["OT"] = "0"

        # Ensure all placeholders are strings
        placeholders = {k: str(v) for k, v in placeholders.items()}

        # Load template and replace placeholders
        output_pdf = None
        output_docx = None
        drive_upload_success = None  # Track Drive upload status: None = not attempted, True = success, False = failed
        try:
            template = Document(template_path)
            for paragraph in template.paragraphs:
                for run in paragraph.runs:
                    for placeholder, value in placeholders.items():
                        if "{{{}}}".format(placeholder) in run.text:
                            run.text = run.text.replace("{{{}}}".format(placeholder), value)

            for table in template.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                for placeholder, value in placeholders.items():
                                    if "{{{}}}".format(placeholder) in run.text:
                                        run.text = run.text.replace("{{{}}}".format(placeholder), value)

            # Save output files
            employee_name = re.sub(r'[^\w\s]', '', placeholders.get("Name", "Employee"))
            output_docx = os.path.join(output_dir, "Salary Slip_{}_{}{}.docx".format(employee_name, month, year))
            template.save(output_docx)
            output_pdf = os.path.join(output_dir, "Salary Slip_{}_{}{}.pdf".format(employee_name, month, year))
            template.save(output_pdf)

            if convert_docx_to_pdf(output_docx, output_pdf):
                # Upload to Google Drive
                try:
                    # Get Google Drive ID
                    folder_id = official_details.get("Google Drive ID")
                    logging.info("Google Drive ID: {}".format(folder_id))
                    if folder_id:
                        logging.info("Found Google Drive ID '{}' for employee {}".format(folder_id, employee_name))
                        upload_success = upload_to_google_drive(output_pdf, folder_id, employee_name, month, year)
                        drive_upload_success = upload_success
                        if not upload_success:
                            warnings.append("Failed to upload to Google Drive")
                            logging.warning("Google Drive upload failed for employee {}".format(employee_name))
                        else:
                            logging.info("Google Drive upload succeeded for employee {}".format(employee_name))
                    else:
                        warnings.append("No Google Drive ID found for employee")
                        logging.warning("No Google Drive ID found for employee: {}. Files will be kept.".format(employee_name))
                        logging.error("Available keys in drive_data: {}".format(list(drive_data[0].keys()) if drive_data else []))
                        # drive_upload_success remains None (not attempted)
                except Exception as e:
                    warnings.append(f"Error processing Google Drive upload: {str(e)}")
                    logging.error("Error processing Google Drive ID: {} {}".format(folder_id, str(e)))
                    drive_upload_success = False  # Mark as failed due to exception

                # If this is part of a multi-month process, collect the PDF but continue with notifications
                if collected_pdfs is not None:
                    collected_pdfs.append(output_pdf)

                # Send email if enabled
                if send_email:
                    try:
                        recipient_email = get_employee_email(placeholders.get("Name"), email_employees)
                        if recipient_email:
                            email_subject = "Salary Slips for {} {} - Bajaj Earths Pvt. Ltd.".format(full_month, full_year) if is_special else "Salary Slip for {} {} - Bajaj Earths Pvt. Ltd.".format(full_month, full_year)
                            months_list = "\n".join(["   -  {} {}".format(month['month'], month['year']) for month in months_data]) if months_data else ""
                            email_body = f"""
                            <html>
                            <body>
                            <p>Dear <b>{placeholders.get('Name')}</b>,</p>
                            <p>Please find attached your <b>salary slip{'s' if is_special else ''}</b> for the following months:</p>
                            <ul>{months_list}</ul>
                            <p>These documents include:</p>
                            <ul>
                            <li>Earnings Breakdown</li>
                            <li>Deductions Summary</li>
                            <li>Net Salary Details</li>
                            </ul>
                            <p>Kindly review the salary slip{'s' if is_special else ''}, and if you have any questions or concerns, please feel free to reach out to the HR department.</p>
                            <p>Thanks & Regards,</p>
                            </body>
                            </html>
                            """
                            logging.info(f"Sending email to {recipient_email}")
                            user_id = session.get('user', {}).get('email') or session.get('user', {}).get('id')
                            if not user_id:
                                errors.append("User session expired. Please log in again.")
                                logging.error("No user ID found in session for email")
                            else:
                                email_success = send_email_gmail_api(user_id, recipient_email, email_subject, email_body, attachment_paths=output_pdf)
                                if email_success == "TOKEN_EXPIRED":
                                    errors.append("Email token expired. Please refresh your credentials.")
                                elif email_success == "USER_NOT_LOGGED_IN":
                                    errors.append("User session expired. Please log in again.")
                                elif email_success == "NO_GMAIL_ACCESS":
                                    errors.append("Gmail access not configured for this user.")
                                elif email_success == "INVALID_RECIPIENT":
                                    errors.append(f"Invalid recipient email address: {recipient_email}")
                                elif email_success == "GMAIL_AUTH_FAILED":
                                    errors.append("Gmail authentication failed. Please check your credentials.")
                                elif email_success == "GMAIL_PERMISSION_DENIED":
                                    errors.append("Gmail permission denied. Please authorize the application.")
                                elif email_success == "GMAIL_API_ERROR":
                                    errors.append("Gmail API error. Please try again later.")
                                elif email_success == "GMAIL_SEND_ERROR":
                                    errors.append(f"Failed to send email to {recipient_email}")
                                elif not email_success:
                                    errors.append(f"Failed to send email to {recipient_email}")
                                else:
                                    logging.info(f"Email sent successfully to {recipient_email}")
                        else:
                            warnings.append(f"No email found for {placeholders.get('Name')}")
                            logging.info(f"No email found for {placeholders.get('Name')}.")
                    except Exception as e:
                        errors.append(f"Error sending email: {str(e)}")
                        logging.error(f"Error sending email: {e}")
                
                # Send WhatsApp message if enabled
                # WhatsApp notifications are now handled by the calling function to prevent duplicates
                # Only send immediate notifications for standalone single month processing
                if send_whatsapp and not is_special and not collected_pdfs:
                    try:
                        contact_name = placeholders.get("Name")
                        whatsapp_number = get_employee_contact(contact_name, contact_employees)
                        if whatsapp_number:
                            try:
                                # Call the imported function directly
                                success = handle_whatsapp_notification(
                                    contact_name=contact_name,
                                    full_month=full_month,
                                    full_year=full_year,
                                    whatsapp_number=whatsapp_number,
                                    file_path=output_pdf,
                                    is_special=False
                                )
                                if success is True:
                                    logging.info(f"WhatsApp notification sent successfully to {contact_name}")
                                elif success == "USER_NOT_LOGGED_IN":
                                    errors.append("User session expired. Please log in again.")
                                elif success == "WHATSAPP_SERVICE_NOT_READY":
                                    errors.append("WhatsApp service is not ready. Please try again later.")
                                elif success == "INVALID_FILE_PATH":
                                    errors.append("Invalid file path for WhatsApp message.")
                                elif success == "INVALID_FILE_PATH_TYPE":
                                    errors.append("Invalid file path type for WhatsApp message.")
                                elif success == "NO_VALID_FILES":
                                    errors.append("No valid files found for WhatsApp message.")
                                elif success == "NO_FILES_FOR_UPLOAD":
                                    errors.append("No files available for WhatsApp upload.")
                                elif success == "WHATSAPP_API_ERROR":
                                    errors.append("WhatsApp API error. Please try again later.")
                                elif success == "WHATSAPP_CONNECTION_ERROR":
                                    errors.append("WhatsApp connection error. Please try again later.")
                                elif success == "WHATSAPP_TIMEOUT_ERROR":
                                    errors.append("WhatsApp timeout error. Please try again later.")
                                elif success == "WHATSAPP_SEND_ERROR":
                                    errors.append(f"Failed to send WhatsApp notification to {contact_name}")
                                else:
                                    errors.append(f"Failed to send WhatsApp notification to {contact_name}")
                                    logging.warning(f"Failed to send WhatsApp notification to {contact_name}")
                            except Exception as e:
                                errors.append(f"Error sending WhatsApp notification to {contact_name}: {e}")
                                logging.error(f"Error sending WhatsApp notification to {contact_name}: {e}")
                        else:
                            warnings.append(f"No WhatsApp number found for {contact_name}")
                    except Exception as e:
                        errors.append(f"Error processing WhatsApp notification: {str(e)}")
                        logging.error(f"Error processing WhatsApp notification: {e}")
                elif send_whatsapp and (is_special or collected_pdfs):
                    # Log that WhatsApp will be handled by calling function
                    logging.info(f"WhatsApp notification for {placeholders.get('Name', 'Unknown')} will be sent by calling function to prevent duplicates")
            else:
                errors.append("Failed to convert DOCX to PDF")
                
        except Exception as e:
            errors.append(f"Error processing salary slip template: {str(e)}")
            logging.error("Error processing salary slip for {}: {}".format(placeholders.get('Name', 'Unknown'), e))
            
    except Exception as e:
        errors.append(f"Critical error in process_salary_slip: {str(e)}")
        logging.error(f"Critical error in process_salary_slip: {e}")
        
    logging.info("Finished process_salary_slip function")
    
    # Return comprehensive result
    result = {
        "success": len(errors) == 0,
        "output_file": output_pdf,
        "errors": errors,
        "warnings": warnings,
        "employee_name": placeholders.get("Name", "Unknown") if 'placeholders' in locals() else "Unknown"
    }
    
    # Add Drive upload status and file paths for conditional deletion
    if 'drive_upload_success' in locals():
        result["drive_upload_success"] = drive_upload_success
    if output_pdf:
        result["pdf_path"] = output_pdf
    if 'output_docx' in locals() and output_docx:
        result["docx_path"] = output_docx
    
    return result

# Generate and process salary slips for multiple employees (batch processing)
def process_salary_slips(template_path, output_dir, employees_data, headers, drive_data, email_employees, contact_employees, month, year, full_month, full_year, send_whatsapp, send_email):
    logging.info("Starting batch process_salary_slips function")
    headers = preprocess_headers(headers)
    
    batch_results = {
        "total_processed": 0,
        "successful": 0,
        "failed": 0,
        "results": [],
        "errors": [],
        "warnings": []
    }
    
    # Process each employee
    for i, employee_data in enumerate(employees_data):
        try:
            logging.info(f"Processing employee {i+1} of {len(employees_data)}")
            
            # Get employee identifier for drive data lookup
            employee_identifier = None
            for j, header in enumerate(headers):
                if header == "Employee\nCode" and j < len(employee_data):
                    employee_identifier = employee_data[j]
                    break
            
            # Process single employee
            result = process_salary_slip(
                template_path=template_path,
                output_dir=output_dir,
                employee_identifier=employee_identifier,
                employee_data=employee_data,
                headers=headers,
                drive_data=drive_data,
                email_employees=email_employees,
                contact_employees=contact_employees,
                month=month,
                year=year,
                full_month=full_month,
                full_year=full_year,
                send_whatsapp=send_whatsapp,
                send_email=send_email,
                is_special=False,
                months_data=None,
                collected_pdfs=None
            )
            
            batch_results["total_processed"] += 1
            batch_results["results"].append(result)
            
            if result["success"]:
                batch_results["successful"] += 1
                logging.info(f"Successfully processed employee: {result['employee_name']}")
            else:
                batch_results["failed"] += 1
                logging.error(f"Failed to process employee: {result['employee_name']}")
                batch_results["errors"].extend([f"{result['employee_name']}: {error}" for error in result["errors"]])
            
            # Collect warnings
            if result["warnings"]:
                batch_results["warnings"].extend([f"{result['employee_name']}: {warning}" for warning in result["warnings"]])
                
        except Exception as e:
            batch_results["total_processed"] += 1
            batch_results["failed"] += 1
            error_msg = f"Critical error processing employee {i+1}: {str(e)}"
            batch_results["errors"].append(error_msg)
            logging.error(error_msg)
            # Continue processing next employee instead of stopping
            continue
            
    logging.info(f"Finished batch processing: {batch_results['successful']}/{batch_results['total_processed']} successful")
    return batch_results

def process_reports(file_path_template):
    
    try:
        file_extension = os.path.splitext(file_path_template)[1].lower()
        file_name = os.path.basename(file_path_template)
        
        # For attachments, just return the filename in square brackets
        if file_extension in ['.pdf', '.png', '.jpg', '.jpeg']:
            return f"[{file_name}]"
            
        if file_extension == '.docx':
            # Read DOCX file
            doc = Document(file_path_template)
            content = []
            
            # Read paragraphs with proper spacing
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    # Add extra newline for paragraphs with specific formatting
                    if para.style.name.startswith('Heading') or para.style.name == 'Title':
                        content.append('\n' + para.text.strip() + '\n')
                    else:
                        content.append(para.text.strip())
            
            # Read tables with proper formatting
            for table in doc.tables:
                content.append('\n')  # Add spacing before table
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        if cell.text.strip():  # Only add non-empty cells
                            row_content.append(cell.text.strip())
                    if row_content:  # Only add non-empty rows
                        content.append(' | '.join(row_content))
                content.append('\n')  # Add spacing after table
            
            # Join content with proper line breaks
            return '\n'.join(content)
            
        elif file_extension in ['.txt', '.csv']:
            # Read text files with proper line breaks
            with open(file_path_template, 'r', encoding='utf-8') as f:
                content = f.read()
                # Ensure proper line breaks and remove extra spaces
                return '\n'.join(line.strip() for line in content.splitlines() if line.strip())
                
        else:
            return f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        logging.error(f"Error reading file {file_path_template}: {e}")
        return f"Error reading file: {str(e)}"

def process_reactor_reports(sheet_id_mapping_data, sheet_recipients_data, table_range_data, input_date, user_id, send_email, send_whatsapp, template_path, output_dir, gspread_client, logger, process_name='reactor-report', google_access_token=None, google_refresh_token=None):
    
    
    
    # Initialize result tracking
    result = {
        "success": False,
        "message": "",
        "generated_files": 0,
        "date_range": "",
        "sheets_processed": 0,
        "notifications_sent": {
            "email": False,
            "whatsapp": False
        },
        "delivery_stats": {
            "total_recipients": 0,
            "successful_deliveries": 0,
            "failed_deliveries": 0,
            "failed_contacts": []
        },
        "output_file": None,
        "errors": [],
        "warnings": []
    }
    
    try:
        REACTOR_CONFIG = {
            'charging_operations': {
                'patterns': [
                    'charging',
                    'charge',
                    'charging (with circulation)',
                    'charging with circulation',
                    'charging operation',
                    'start charging',
                    'reactor charging',
                    'material charging'
                ],
                'date_column_keywords': [
                    'Start Date & Time',
                    'Start Date',
                    'Start Time',
                    'Start Date\n& Time',
                    'Charging Start Date',
                    'Operation Start Date',
                    'Process Start Date',
                    'Begin Date',
                    'Initiation Date'
                ]
            },
            'drain_valve_operations': {
                'patterns': [
                    'ml & a & b drain valve',
                    'ml and a and b drain valve',
                    'drain valve',
                    'ml drain valve',
                    'a & b drain valve',
                    'drain operation',
                    'valve drain',
                    'ml drain',
                    'a b drain',
                    'drain process'
                ],
                'date_column_keywords': [
                    'End Date & Time',
                    'End Date',
                    'End Time',
                    'End Date\n& Time',
                    'Drain End Date',
                    'Operation End Date',
                    'Process End Date',
                    'Completion Date',
                    'Finish Date',
                    'Termination Date'
                ]
            },
            'required_columns': [
                'Particulars',
                'Description',
                'Operation',
                'Process Step'
            ]
        }
        
        # Helper for robust header lookup
        def find_header(headers, name):
            for i, h in enumerate(headers):
                if h.strip().lower() == name.strip().lower():
                    return i
            raise ValueError(f"{name} is not in list: {headers}")

        # Helper for flexible row finding with multiple patterns
        def find_row_by_patterns(data, particulars_idx, patterns, case_sensitive=False):
            
            for row in data[2:]:  # Skip Row 1 (non-data) and Row 2 (headers), start from Row 3 (index 2)
                if len(row) <= particulars_idx:
                    continue
                particulars_text = row[particulars_idx].strip()
                if not case_sensitive:
                    particulars_text = particulars_text.lower()
                
                for pattern in patterns:
                    if case_sensitive:
                        if pattern in particulars_text:
                            return row
                    else:
                        if pattern.lower() in particulars_text:
                            return row
            return None

        # Helper for flexible date column finding
        def find_date_column(headers, date_keywords):
            for keyword in date_keywords:
                try:
                    return find_header(headers, keyword)
                except ValueError:
                    continue
            return None

        # Helper for robust date parsing and comparison
        def normalize_date_for_comparison(date_str, input_date):
            if not date_str or not input_date:
                return None, None
            
            # Common date separators and formats
            date_formats = [
                "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y", "%Y/%m/%d",
                "%d/%m/%y", "%m/%d/%y", "%d-%m-%y", "%y-%m-%d"
            ]
            
            # Clean the date string (remove time part if present)
            date_part = date_str.split()[0] if ' ' in date_str else date_str
            
            # Try to parse the date
            for fmt in date_formats:
                try:
                    parsed_date = datetime.strptime(date_part, fmt)
                    # Normalize to YYYY-MM-DD format
                    normalized_date = parsed_date.strftime("%Y-%m-%d")
                    return normalized_date, parsed_date
                except ValueError:
                    continue
            
            return None, None

        # Helper for debugging sheet structure
        def log_sheet_structure(headers, data, sheet_name, logger):
            
            logger.info(f"Sheet '{sheet_name}' structure:")
            logger.info(f"Headers (Row 2): {headers}")
            logger.info(f"Total data rows: {len(data) - 1}")  # Subtract 1 for headers
            if len(data) > 1:
                logger.info(f"Sample data rows (starting from Row 3):")
                for i, row in enumerate(data[1:4]):  # Show first 3 data rows (Row 3, 4, 5)
                    logger.info(f"  Row {i+3}: {row[:5]}...")  # Show first 5 columns, adjust row number


        # Parse input date
        date_formats = ["%d/%m/%y", "%d/%m/%Y", "%m/%d/%Y", "%Y-%m-%d"]
        dt_input = None
        for fmt in date_formats:
            try:
                dt_input = datetime.strptime(input_date, fmt)
                break
            except Exception:
                continue
        
        if not dt_input:
            result["errors"].append(f"Could not parse input date: {input_date}")
            logger.error(f"Could not parse input date: {input_date}")
            # Continue with partial processing if possible
            result["warnings"].append("Using current date as fallback")
            dt_input = datetime.now()
        
        # Build date->sheet_id mapping
        date_sheet_map = {}
        for row in sheet_id_mapping_data[1:]:
            try:
                date_str = row[0].strip()
                sheet_id = row[1].strip() if len(row) > 1 else None
                if not date_str or not sheet_id:
                    continue
                for fmt in date_formats:
                    try:
                        dt = datetime.strptime(date_str, fmt)
                        date_sheet_map[dt.date()] = sheet_id
                        break
                    except Exception:
                        continue
            except Exception as e:
                result["warnings"].append(f"Error processing sheet mapping row: {e}")
                continue
        
        # Get list of dates: input date and 5 previous dates
        dates_to_check = [dt_input.date() - timedelta(days=i) for i in range(0, 6)]
        sheets_to_process = []
        
        for idx, d in enumerate(dates_to_check):
            sheet_id = date_sheet_map.get(d)
            if not sheet_id:
                if idx == 0:  # Only warn for input date
                    result["warnings"].append(f"No sheet found for input date {d}")
                continue
                
            try:
                spreadsheet = gspread_client.open_by_key(sheet_id)
                worksheet = spreadsheet.sheet1  # Assume first worksheet
                data = worksheet.get_all_values()
                if not data or len(data) < 3:  # Need at least 3 rows: Row 1 (non-data) + Row 2 (headers) + Row 3 (data)
                    logger.warning(f"Sheet for date {d} has insufficient data: {len(data) if data else 0} rows (need at least 3)")
                    result["warnings"].append(f"Sheet for date {d} has insufficient data")
                    continue
                
                # Headers are in Row 2 (index 1), data starts from Row 3 (index 2)
                headers = [h.strip() for h in data[1]]  # Row 2 contains headers
                
                # Log sheet structure for debugging (only for first few sheets to avoid spam)
                if idx < 2:
                    log_sheet_structure(headers, data[1:], f"Date {d}", logger)  # Pass data starting from Row 2
                
                # Find operation description column dynamically (try multiple possible names)
                idx_particulars = None
                for col_name in REACTOR_CONFIG['required_columns']:
                    try:
                        idx_particulars = find_header(headers, col_name)
                        logger.info(f"Found operation column: '{col_name}' at index {idx_particulars}")
                        break
                    except ValueError:
                        continue
                
                if idx_particulars is None:
                    logger.warning(f"No operation description column found in sheet for date {d}. Available columns: {headers}")
                    result["warnings"].append(f"No operation description column found in sheet for date {d}")
                    continue
                
                if idx == 0:
                    # For input date sheet: Always include, no conditions
                    sheets_to_process.append((d, sheet_id, worksheet))
                    logger.info(f"Input date sheet {d} added (no conditions applied)")
                else:
                    # For previous 5 sheets: Only include if drain valve row's end date is blank or matches input_date
                    drain_config = REACTOR_CONFIG['drain_valve_operations']
                    
                    idx_end_date = find_date_column(headers, drain_config['date_column_keywords'])
                    if idx_end_date is None:
                        logger.warning(f"End date column not found in sheet for date {d}. Available columns: {headers}")
                        result["warnings"].append(f"End date column not found in sheet for date {d}")
                        continue
                    
                    drain_valve_row = find_row_by_patterns(data, idx_particulars, drain_config['patterns'])
                    if drain_valve_row:
                        end_date_val = drain_valve_row[idx_end_date].strip()
                        if not end_date_val:
                            sheets_to_process.append((d, sheet_id, worksheet))
                            logger.info(f"Previous date sheet {d} added - drain valve end date is blank")
                        else:
                            # Use robust date comparison
                            end_date_norm, _ = normalize_date_for_comparison(end_date_val, input_date)
                            input_date_norm, _ = normalize_date_for_comparison(input_date, input_date)
                            
                            if end_date_norm and input_date_norm and end_date_norm == input_date_norm:
                                sheets_to_process.append((d, sheet_id, worksheet))
                                logger.info(f"Previous date sheet {d} added - drain valve end date matches: {end_date_norm}")
                            else:
                                logger.info(f"Previous date sheet {d} skipped - drain valve end date '{end_date_val}' doesn't match input date '{input_date}'")
                    else:
                        logger.info(f"No drain valve operation row found in sheet for date {d}")
                        
            except Exception as e:
                logger.error(f"Error processing sheet for date {d}: {e}")
                result["errors"].append(f"Error processing sheet for date {d}: {e}")
                continue
        
        # Log summary of sheets found
        logger.info(f"Processing summary: Found {len(sheets_to_process)} sheets to process out of {len(dates_to_check)} dates checked")
        for d, sheet_id, worksheet in sheets_to_process:
            logger.info(f"  - Date {d}: Sheet ID {sheet_id}")
        
        if not sheets_to_process:
            result["errors"].append("No sheets found to process")
            return result

        # Parse table_range_data header indices robustly (excluding Sheet Name)
        try:
            tr_headers = [h.strip() for h in table_range_data[1]]
            idx_no = next(i for i, h in enumerate(tr_headers) if h.lower() == 'table no.')
            idx_name = next(i for i, h in enumerate(tr_headers) if h.lower() == 'table name')
            idx_start = next(i for i, h in enumerate(tr_headers) if h.lower() == 'start range')
            idx_end = next(i for i, h in enumerate(tr_headers) if h.lower() == 'end range')
        except Exception as e:
            logger.error(f"Error finding headers in table_range_data: {e}")
            result["errors"].append(f"Header error in table_range_data: {e}")
            # Continue with default values if possible
            try:
                idx_no, idx_name, idx_start, idx_end = 0, 1, 2, 3
                result["warnings"].append("Using default column indices for table_range_data")
            except:
                return result

        # Now, for each sheet in sheets_to_process, extract tables and add to doc
        try:
            doc = Document(template_path)
            content_added = False
            first_sheet = True

            for d, sheet_id, worksheet in sheets_to_process:
                try:
                    sheet_name = worksheet.title if hasattr(worksheet, 'title') else str(d)
                    
                    # Write to the first line if it's empty, else add a new paragraph
                    if content_added and not first_sheet:
                        doc.add_page_break()
                    first_sheet = False
                    
                    if not doc.paragraphs or not doc.paragraphs[0].text.strip():
                        para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
                        para.text = f"Reactor: {sheet_name}"
                    else:
                        para = doc.add_paragraph(f"Reactor: {sheet_name}")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(14)

                    # Extract tables as per table_range_data
                    table_defs = []
                    try:
                        for row in table_range_data[2:]:
                            if len(row) > max(idx_no, idx_name, idx_start, idx_end) and row[idx_name].strip() and row[idx_start].strip() and row[idx_end].strip():
                                table_defs.append({
                                    'no': int(row[idx_no]) if row[idx_no].strip().isdigit() else 0,
                                    'name': row[idx_name].strip(),
                                    'start': row[idx_start].strip(),
                                    'end': row[idx_end].strip()
                                })
                        table_defs.sort(key=lambda x: x['no'])
                    except Exception as e:
                        result["warnings"].append(f"Error parsing table definitions: {e}")
                        continue

                    for i, table_def in enumerate(table_defs):
                        try:
                            # Get data from the specified range
                            data = worksheet.get(f"{table_def['start']}:{table_def['end']}")
                            if not data or len(data) < 2:  # Need at least header + 1 data row
                                logger.warning(f"Table {table_def['name']} has insufficient data: {len(data) if data else 0} rows")
                                result["warnings"].append(f"Table {table_def['name']} has insufficient data")
                                continue
                           
                            if should_start_table_on_new_page(doc, len(data), len(table_def['name']), logger):
                                # Add page break before table to ensure it starts on new page
                                doc.add_page_break()
                                logger.info(f"Added page break before table '{table_def['name']}' to prevent breaking")
                            
                            # Add table name with formatting
                            doc.add_paragraph()
                            table_name_para = doc.add_paragraph(table_def['name'])
                            table_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in table_name_para.runs:
                                run.bold = True
                                run.font.size = Pt(11)
                            
                            # Create table with proper formatting
                            max_cols = max(len(row) for row in data)
                            table = doc.add_table(rows=len(data), cols=max_cols)
                            
                            # Fill table data with proper formatting
                            for row_idx, row in enumerate(data):
                                for col_idx, cell_value in enumerate(row):
                                    if col_idx < max_cols:
                                        cell = table.cell(row_idx, col_idx)
                                        cell.text = str(cell_value)
                            
                            # Apply professional formatting to the table (includes page break prevention)
                            format_table_professional(table, is_header=True, logger=logger)
                            
                            # Add 2 lines gap between tables (but not after the last table)
                            if i < len(table_defs) - 1:
                                doc.add_paragraph()
                                
                            content_added = True
                            
                        except Exception as e:
                            logger.error(f"Error extracting table {table_def['name']} from {sheet_name}: {e}")
                            result["errors"].append(f"Error extracting table {table_def['name']} from {sheet_name}: {e}")
                            # Continue with next table instead of failing completely
                            continue
                            
                    result["sheets_processed"] += 1
                    
                except Exception as e:
                    logger.error(f"Error processing sheet {sheet_id}: {e}")
                    result["errors"].append(f"Error processing sheet {sheet_id}: {e}")
                    continue
            
            # Clean up completely blank pages only
            def remove_completely_blank_pages(doc):
                # Check if the last paragraph is a page break that creates a blank page
                if doc.paragraphs:
                    last_para = doc.paragraphs[-1]
                    # If the last paragraph has no text and contains only a page break
                    if not last_para.text.strip() and last_para.runs:
                        last_run = last_para.runs[0]
                        # Check if this run contains only a page break (no other content)
                        if (hasattr(last_run, '_element') and 
                            last_run._element.xml.count('<w:br') > 0 and
                            not last_run.text.strip()):
                            # Remove this paragraph to eliminate the blank page
                            last_para._element.getparent().remove(last_para._element)
                            logger.info("Removed completely blank page at the end of document")
            
            # Apply cleanup for completely blank pages only
            remove_completely_blank_pages(doc)
            
            # Save the document
            output_filename = f"reactor_report_{input_date.replace('/', '_')}.docx"
            output_path = os.path.join(output_dir, output_filename)
            doc.save(output_path)
            
            # Convert to PDF
            pdf_filename = f"reactor_report_{input_date.replace('/', '_')}.pdf"
            pdf_path = os.path.join(output_dir, pdf_filename)
            logger.info(f"Generated PDF filename: {pdf_filename}")
            logger.info(f"Generated PDF path: {pdf_path}")
            logger.info(f"PDF file exists: {os.path.exists(pdf_path)}")
            pdf_conversion_success = convert_docx_to_pdf(output_path, pdf_path)
            if pdf_conversion_success:
                logger.info("Successfully converted DOCX to PDF")
                logger.info(f"PDF file exists after conversion: {os.path.exists(pdf_path)}")
                result["output_file"] = pdf_path
            else:
                logger.warning("PDF conversion failed, using DOCX file")
                result["warnings"].append("PDF conversion failed, using DOCX file")
                result["output_file"] = output_path

            # Store both file paths for later cleanup
            result["docx_path"] = output_path
            result["pdf_path"] = pdf_path  # Store path regardless of conversion success

            result["generated_files"] = 1
            result["date_range"] = input_date

        except Exception as e:
            logger.error(f"Error creating document: {e}")
            result["errors"].append(f"Error creating document: {e}")
            return result

        # Send email notifications
        if send_email and result["output_file"]:
            try:
                # Parse recipients from sheet_recipients_data
                headers = [h.strip() for h in sheet_recipients_data[0]]
                to_idx = headers.index('Email ID - To') if 'Email ID - To' in headers else None
                cc_idx = headers.index('Email ID - CC') if 'Email ID - CC' in headers else None
                bcc_idx = headers.index('Email ID - BCC') if 'Email ID - BCC' in headers else None
                
                recipients_to = []
                recipients_cc = []
                recipients_bcc = []
                
                for row in sheet_recipients_data[1:]:
                    try:
                        if to_idx is not None and len(row) > to_idx and row[to_idx].strip():
                            email = row[to_idx].strip()
                            if is_valid_email(email):
                                recipients_to.append(email)
                        if cc_idx is not None and len(row) > cc_idx and row[cc_idx].strip():
                            email = row[cc_idx].strip()
                            if is_valid_email(email):
                                recipients_cc.append(email)
                        if bcc_idx is not None and len(row) > bcc_idx and row[bcc_idx].strip():
                            email = row[bcc_idx].strip()
                            if is_valid_email(email):
                                recipients_bcc.append(email)
                    except Exception as e:
                        result["warnings"].append(f"Error processing recipient row: {e}")
                        continue
                
                if not recipients_to:
                    result["warnings"].append("No valid recipient emails found in Recipients sheet")
                else:
                    email_subject = "Reactor Report - Daily Operations Summary"
                    email_body = f"""
                        <html>
                        <body>
                        <h2>Reactor Report</h2>
                        <p>Please find attached the reactor report for the period from {input_date} to {input_date}.</p>
                        <p>This report contains snapshots of all reactor operations data for the specified period.</p>
                        <br>
                        <p>Best regards,<br>Reactor Automation System</p>
                        </body>
                        </html>
                        """
                    # Send email to each recipient individually to avoid syntax errors
                    success_count = 0
                    total_email_recipients = len(recipients_to)
                    
                    # Update delivery stats for email recipients
                    result["delivery_stats"]["total_recipients"] += total_email_recipients
                    
                    for recipient in recipients_to:
                        try:
                            if process_name == "kr_reactor-report":
                                
                                
                                # If Google tokens are provided, use them; otherwise, function will attempt DB/session tokens
                                logger.info(f"KR reactor: Sending via Gmail API only (no SMTP fallback)")
                                logger.info(f"Email attachment path: {result['output_file']}")
                                logger.info(f"Attachment file exists: {os.path.exists(result['output_file'])}")
                                success = send_email_gmail_api(
                                    user_email=user_id,
                                    recipient_email=recipient,
                                    subject=email_subject,
                                    body=email_body,
                                    attachment_paths=[result["output_file"]],
                                    cc=','.join(recipients_cc) if recipients_cc else None,
                                    bcc=','.join(recipients_bcc) if recipients_bcc else None,
                                    access_token=google_access_token,
                                    refresh_token=google_refresh_token
                                )
                            else:
                                success = send_email_gmail_api(
                                    user_email=user_id,
                                    recipient_email=recipient,
                                    subject=email_subject,
                                    body=email_body,
                                    attachment_paths=[result["output_file"]],
                                    cc=','.join(recipients_cc) if recipients_cc else None,
                                    bcc=','.join(recipients_bcc) if recipients_bcc else None
                                )
                            
                            # Track email delivery results
                            if success is True:
                                success_count += 1
                                result["delivery_stats"]["successful_deliveries"] += 1
                                logger.info(f"Email sent successfully to {recipient}")
                            else:
                                result["delivery_stats"]["failed_deliveries"] += 1
                                failure_reason = "Unknown error"
                                
                                if success == "USER_NOT_LOGGED_IN":
                                    failure_reason = "User session expired"
                                    result["errors"].append(f"User session expired for {recipient}. Please log in again.")
                                    logger.error(f"User session expired for {recipient}")
                                elif success == "NO_GMAIL_ACCESS":
                                    failure_reason = "Gmail access not configured"
                                    result["errors"].append(f"Gmail access not configured for {recipient}. Please configure Google OAuth.")
                                    logger.error(f"Gmail access not configured for {recipient}")
                                elif success == "INVALID_RECIPIENT":
                                    failure_reason = "Invalid recipient email address"
                                    result["errors"].append(f"Invalid recipient email address: {recipient}")
                                    logger.error(f"Invalid recipient email address: {recipient}")
                                elif success == "GMAIL_AUTH_FAILED":
                                    failure_reason = "Gmail authentication failed"
                                    result["errors"].append(f"Gmail authentication failed for {recipient}. Please check your credentials.")
                                    logger.error(f"Gmail authentication failed for {recipient}")
                                elif success == "GMAIL_PERMISSION_DENIED":
                                    failure_reason = "Gmail permission denied"
                                    result["errors"].append(f"Gmail permission denied for {recipient}. Please authorize the application.")
                                    logger.error(f"Gmail permission denied for {recipient}")
                                elif success == "GMAIL_API_ERROR":
                                    failure_reason = "Gmail API error"
                                    result["errors"].append(f"Gmail API error for {recipient}. Please try again later.")
                                    logger.error(f"Gmail API error for {recipient}")
                                elif success == "GMAIL_SEND_ERROR":
                                    failure_reason = "Email send error"
                                    result["errors"].append(f"Failed to send email to {recipient}")
                                    logger.error(f"Failed to send email to {recipient}")
                                else:
                                    failure_reason = f"Unknown error: {success}"
                                    result["errors"].append(f"Failed to send email to {recipient}")
                                    logger.error(f"Failed to send email to {recipient}")
                                
                                # Add to failed contacts list
                                result["delivery_stats"]["failed_contacts"].append({
                                    "name": recipient,
                                    "contact": recipient,
                                    "reason": failure_reason
                                })
                                
                        except Exception as e:
                            result["delivery_stats"]["failed_deliveries"] += 1
                            result["delivery_stats"]["failed_contacts"].append({
                                "name": recipient,
                                "contact": recipient,
                                "reason": f"Exception: {str(e)}"
                            })
                            result["errors"].append(f"Error sending email to {recipient}: {e}")
                            logger.error(f"Error sending email to {recipient}: {e}")
                    
                    if success_count > 0:
                        result["notifications_sent"]["email"] = True
                        logger.info(f"Successfully sent emails to {success_count} recipients")
                    else:
                        result["errors"].append("Failed to send email to any recipients")
                        
            except Exception as e:
                logger.error(f"Error processing email notifications: {e}")
                result["errors"].append(f"Error processing email notifications: {e}")

        # Send WhatsApp messages if enabled
        if send_whatsapp and result["output_file"]:
            try:
                # Track delivery statistics for WhatsApp notifications
                whatsapp_result = handle_reactor_report_notification_with_stats(
                    recipients_data=sheet_recipients_data,
                    input_date=input_date,
                    file_path=result["output_file"],
                    sheets_processed=result["sheets_processed"],
                    user_email=user_id
                )
                
                if isinstance(whatsapp_result, dict) and "delivery_stats" in whatsapp_result:
                    # Update delivery stats from WhatsApp notifications
                    result["delivery_stats"]["total_recipients"] += whatsapp_result["delivery_stats"]["total_recipients"]
                    result["delivery_stats"]["successful_deliveries"] += whatsapp_result["delivery_stats"]["successful_deliveries"]
                    result["delivery_stats"]["failed_deliveries"] += whatsapp_result["delivery_stats"]["failed_deliveries"]
                    result["delivery_stats"]["failed_contacts"].extend(whatsapp_result["delivery_stats"]["failed_contacts"])
                    
                    if whatsapp_result["delivery_stats"]["successful_deliveries"] > 0:
                        result["notifications_sent"]["whatsapp"] = True
                        logger.info("Reactor report WhatsApp notifications sent successfully")
                    else:
                        result["warnings"].append("No WhatsApp notifications were sent successfully.")
                        logger.warning("No WhatsApp notifications were sent successfully")
                else:
                    # Handle legacy return values
                    success = whatsapp_result
                    if success is True:
                        result["notifications_sent"]["whatsapp"] = True
                        logger.info("Reactor report WhatsApp notifications sent successfully")
                    elif success == "USER_NOT_LOGGED_IN":
                        result["errors"].append("User session expired. Please log in again.")
                        logger.error("User session expired for WhatsApp notifications")
                    elif success == "WHATSAPP_SERVICE_NOT_READY":
                        result["warnings"].append("WhatsApp service is not ready. Please try again later.")
                        logger.warning("WhatsApp service is not ready")
                    elif success == "NO_RECIPIENTS_DATA":
                        result["warnings"].append("No recipients data provided for WhatsApp notifications.")
                        logger.warning("No recipients data provided for WhatsApp notifications")
                    elif success == "MISSING_REQUIRED_COLUMNS":
                        result["warnings"].append("Required columns for WhatsApp notifications not found.")
                        logger.warning("Required columns for WhatsApp notifications not found")
                    elif success == "NO_SUCCESSFUL_NOTIFICATIONS":
                        result["warnings"].append("No WhatsApp notifications were sent successfully.")
                        logger.warning("No WhatsApp notifications were sent successfully")
                    elif success == "REACTOR_NOTIFICATION_ERROR":
                        result["warnings"].append("Error processing WhatsApp notifications.")
                        logger.warning("Error processing WhatsApp notifications")
                    else:
                        result["warnings"].append("Some or all reactor report WhatsApp notifications failed")
                        logger.warning("Some or all reactor report WhatsApp notifications failed")
            except Exception as e:
                logger.error(f"Error processing WhatsApp notifications: {e}")
                result["errors"].append(f"Error processing WhatsApp notifications: {e}")

        # Set final success status
        if result["generated_files"] > 0:
            result["success"] = True
            if not result["errors"]:
                result["message"] = "Reactor reports generated and sent successfully"
            else:
                result["message"] = f"Reactor reports generated with {len(result['errors'])} errors"
        else:
            result["message"] = "Failed to generate reactor reports"
            
        # Send log report to the user who generated the report (only if notifications were sent)
        if (send_email or send_whatsapp) and result["generated_files"] > 0:
            try:
                # Import the log report function from app.py
                from app import send_log_report_to_user
                send_log_report_to_user(user_id, result["delivery_stats"], send_email, send_whatsapp, logger)
                logger.info("Log report sent successfully to user")
            except Exception as e:
                logger.error(f"Error sending log report to user: {e}")
        
        # Upload PDF to Google Drive before deletion (only if document was generated and notifications were attempted)
        if result["generated_files"] > 0 and (send_email or send_whatsapp) and result.get("pdf_path"):
            try:
                # Base folder ID for reactor reports
                base_folder_id = "1cuL5gdl5GncegK2-FItKux7pg-D2sT--"
                
                logger.info(f"Uploading reactor report PDF to Google Drive (base folder: {base_folder_id})")
                upload_success, file_id, folder_id, upload_error = upload_reactor_report_to_drive(
                    pdf_path=result["pdf_path"],
                    base_folder_id=base_folder_id,
                    input_date=input_date,
                    logger=logger
                )
                
                if upload_success:
                    logger.info(f"Successfully uploaded reactor report to Google Drive. File ID: {file_id}, Month Folder ID: {folder_id}")
                    result["drive_upload"] = {
                        "success": True,
                        "file_id": file_id,
                        "folder_id": folder_id
                    }
                else:
                    logger.warning(f"Failed to upload reactor report to Google Drive: {upload_error}")
                    result["warnings"].append(f"Google Drive upload failed: {upload_error}")
                    result["drive_upload"] = {
                        "success": False,
                        "error": upload_error
                    }
            except Exception as e:
                logger.error(f"Error during Google Drive upload: {e}")
                result["warnings"].append(f"Error during Google Drive upload: {e}")
                result["drive_upload"] = {
                    "success": False,
                    "error": str(e)
                }
        
        # Delete generated documents after notifications are completed
        # Only delete if document was generated, notifications were attempted, AND Drive upload succeeded
        if result["generated_files"] > 0 and (send_email or send_whatsapp):
            try:
                # Check Drive upload status
                drive_upload_status = result.get("drive_upload", {}).get("success")
                
                # Use generic deletion function
                deletion_result = delete_generated_files(
                    file_paths=result,
                    drive_upload_success=drive_upload_status,
                    logger=logger
                )
                
                if deletion_result['success']:
                    logger.info(f"File cleanup completed: Deleted {deletion_result['total_deleted']} file(s) after successful Drive upload")
                elif deletion_result['skipped_reason']:
                    logger.info(f"File deletion skipped: {deletion_result['skipped_reason']}")
                    result["warnings"].append(deletion_result['skipped_reason'])
                else:
                    if deletion_result['failed_files']:
                        for failed in deletion_result['failed_files']:
                            warning_msg = f"Failed to delete {failed['path']}: {failed['reason']}"
                            logger.warning(warning_msg)
                            result["warnings"].append(warning_msg)
                    
            except Exception as e:
                logger.error(f"Error during document cleanup: {e}")
                result["warnings"].append(f"Error during document cleanup: {e}")
            
    except Exception as e:
        logger.error(f"Critical error in process_reactor_reports: {e}")
        result["errors"].append(f"Critical error: {e}")
        result["message"] = f"Critical error in reactor report processing: {e}"
    
    return result

# Google Sheets Material Data Integration Functions

# Helper functions for plant data (now using data from frontend)
def get_plant_name_by_id(plant_id, plant_data):
    """Get plant name by plant ID from plant data"""
    for plant in plant_data:
        if plant.get('material_sheet_id') == plant_id:
            return plant.get('name', 'Unknown Plant')
    return 'Unknown Plant'

def get_plant_document_name_by_id(plant_id, plant_data):
    """Get plant document name by plant ID from plant data"""
    for plant in plant_data:
        if plant.get('material_sheet_id') == plant_id:
            return plant.get('document_name', 'UNKNOWN')
    return 'UNKNOWN'

def generate_material_key(category, sub_category, material_name, specifications):
    """
    Generate unique key for material comparison
    Uses all 4 fields: category|subCategory|materialName|specifications
    Empty values are treated as empty strings for consistency
    """
    # Normalize empty values to empty strings (not None)
    sub_cat = str(sub_category).strip() if sub_category else ''
    specs = str(specifications).strip() if specifications else ''
    cat = str(category).strip() if category else ''
    mat_name = str(material_name).strip() if material_name else ''
    
    return f"{cat}|{sub_cat}|{mat_name}|{specs}"

def get_sheet_name_by_id(plant_id, plant_data):
    """Get sheet name by plant ID from plant data"""
    for plant in plant_data:
        if plant.get('material_sheet_id') == plant_id:
            sheet_name = plant.get('sheet_name', 'Material List')
            
            # Handle both old string format and new object format for backward compatibility
            if isinstance(sheet_name, str):
                return sheet_name
            elif isinstance(sheet_name, dict):
                # Always return 'Material List' for MaterialList sheet
                return sheet_name.get('MaterialList', 'Material List')
            
            return 'Material List'
    return 'Material List'

def get_plant_material_data_from_sheets(plant_id, plant_data):
    """Fetch material data from Google Sheets for a specific plant"""
    try:
        if not plant_id or not plant_data:
            logging.error("No plant_id or plant_data provided")
            return {}
        
        logging.info(f"Attempting to fetch material data for plant_id: {plant_id}")
        
        # Initialize gspread client
        try:
            client = gspread.authorize(creds)
            logging.info("Successfully authorized gspread client")
        except Exception as e:
            logging.error(f"Failed to authorize gspread client: {e}")
            raise
        
        # Open the spreadsheet by ID
        try:
            spreadsheet = client.open_by_key(plant_id)
            logging.info(f"Successfully opened spreadsheet: {spreadsheet.title}")
        except Exception as e:
            logging.error(f"Failed to open spreadsheet with ID {plant_id}: {e}")
            raise
        
        # Try to find the correct worksheet using plant data from frontend
        worksheet = None
        sheet_name = get_sheet_name_by_id(plant_id, plant_data)
        try:
            worksheet = spreadsheet.worksheet(sheet_name)
            logging.info(f"Found '{sheet_name}' worksheet")
        except gspread.WorksheetNotFound:
            # If specified sheet not found, use the first worksheet
            worksheet = spreadsheet.sheet1
            logging.warning(f"'{sheet_name}' sheet not found, using first worksheet: {worksheet.title}")
        except Exception as e:
            logging.error(f"Error accessing worksheet: {e}")
            raise
        
        # Get all data from the sheet
        try:
            all_data = worksheet.get_all_values()
            logging.info(f"Retrieved {len(all_data)} rows from worksheet")
        except Exception as e:
            logging.error(f"Error getting data from worksheet: {e}")
            raise
        
        if not all_data or len(all_data) < 3:
            logging.warning(f"Insufficient data in sheet for plant {plant_id}. Found {len(all_data) if all_data else 0} rows, need at least 3")
            return {}
        
        # Headers are on row 2 (index 1), data starts from row 3 (index 2)
        headers = [header.strip() for header in all_data[1]]  # Row 2 contains headers
        logging.info(f"Found headers: {headers}")
        
        # Expected headers: Category, Sub Category, Specifications, Material Name, UOM, Initial\nQuantity
        expected_headers = ['Category', 'Sub Category', 'Specifications', 'Material Name', 'UOM', 'Initial\nQuantity']
        
        # Find column indices for expected headers
        header_indices = {}
        for expected_header in expected_headers:
            for i, header in enumerate(headers):
                if expected_header.lower() in header.lower():
                    header_indices[expected_header] = i
                    logging.info(f"Found header '{expected_header}' at index {i} (actual: '{header}')")
                    break
        
        logging.info(f"Header indices found: {header_indices}")
        logging.info(f"Looking for 'Initial\\nQuantity' header: {'Initial\\nQuantity' in header_indices}")
        if 'Initial\nQuantity' not in header_indices:
            logging.warning("'Initial\\nQuantity' header not found! Available headers: " + str(headers))
        
        # Check for required headers (Category, Material Name, UOM are compulsory)
        required_headers = ['Category', 'Material Name', 'UOM']
        missing_required = [h for h in required_headers if h not in header_indices]
        
        if missing_required:
            logging.error(f"Required headers not found in sheet: {missing_required}. Found headers: {headers}")
            return {}
        
        # Process data rows (skip rows 1 and 2, start from row 3 - index 2)
        material_data = {}
        skipped_rows = []
        skipped_reasons = {}
        total_processed = 0
        
        for row_index, row in enumerate(all_data[2:], start=3):  # Start from row 3 (index 2), but row_index starts from 3
            total_processed += 1
            
            if len(row) < max(header_indices.values()) + 1:
                skipped_rows.append({
                    'rowNumber': row_index,
                    'reason': 'Incomplete row data'
                })
                skipped_reasons['Incomplete row data'] = skipped_reasons.get('Incomplete row data', 0) + 1
                continue  # Skip incomplete rows
            
            try:
                category = row[header_indices.get('Category', 0)].strip()
                sub_category = row[header_indices.get('Sub Category', 1)].strip() if 'Sub Category' in header_indices else ''
                specifications = row[header_indices.get('Specifications', 2)].strip() if 'Specifications' in header_indices else ''
                material_name = row[header_indices.get('Material Name', 3)].strip()
                uom = row[header_indices.get('UOM', 4)].strip()
                initial_quantity = row[header_indices.get('Initial\nQuantity', 5)].strip() if 'Initial\nQuantity' in header_indices else '0'
                
                # Debug logging for initial quantity
                logging.info(f"Row {row_index}: initial_quantity='{initial_quantity}', has_header={('Initial\nQuantity' in header_indices)}")
                
                # Validate compulsory fields and track specific missing fields
                missing_fields = []
                if not category:
                    missing_fields.append('Category')
                if not material_name:
                    missing_fields.append('Material Name')
                if not uom:
                    missing_fields.append('UOM')
                
                if missing_fields:
                    reason = f"Missing: {', '.join(missing_fields)}"
                    logging.warning(f"Skipping row {row_index} due to missing compulsory fields: Category='{category}', Material Name='{material_name}', UOM='{uom}'")
                    skipped_rows.append({
                        'rowNumber': row_index,
                        'reason': reason,
                        'category': category,
                        'materialName': material_name,
                        'uom': uom
                    })
                    skipped_reasons[reason] = skipped_reasons.get(reason, 0) + 1
                    continue  # Skip rows without essential data
                
                # Set default quantity to 0 if not provided
                if not initial_quantity or initial_quantity.strip() == '':
                    initial_quantity = '0'
                
                # Initialize category if not exists
                if category not in material_data:
                    material_data[category] = {
                        'subCategories': set(),
                        'specifications': set(),
                        'materialNames': []
                    }
                
                # Add sub category
                if sub_category:
                    material_data[category]['subCategories'].add(sub_category)
                
                # Add specifications
                if specifications:
                    material_data[category]['specifications'].add(specifications)
                
                # Add material name with details (following standard field sequence)
                material_info = {
                    'subCategory': sub_category,
                    'name': material_name,
                    'specifications': specifications,
                    'uom': uom,
                    'initialQuantity': initial_quantity
                }
                material_data[category]['materialNames'].append(material_info)
                
            except Exception as e:
                logging.warning(f"Error processing row: {e}")
                continue
        
        # Convert sets to lists and restructure materialNames for KR_PlaceOrder.jsx compatibility
        for category, category_data in material_data.items():
            sub_categories_list = list(category_data['subCategories'])
            specifications_list = list(category_data['specifications'])
            category_data['subCategories'] = sub_categories_list
            category_data['specifications'] = specifications_list
            
            materials_list = category_data['materialNames']
            
            def _normalize(value):
                return (value or '').strip().lower()
            
            has_subcategories = any(sc and sc.strip() for sc in sub_categories_list)
            has_specifications = any(spec and spec.strip() for spec in specifications_list)
            
            if has_subcategories:
                nested_materials = {}
                for sub_cat in sub_categories_list:
                    if not sub_cat or not sub_cat.strip():
                        continue
                    materials_for_subcat = [
                        mat for mat in materials_list
                        if _normalize(mat.get('subCategory')) == _normalize(sub_cat)
                    ]
                    if materials_for_subcat:
                        nested_materials[sub_cat] = materials_for_subcat
                
                # Handle materials without subcategory
                materials_without_subcat = [
                    mat for mat in materials_list
                    if not mat.get('subCategory') or not mat.get('subCategory').strip()
                ]
                if materials_without_subcat:
                    nested_materials[''] = materials_without_subcat
                
                category_data['materialNames'] = nested_materials or materials_list
            elif has_specifications:
                nested_materials = {}
                for specification in specifications_list:
                    if not specification or not specification.strip():
                        continue
                    materials_for_specification = [
                        mat for mat in materials_list
                        if _normalize(mat.get('specifications')) == _normalize(specification)
                    ]
                    if materials_for_specification:
                        nested_materials[specification] = materials_for_specification
                
                category_data['materialNames'] = nested_materials or materials_list
            else:
                category_data['materialNames'] = materials_list
        
        # Calculate total materials processed successfully
        total_materials = 0
        for category_data in material_data.values():
            material_names = category_data.get('materialNames', [])
            if isinstance(material_names, list):
                # Simple array structure
                total_materials += len(material_names)
            elif isinstance(material_names, dict):
                # Nested structure - count all materials in nested objects
                for sub_cat_or_spec, materials in material_names.items():
                    if isinstance(materials, list):
                        total_materials += len(materials)
                    elif isinstance(materials, dict):
                        # Further nested (subCategory -> specifications -> materials)
                        for spec_or_mat, mat_list in materials.items():
                            if isinstance(mat_list, list):
                                total_materials += len(mat_list)
        
        logging.info(f"Successfully fetched material data for plant {plant_id}: {len(material_data)} categories, {total_materials} materials processed, {len(skipped_rows)} rows skipped")
        
        return {
            'material_data': material_data,
            'total_processed': total_processed,
            'total_materials': total_materials,
            'skipped_rows': skipped_rows,
            'skipped_count': len(skipped_rows),
            'skipped_reasons': skipped_reasons
        }
        
    except Exception as e:
        logging.error(f"Error fetching material data from Google Sheets: {e}")
        logging.error(f"Error type: {type(e).__name__}")
        logging.error(f"Error details: {str(e)}")
        logging.error(f"Full traceback: {traceback.format_exc()}")
        return {}

def get_designation_from_authority_sheet(factory, given_by_name, logger=None):
    """
    Fetch designation from authority sheet based on givenBy name (authority name).
    
    Args:
        factory: Factory identifier (e.g., 'KR', 'GB', 'HB', etc.)
        given_by_name: Name of the authority (Given By name)
        logger: Optional logger instance
    
    Returns:
        str: Designation value or 'N/A' if not found
    """
    log = logger if logger else logging
    
    try:
        # Lazy import to avoid circular dependency
        import sys
        if 'app' in sys.modules:
            PLANT_DATA = sys.modules['app'].PLANT_DATA
        else:
            from app import PLANT_DATA
        
        # Find plant configuration
        plant_config = None
        for plant in PLANT_DATA:
            if plant.get("document_name") == factory:
                plant_config = plant
                break
        
        if not plant_config:
            log.warning(f"Plant configuration not found for factory: {factory}")
            return 'N/A'
        
        # Get sheet ID and sheet name for authority list
        sheet_id = plant_config.get("material_sheet_id")
        if not sheet_id:
            log.warning(f"No material_sheet_id found for factory: {factory}")
            return 'N/A'
        
        # Get authority sheet name (usually "List")
        sheet_names = plant_config.get("sheet_name", {})
        authority_sheet_name = sheet_names.get("AuthorityList", "List")
        
        # Import fetch function
        from Utils.fetch_data import fetch_google_sheet_data
        
        # Fetch authority sheet data
        authority_data = fetch_google_sheet_data(sheet_id, authority_sheet_name)
        
        if not authority_data or len(authority_data) < 2:
            log.warning(f"No authority data found in sheet {authority_sheet_name} for factory {factory}")
            return 'N/A'
        
        # Sheet structure:
        # - First row is blank
        # - Second row (index 1) contains headers
        # - Data starts from third row (index 2)
        headers = [str(h).strip() for h in authority_data[1]] if len(authority_data) > 1 else []
        data_rows = authority_data[2:] if len(authority_data) > 2 else []
        
        # Find column indices for "Given By" and "Designation"
        def find_column_index(possible_terms):
            for term in possible_terms:
                for idx, h in enumerate(headers):
                    if str(h).strip().lower() == term.lower():
                        return idx
            return None
        
        given_by_col_idx = find_column_index(["given by", "authority name", "name"])
        designation_col_idx = find_column_index(["designation"])
        
        if given_by_col_idx is None:
            log.warning(f"Could not find 'Given By' column in authority sheet for factory {factory}")
            return 'N/A'
        
        if designation_col_idx is None:
            log.warning(f"Could not find 'Designation' column in authority sheet for factory {factory}")
            return 'N/A'
        
        # Search for the given_by_name and get its designation
        given_by_name_clean = str(given_by_name).strip() if given_by_name else ""
        if not given_by_name_clean:
            return 'N/A'
        
        for row in data_rows:
            if not row or len(row) <= max(given_by_col_idx, designation_col_idx):
                continue
            
            row_given_by = str(row[given_by_col_idx]).strip()
            row_designation = str(row[designation_col_idx]).strip() if designation_col_idx < len(row) else ""
            
            # Case-insensitive comparison
            if row_given_by.lower() == given_by_name_clean.lower():
                return row_designation if row_designation else 'N/A'
        
        log.warning(f"Designation not found for authority '{given_by_name}' in factory {factory}")
        return 'N/A'
        
    except Exception as e:
        log.error(f"Error fetching designation from authority sheet for factory {factory}, givenBy {given_by_name}: {e}")
        return 'N/A'

def sync_plant_material_to_firebase(plant_id, plant_name, plant_data, sync_description, sync_timestamp, synced_by):
    """
    Smart sync material data from Google Sheets to Firebase
    Only adds new materials, preserves existing ones based on exact 4-field match
    """
    try:
        # First, run migration to ensure data is in nested structure
        from Utils.firebase_utils import migrate_individual_documents_to_nested_structure
        migration_result = migrate_individual_documents_to_nested_structure()
        if migration_result['success'] and migration_result['migrated_count'] > 0:
            logging.info(f"Migration completed before sync: {migration_result['message']}")
        
        # Then, get the data from Google Sheets
        sheet_result = get_plant_material_data_from_sheets(plant_id, plant_data)
        
        if not sheet_result or not sheet_result.get('material_data'):
            return {
                'success': False,
                'message': 'No data found in Google Sheets'
            }
        
        # Extract data from the result
        sheet_data = sheet_result['material_data']
        total_processed = sheet_result.get('total_processed', 0)
        total_materials_in_sheets = sheet_result.get('total_materials', 0)
        skipped_rows = sheet_result.get('skipped_rows', [])
        skipped_count = sheet_result.get('skipped_count', 0)
        skipped_reasons = sheet_result.get('skipped_reasons', {})
        
        # Get the correct document name for Firebase storage
        document_name = get_plant_document_name_by_id(plant_id, plant_data)
        logging.info(f"Using document name '{document_name}' for plant '{plant_name}' (ID: {plant_id})")
        
        # Step 1: Fetch existing materials from Firebase
        plant_ref = db.collection('MATERIAL').document(document_name)
        plant_doc = plant_ref.get()
        
        existing_materials_map = {}
        existing_materials_list = []
        
        if plant_doc.exists:
            existing_data = plant_doc.to_dict()
            
            # Try to get materials from subcollection first (new format)
            materials_subcollection = plant_ref.collection('materials')
            materials_docs = materials_subcollection.get()
            
            existing_materials_list = []
            
            if materials_docs:
                # New format: materials in subcollection
                for mat_doc in materials_docs:
                    mat_data = mat_doc.to_dict()
                    existing_materials_list.append(mat_data)
                    
                    # Build lookup map with unique keys
                    key = generate_material_key(
                        mat_data.get('category', ''),
                        mat_data.get('subCategory', ''),
                        mat_data.get('materialName', ''),
                        mat_data.get('specifications', '')
                    )
                    existing_materials_map[key] = mat_data
                
                logging.info(f"Found {len(existing_materials_list)} existing materials in subcollection for {document_name}")
            else:
                # Old format: materials as array in main document
                existing_materials_list = existing_data.get('materials', [])
                
                for material in existing_materials_list:
                    # Build lookup map with unique keys
                    key = generate_material_key(
                        material.get('category', ''),
                        material.get('subCategory', ''),
                        material.get('materialName', ''),
                        material.get('specifications', '')
                    )
                    existing_materials_map[key] = material
                
                logging.info(f"Found {len(existing_materials_list)} existing materials in main document for {document_name}")
            
            logging.info(f"Sample existing material keys: {list(existing_materials_map.keys())[:3]}")
        else:
            logging.info(f"No existing materials found in Firebase for {document_name}")
        
        # Step 2: Process Google Sheets materials and identify new ones
        new_materials = []
        existing_materials_skipped = []
        successfully_processed_count = 0  # Count of materials that passed validation
        
        for category, category_data in sheet_data.items():
            # Handle nested materialNames structure
            material_names = category_data.get('materialNames', [])
            materials_to_process = []
            
            if isinstance(material_names, list):
                # Simple array structure
                materials_to_process = material_names
            elif isinstance(material_names, dict):
                # Nested structure - flatten all materials
                for sub_cat_or_spec, materials in material_names.items():
                    if isinstance(materials, list):
                        materials_to_process.extend(materials)
                    elif isinstance(materials, dict):
                        # Further nested (subCategory -> specifications -> materials)
                        for spec_or_mat, mat_list in materials.items():
                            if isinstance(mat_list, list):
                                materials_to_process.extend(mat_list)
            
            for material_info in materials_to_process:
                # Extract material details
                category_name = category
                sub_category = material_info.get('subCategory', '')
                material_name = material_info.get('name', '')
                specifications = material_info.get('specifications', '')
                uom = material_info.get('uom', '')
                initial_qty = float(material_info.get('initialQuantity', '0') or '0')
                
                # Count this as successfully processed (passed validation)
                successfully_processed_count += 1
                
                # Generate unique key for this material
                material_key = generate_material_key(
                    category_name,
                    sub_category,
                    material_name,
                    specifications
                )
                
                # Debug logging for first few materials
                if successfully_processed_count <= 3:
                    logging.info(f"Processing material {successfully_processed_count}: key='{material_key}', exists_in_firebase={material_key in existing_materials_map}")
                
                # Check if material already exists in Firebase
                if material_key in existing_materials_map:
                    # Material exists - skip it (preserve existing data)
                    existing_materials_skipped.append({
                        'category': category_name,
                        'subCategory': sub_category,
                        'materialName': material_name,
                        'specifications': specifications,
                        'reason': 'Already exists in Firebase (preserved)'
                    })
                    logging.info(f"Skipping existing material: {material_key}")
                else:
                    # New material - add it (following standard field sequence)
                    material_entry = {
                        'category': category_name,
                        'subCategory': sub_category,
                        'materialName': material_name,
                        'specifications': specifications,
                        'uom': uom,
                        'initialQuantity': initial_qty,
                        'currentQuantity': initial_qty,
                        'syncedAt': sync_timestamp,
                        'syncedBy': synced_by,
                        'syncDescription': sync_description,
                        'createdAt': sync_timestamp
                    }
                    new_materials.append(material_entry)
                    logging.info(f"Adding new material: {material_key}")
        
        # Verify our counts make sense
        total_successfully_processed = len(new_materials) + len(existing_materials_skipped)
        logging.info(f"Count verification: {len(new_materials)} new + {len(existing_materials_skipped)} existing = {total_successfully_processed} total processed")
        logging.info(f"Expected successfully processed: {successfully_processed_count}")
        
        # Step 3: Combine existing and new materials
        final_materials_list = existing_materials_list + new_materials
        
        # Step 4: Update Firebase with new nested structure
        # Use the new nested structure storage function
        from Utils.firebase_utils import store_materials_in_nested_structure
        
        sync_metadata = {
            'timestamp': sync_timestamp,
            'synced_by': synced_by,
            'description': sync_description
        }
        
        # Store materials in the correct nested structure for cascading dropdowns
        storage_result = store_materials_in_nested_structure(
            factory_name=document_name,
            materials_list=final_materials_list,
            sync_metadata=sync_metadata
        )
        
        if not storage_result['success']:
            return {
                'success': False,
                'message': f"Failed to store materials in nested structure: {storage_result['message']}"
            }
        
        # Step 5: Sync history is now handled by store_materials_in_nested_structure function
        
        logging.info(f"Smart sync completed for {plant_name}: {len(new_materials)} new materials added, {len(existing_materials_skipped)} existing materials preserved")
        
        # Step 6: Prepare detailed response
        return {
            'success': True,
            'message': f'Smart sync completed: {len(new_materials)} new materials added, {len(existing_materials_skipped)} existing materials preserved',
            'total_processed': total_processed,  # Total rows from Google Sheets (including skipped)
            'total_synced': len(new_materials),  # New materials added to Firebase
            'existing_materials_preserved': len(existing_materials_skipped),  # Existing materials preserved
            'skipped_count': skipped_count,  # Materials skipped due to validation errors
            'skipped_rows': skipped_rows,
            'skipped_reasons': skipped_reasons,
            'data': {
                'materialsCount': len(final_materials_list),
                'newMaterialsAdded': len(new_materials),
                'existingMaterialsPreserved': len(existing_materials_skipped),
                'categories': len(sheet_data),
                'plantName': plant_name,
                'documentName': document_name
            }
        }
        
    except Exception as e:
        logging.error(f"Error syncing material data to Firebase: {e}")
        return {
            'success': False,
            'message': f'Error syncing material data: {str(e)}'
        }

def process_general_reports(template_files, attachment_files, file_sequence, sheet_id, sheet_name, send_whatsapp, send_email, mail_subject, use_template_as_caption, user_id, output_dir, logger, send_email_func, send_whatsapp_message, validate_sheet_id_func, prepare_file_paths_func, fetch_google_sheet_data_func, process_template_func, send_log_report_to_user_func):
    """
    Process general reports using reactor report template logic with table and column processing
    This function moves the heavy functionality from app.py to process_utils.py
    """
    try:
        # Initialize result tracking
        result = {
            "success": False,
            "message": "",
            "generated_files": 0,
            "notifications_sent": {
                "email": False,
                "whatsapp": False
            },
            "delivery_stats": {
                "total_recipients": 0,
                "successful_deliveries": 0,
                "failed_deliveries": 0,
                "failed_contacts": []
            },
            "errors": [],
            "warnings": []
        }

        if not template_files:
            result["errors"].append("No template files provided")
            return result

        if not sheet_id:
            result["errors"].append("Google Sheet ID is required")
            return result

        if not sheet_name:
            result["errors"].append("Sheet name is required")
            return result

        # Validate sheet ID format
        if not validate_sheet_id_func(sheet_id):
            result["errors"].append("Invalid Google Sheet ID format")
            return result

        # Get user-specific temporary directory for attachments
        from .temp_manager import get_user_temp_dir
        temp_dir = get_user_temp_dir(user_id, output_dir)
        
        # Save attachment files temporarily and store their paths
        attachment_paths = prepare_file_paths_func(attachment_files, user_email=user_id, base_output_dir=output_dir, is_upload=True)

        try:
            # Fetch data from Google Sheet
            sheet_data = fetch_google_sheet_data_func(sheet_id, sheet_name)
            if not sheet_data or len(sheet_data) < 2:  # Check if we have headers and at least one row
                result["errors"].append("No data found in the Google Sheet")
                return result
        except Exception as e:
            logger.error("Error fetching Google Sheet data: {}".format(e))
            result["errors"].append("Failed to fetch data from Google Sheet")
            return result

        # Process headers and data
        headers = sheet_data[0]
        data_rows = sheet_data[1:]

        # Create output directory if it doesn't exist
        reports_output_dir = os.path.join(output_dir, "reports")
        os.makedirs(reports_output_dir, exist_ok=True)

        # Initialize delivery statistics tracking
        delivery_stats = {
            "total_recipients": len(data_rows),
            "successful_deliveries": 0,
            "failed_deliveries": 0,
            "failed_contacts": []
        }

        # STEP 1: Pre-process all template files and extract their content
        template_data = []  # List to store template info: {filename, temp_path, content}
        
        for template_file in template_files:
            if not template_file.filename.endswith('.docx'):
                logger.warning(f"Skipping non-docx template: {template_file.filename}")
                continue

            try:
                # Save template temporarily
                temp_template_path = os.path.join(reports_output_dir, "temp_{}".format(template_file.filename))
                template_file.save(temp_template_path)

                # Read template content for messages
                from docx import Document
                doc = Document(temp_template_path)
                template_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                template_data.append({
                    'filename': template_file.filename,
                    'temp_path': temp_template_path,
                    'raw_content': template_content
                })
                
                logger.info(f"Pre-processed template: {template_file.filename}")
                
            except Exception as e:
                logger.error("Error pre-processing template {}: {}".format(template_file.filename, e))
                result["errors"].append(f"Failed to pre-process template: {template_file.filename}")
                continue
        
        if not template_data:
            result["errors"].append("No valid template files to process")
            return result
        
        logger.info(f"Pre-processed {len(template_data)} template files")
        
        # STEP 2: Process each recipient with ALL templates in file_sequence order
        generated_files = []
        
        for row in data_rows:
            try:
                # Create data dictionary from headers and row
                data_dict = dict(zip(headers, row))
                recipient_name = data_dict.get('Name', 'unknown')
                
                logger.info(f"Processing recipient: {recipient_name}")
                
                # STEP 2a: Process each template for this recipient
                recipient_template_contents = {}  # Map filename -> processed content
                
                for template_info in template_data:
                    try:
                        # Replace placeholders in template content for this recipient
                        processed_content = template_info['raw_content']
                        
                        for key, value in data_dict.items():
                            placeholder = "{{{}}}".format(key)
                            processed_content = processed_content.replace(placeholder, str(value))
                        
                        recipient_template_contents[template_info['filename']] = processed_content
                        
                    except Exception as e:
                        logger.error(f"Error processing template {template_info['filename']} for {recipient_name}: {e}")
                        continue
                
                # STEP 2b: Prepare message content for Email (combine all templates)
                # For WhatsApp: Templates will be sent separately in sequence
                # For Email: Combine all templates into one email body
                email_content = ""
                
                if file_sequence and len(file_sequence) > 0:
                    # Find all message-type items in file_sequence (sorted by sequence_no)
                    message_items = sorted(
                        [item for item in file_sequence if item.get('file_type') == 'message'],
                        key=lambda x: x.get('sequence_no', 0)
                    )
                    
                    if message_items:
                        # Combine all templates in sequence order for EMAIL
                        combined_messages = []
                        
                        for message_item in message_items:
                            message_filename = message_item.get('file_name')
                            
                            if message_filename in recipient_template_contents:
                                combined_messages.append(recipient_template_contents[message_filename])
                                logger.info(f"Added template '{message_filename}' for {recipient_name}")
                            else:
                                logger.warning(f"Template '{message_filename}' not found for {recipient_name}")
                        
                        # Join all templates with double newline separator for email
                        email_content = "\n\n".join(combined_messages) if combined_messages else ""
                        
                        logger.info(f"Prepared {len(combined_messages)} template(s) for {recipient_name}")
                    else:
                        # No message items in sequence, use first template
                        email_content = list(recipient_template_contents.values())[0] if recipient_template_contents else ""
                else:
                    # No file_sequence provided, use first template
                    email_content = list(recipient_template_contents.values())[0] if recipient_template_contents else ""
                
                # Replace placeholders in mail subject
                processed_mail_subject = mail_subject
                for key, value in data_dict.items():
                    placeholder = "{{{}}}".format(key)
                    processed_mail_subject = processed_mail_subject.replace(placeholder, str(value))

                # STEP 2c: Get contact details from Google Sheet data
                recipient_email = data_dict.get('Email ID - To')
                cc_email = data_dict.get('Email ID - CC', '')
                bcc_email = data_dict.get('Email ID - BCC', '')
                
                # Helper to split emails by comma or newline and join as comma-separated string
                def clean_emails(email_str):
                    if not email_str:
                        return None
                    import re
                    emails = [e.strip() for e in re.split(r'[\n,]+', email_str) if e.strip()]
                    return ','.join(emails) if emails else None
                
                recipient_email = clean_emails(recipient_email)
                cc_email = clean_emails(cc_email)
                bcc_email = clean_emails(bcc_email)
                
                country_code = data_dict.get('Country Code', '').strip()
                phone_no = data_dict.get('Contact No.', '').strip()
                # Format phone number properly: remove spaces and combine country code + number
                recipient_phone = f"{country_code}{phone_no}".replace(' ', '')

                # STEP 2d: Handle WhatsApp notifications with validation and file sequencing
                if send_whatsapp:
                    # For WhatsApp: Handle file_sequence based on caption mode
                    whatsapp_file_sequence = []
                    
                    if use_template_as_caption:
                        # Caption mode: Remove message items (first template will be used as caption)
                        whatsapp_file_sequence = [
                            item for item in file_sequence 
                            if item.get('file_type') == 'file'
                        ] if file_sequence else []
                        
                        # Use first template as caption message
                        caption_message = list(recipient_template_contents.values())[0] if recipient_template_contents else ""
                        
                        logger.info(f"Caption mode: Using template as caption, {len(whatsapp_file_sequence)} files in sequence")
                    else:
                        # Normal mode: Keep ALL message items and file items
                        # Each template will be sent as a SEPARATE message in sequence
                        whatsapp_file_sequence = file_sequence if file_sequence else []
                        caption_message = ""  # Not used in normal mode
                        
                        logger.info(f"Normal mode: {len(whatsapp_file_sequence)} items in sequence")
                    
                    # Re-number the sequence to be continuous
                    for idx, item in enumerate(whatsapp_file_sequence, start=1):
                        item['sequence_no'] = idx
                    
                    logger.info(f"WhatsApp file sequence for {recipient_name}: {whatsapp_file_sequence}")
                    
                    # Pass recipient_template_contents so each message item can get its content
                    success, failure_reason = handle_whatsapp_validation_and_sending_v2(
                        recipient_name, country_code, phone_no, recipient_phone,
                        recipient_template_contents, attachment_paths, whatsapp_file_sequence,
                        use_template_as_caption, caption_message, send_whatsapp_message, logger
                    )
                    
                    if not success:
                        delivery_stats["failed_deliveries"] += 1
                        delivery_stats["failed_contacts"].append({
                            "name": recipient_name,
                            "contact": recipient_phone if recipient_phone else (phone_no if phone_no else country_code),
                            "reason": failure_reason
                        })
                    else:
                        delivery_stats["successful_deliveries"] += 1

                # STEP 2e: Handle email notifications with all attachments
                if send_email:
                    success, failure_reason = handle_email_sending(
                        recipient_name, recipient_email, cc_email, bcc_email,
                        processed_mail_subject, email_content, attachment_paths,
                        user_id, send_email_func, logger
                    )
                    
                    if not success:
                        logger.warning(f"Email sending failed for {recipient_name}: {failure_reason}")

            except Exception as e:
                logger.error("Error processing row for recipient {}: {}".format(recipient_name if 'recipient_name' in locals() else 'unknown', e))
                continue
        
        # STEP 3: Clean up temporary template files
        for template_info in template_data:
            try:
                os.remove(template_info['temp_path'])
                logger.info(f"Cleaned up temporary template: {template_info['filename']}")
            except Exception as e:
                logger.error("Error removing temporary template {}: {}".format(template_info['filename'], e))

        # Clean up user-specific temporary directory
        from .temp_manager import cleanup_user_temp_dir
        cleanup_user_temp_dir(user_id, output_dir)

        # Update result with delivery stats
        result["delivery_stats"] = delivery_stats
        result["generated_files"] = len(generated_files)
        result["success"] = True
        result["message"] = "Reports generated successfully"
        result["notifications_sent"]["email"] = send_email
        result["notifications_sent"]["whatsapp"] = send_whatsapp

        # Send log report to the user who generated the report
        try:
            send_log_report_to_user_func(user_id, delivery_stats, send_email, send_whatsapp, logger)
        except Exception as e:
            logger.error("Error sending log report to user: {}".format(e))

        return result

    except Exception as e:
        logger.error("Error in process_general_reports: {}".format(e))
        return {
            "success": False,
            "message": str(e),
            "errors": [str(e)]
        }

def process_template_with_reactor_logic(template_path, output_path, data_dict, logger):
    """
    Process template using reactor report table and column logic
    Uses reactorreportformat.docx template with same table and column processing as reactor reports
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        import os
        
        # Use reactorreportformat.docx as the base template
        reactor_template_path = os.path.join(os.path.dirname(__file__), "..", "reactorreportformat.docx")
        
        # Check if reactor template exists, otherwise use the provided template
        if os.path.exists(reactor_template_path):
            logger.info(f"Using reactor template: {reactor_template_path}")
            doc = Document(reactor_template_path)
        else:
            logger.info(f"Reactor template not found, using provided template: {template_path}")
            doc = Document(template_path)
        
        # Replace placeholders in the document using reactor report logic
        for paragraph in doc.paragraphs:
            for key, value in data_dict.items():
                placeholder = "{{{}}}".format(key)
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Process tables with reactor report formatting logic
        for table in doc.tables:
            # Apply professional table formatting using shared function
            format_table_professional(table, is_header=True, logger=logger)
            
            # Replace placeholders in table cells
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data_dict.items():
                        placeholder = "{{{}}}".format(key)
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        # Add data as tables using reactor report logic
        add_data_as_reactor_tables(doc, data_dict, logger)
        
        # Save the processed document
        doc.save(output_path)
        return True
        
    except Exception as e:
        logger.error("Error processing template with reactor logic: {}".format(e))
        return False

def add_data_as_reactor_tables(doc, data_dict, logger):
    """
    Add data as tables using reactor report table logic
    """
    try:
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.shared import OxmlElement, qn
        
        # Add a page break before adding data tables
        doc.add_page_break()
        
        # Add header for the data section
        header_para = doc.add_paragraph("Report Data")
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.bold = True
            run.font.size = Pt(14)
        
        # Create a table for the data with reactor report styling
        # Convert data_dict to a list of key-value pairs
        data_items = [(key, value) for key, value in data_dict.items() if value]
        
        if data_items:
            # Create table with headers
            table = doc.add_table(rows=len(data_items) + 1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table width
            table.allow_autofit = False
            table.autofit = False
            
            # Add header row
            header_row = table.rows[0]
            header_cells = header_row.cells
            header_cells[0].text = "Field"
            header_cells[1].text = "Value"
            
            # Format header cells
            for cell in header_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            
            # Add data rows
            for i, (key, value) in enumerate(data_items, 1):
                row = table.rows[i]
                cells = row.cells
                cells[0].text = str(key)
                cells[1].text = str(value)
                
                # Format data cells
                for cell in cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            # Apply professional table formatting using shared function
            format_table_professional(table, is_header=True, logger=logger)
            
            logger.info(f"Added data table with {len(data_items)} fields to document")
        
    except Exception as e:
        logger.error(f"Error adding data as reactor tables: {e}")


def handle_whatsapp_validation_and_sending(recipient_name, country_code, phone_no, recipient_phone, message_content, attachment_paths, file_sequence, use_template_as_caption, send_whatsapp_message, logger):
    """
    Handle WhatsApp validation and sending with proper error tracking
    """
    try:
        # Validate phone number components
        if not country_code or not country_code.strip():
            return False, "Missing Country Code"
        
        if not phone_no or not phone_no.strip():
            return False, "Missing Contact No."
        
        if not recipient_phone or len(recipient_phone.replace(' ', '')) < 4:
            return False, f"Invalid phone number format (Country Code: {country_code}, Contact No.: {phone_no})"
        
        logger.info("Valid phone number for {}: Country Code '{}', Contact No. '{}' -> Formatted: '{}'".format(
            recipient_name, country_code, phone_no, recipient_phone))

        # Send WhatsApp message with attachments
        options = {
            'use_template_as_caption': use_template_as_caption
        }
        success = send_whatsapp_message(
            contact_name=recipient_name,
            message=message_content,
            file_paths=attachment_paths,
            file_sequence=file_sequence,
            whatsapp_number=recipient_phone,
            process_name="report",
            options=options
        )
        
        # Handle different success/failure scenarios
        if success == "USER_NOT_LOGGED_IN":
            return False, "User session expired"
        elif success == "WHATSAPP_SERVICE_NOT_READY":
            return False, "WhatsApp service is not ready"
        elif success == "INVALID_FILE_PATH":
            return False, "Invalid file path for WhatsApp message"
        elif success == "INVALID_FILE_PATH_TYPE":
            return False, "Invalid file path type for WhatsApp message"
        elif success == "NO_VALID_FILES":
            return False, "No valid files found for WhatsApp message"
        elif success == "NO_FILES_FOR_UPLOAD":
            return False, "No files available for WhatsApp upload"
        elif success == "WHATSAPP_API_ERROR":
            return False, "WhatsApp API error"
        elif success == "WHATSAPP_CONNECTION_ERROR":
            return False, "WhatsApp connection error"
        elif success == "WHATSAPP_TIMEOUT_ERROR":
            return False, "WhatsApp timeout error"
        elif success == "WHATSAPP_SEND_ERROR":
            return False, "Failed to send WhatsApp message"
        elif not success:
            return False, "Failed to send WhatsApp message"
        else:
            logger.info("WhatsApp message sent successfully to {}".format(recipient_name))
            return True, None

    except Exception as e:
        logger.error("Error sending WhatsApp message to {}: {}".format(recipient_phone, e))
        return False, f"Exception: {str(e)}"

def handle_whatsapp_validation_and_sending_v2(recipient_name, country_code, phone_no, recipient_phone, recipient_template_contents, attachment_paths, file_sequence, use_template_as_caption, caption_message, send_whatsapp_message, logger):
    """
    Handle WhatsApp validation and sending with multiple templates as separate messages
    Each template is sent as a separate WhatsApp message in sequence order
    """
    try:
        # Validate phone number components
        if not country_code or not country_code.strip():
            return False, "Missing Country Code"
        
        if not phone_no or not phone_no.strip():
            return False, "Missing Contact No."
        
        if not recipient_phone or len(recipient_phone.replace(' ', '')) < 4:
            return False, f"Invalid phone number format (Country Code: {country_code}, Contact No.: {phone_no})"
        
        logger.info("Valid phone number for {}: Country Code '{}', Contact No. '{}' -> Formatted: '{}'".format(
            recipient_name, country_code, phone_no, recipient_phone))

        # Sort file_sequence by sequence_no
        sorted_sequence = sorted(file_sequence, key=lambda x: x.get('sequence_no', 0)) if file_sequence else []
        
        logger.info(f"Processing {len(sorted_sequence)} items in sequence for {recipient_name}")
        
        # Send each item in sequence
        import time
        for seq_item in sorted_sequence:
            try:
                item_type = seq_item.get('file_type')
                sequence_no = seq_item.get('sequence_no')
                
                if item_type == 'message':
                    # Send template message
                    template_filename = seq_item.get('file_name')
                    template_content = recipient_template_contents.get(template_filename, "")
                    
                    if not template_content:
                        logger.warning(f"No content found for template '{template_filename}' in sequence {sequence_no}")
                        continue
                    
                    logger.info(f"Sending template '{template_filename}' as message (sequence {sequence_no})")
                    
                    # Send just the message (no files, no sequence)
                    options = {'use_template_as_caption': False}  # Don't use caption for separate messages
                    success = send_whatsapp_message(
                        contact_name=recipient_name,
                        message=template_content,
                        file_paths=[],  # No files for template messages
                        file_sequence=[],  # No sequence - sending one item at a time
                        whatsapp_number=recipient_phone,
                        process_name="report",
                        options=options
                    )
                    
                    if success is not True:
                        logger.error(f"Failed to send template '{template_filename}' to {recipient_name}: {success}")
                        return False, f"Failed to send template message: {success}"
                    
                    logger.info(f"Template '{template_filename}' sent successfully (sequence {sequence_no})")
                    
                elif item_type == 'file':
                    # Send file attachment
                    file_name = seq_item.get('file_name')
                    
                    # Find the file path in attachment_paths
                    import os
                    file_path = next((path for path in attachment_paths if os.path.basename(path) == file_name), None)
                    
                    if not file_path:
                        logger.warning(f"File '{file_name}' not found in attachment_paths for sequence {sequence_no}")
                        continue
                    
                    logger.info(f"Sending file '{file_name}' (sequence {sequence_no})")
                    
                    # Check if we should use caption (only in caption mode)
                    if use_template_as_caption and caption_message:
                        options = {'use_template_as_caption': True}
                        message_to_send = caption_message
                    else:
                        options = {'use_template_as_caption': False}
                        message_to_send = ""  # No message, just file
                    
                    success = send_whatsapp_message(
                        contact_name=recipient_name,
                        message=message_to_send,
                        file_paths=[file_path],  # Single file
                        file_sequence=[],  # No sequence - sending one item at a time
                        whatsapp_number=recipient_phone,
                        process_name="report",
                        options=options
                    )
                    
                    if success is not True:
                        logger.error(f"Failed to send file '{file_name}' to {recipient_name}: {success}")
                        return False, f"Failed to send file: {success}"
                    
                    logger.info(f"File '{file_name}' sent successfully (sequence {sequence_no})")
                
                # Add delay between items (1 second)
                time.sleep(1)
                
            except Exception as e:
                logger.error(f"Error sending sequence item {sequence_no} to {recipient_name}: {e}")
                continue
        
        logger.info(f"All {len(sorted_sequence)} items sent successfully to {recipient_name}")
        return True, None

    except Exception as e:
        logger.error("Error in WhatsApp sending to {}: {}".format(recipient_phone, e))
        return False, f"Exception: {str(e)}"

def handle_email_sending(recipient_name, recipient_email, cc_email, bcc_email, mail_subject, email_content, attachment_paths, user_id, send_email_func, logger):
    """
    Handle email sending with proper error tracking
    """
    try:
        if not recipient_email:
            return False, "No email found for recipient"
        
        email_subject_formatted = mail_subject
        email_body = f"""
        <html>
        <body>
        {email_content.replace('\n', '<br>')} 
        </body>
        </html>
        """
        
        success = send_email_func(
            user_email=user_id,
            recipient_email=recipient_email,
            subject=email_subject_formatted,
            body=email_body,
            attachment_paths=attachment_paths,
            cc=cc_email,
            bcc=bcc_email
        )
        
        if success == "TOKEN_EXPIRED":
            return False, "Email token expired"
        elif success == "USER_NOT_LOGGED_IN":
            return False, "User session expired"
        elif success == "NO_GMAIL_ACCESS":
            return False, "Gmail access not configured for this user"
        elif success == "NO_GOOGLE_ACCESS_TOKEN":
            return False, "No Google access token found"
        elif success == "INVALID_RECIPIENT":
            return False, "Invalid recipient email address"
        elif success == "GMAIL_AUTH_FAILED":
            return False, "Gmail authentication failed"
        elif success == "GMAIL_PERMISSION_DENIED":
            return False, "Gmail permission denied"
        elif success == "GMAIL_API_ERROR":
            return False, "Gmail API error"
        elif success == "GMAIL_SEND_ERROR":
            return False, "Failed to send email"
        elif not success:
            return False, "Failed to send email"
        else:
            logger.info("Email sent successfully to {}".format(recipient_email))
            return True, None

    except Exception as e:
        logger.error("Error sending email to {}: {}".format(recipient_email, e))
        return False, f"Exception: {str(e)}"

def generate_log_report_pdf(delivery_stats, output_dir, logger):
    """
    Generate a PDF log report from delivery statistics
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        import os
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('📊 Delivery Log Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        timestamp_para = doc.add_paragraph(f"Generated on: {timestamp}")
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Extract statistics
        total_recipients = delivery_stats.get("total_recipients", 0)
        successful_deliveries = delivery_stats.get("successful_deliveries", 0)
        failed_deliveries = delivery_stats.get("failed_deliveries", 0)
        failed_contacts = delivery_stats.get("failed_contacts", [])
        
        # Add summary section
        summary_heading = doc.add_heading('📋 Summary', level=1)
        
        # Create summary table
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for row in summary_table.rows:
            for i, cell in enumerate(row.cells):
                if i == 0:
                    cell.width = Inches(3)
                else:
                    cell.width = Inches(1.5)
        
        # Add summary data
        summary_data = [
            ("📋 Total messages to be delivered:", str(total_recipients)),
            ("✅ Messages delivered successfully:", str(successful_deliveries)),
            ("❌ Failed messages:", str(failed_deliveries)),
            ("📊 Success Rate:", f"{(successful_deliveries/total_recipients*100):.1f}%" if total_recipients > 0 else "0%")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Format the cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()
        
        # Add failed contacts section if there are any
        if failed_contacts:
            failed_heading = doc.add_heading('📋 Failed Contacts', level=1)
            
            # Create failed contacts table
            failed_table = doc.add_table(rows=1, cols=3)
            failed_table.style = 'Table Grid'
            failed_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            for row in failed_table.rows:
                for i, cell in enumerate(row.cells):
                    if i == 0:
                        cell.width = Inches(2)
                    elif i == 1:
                        cell.width = Inches(2)
                    else:
                        cell.width = Inches(3)
            
            # Add header row
            header_row = failed_table.rows[0]
            header_row.cells[0].text = "Name"
            header_row.cells[1].text = "Contact Number"
            header_row.cells[2].text = "Reason"
            
            # Add data rows
            for contact in failed_contacts:
                row = failed_table.add_row()
                row.cells[0].text = contact.get("name", "Unknown")
                row.cells[1].text = contact.get("contact", "N/A")
                row.cells[2].text = contact.get("reason", "Unknown error")
            
            # Apply simple table formatting for log report
            format_table_simple(failed_table, is_header=True, logger=logger)
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph("Generated by Bajaj Earths Automation System")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.italic = True
        
        # Save the document
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"log_report_{timestamp_str}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        doc.save(docx_path)
        
        # Convert to PDF
        pdf_filename = f"log_report_{timestamp_str}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        if convert_docx_to_pdf(docx_path, pdf_path):
            logger.info(f"Successfully generated PDF log report: {pdf_path}")
            # Clean up DOCX file
            try:
                os.remove(docx_path)
                logger.info(f"Cleaned up temporary DOCX file: {docx_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up DOCX file {docx_path}: {e}")
            
            return pdf_path
        else:
            logger.error(f"Failed to convert DOCX to PDF for log report")
            return None
            
    except Exception as e:
        logger.error(f"Error generating log report PDF: {e}")
        return None

def process_order_notification(order_id, order_data, recipients, method, factory, template_path, output_dir, user_email, logger, send_whatsapp_message):
    """
    Process order notification by creating a formatted document and sending via email/WhatsApp
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        import os
        
        # Initialize result tracking
        result = {
            "success": False,
            "message": "",
            "delivery_stats": {
                "total_recipients": len(recipients),
                "successful_deliveries": 0,
                "failed_deliveries": 0,
                "email_successful": 0,
                "whatsapp_successful": 0,
                "failed_contacts": []
            },
            "errors": [],
            "warnings": []
        }
        
        logger.info(f"Processing order notification for Order ID: {order_id}")
        logger.info(f"Recipients count: {len(recipients)}, Method: {method}, Factory: {factory}")
        logger.info(f"Order data keys: {list(order_data.keys()) if order_data else 'None'}")
        logger.info(f"Order items count: {len(order_data.get('orderItems', [])) if order_data else 0}")
        
        # Create document
        doc = None
        try:
            if template_path and os.path.exists(template_path):
                logger.info(f"Attempting to use template: {template_path}")
                doc = Document(template_path)
                logger.info(f"Template loaded successfully: {template_path}")
            else:
                logger.info("Template not found, creating new document")
                doc = Document()
        except Exception as e:
            logger.warning(f"Error loading template, falling back to new document: {e}")
            try:
                doc = Document()
                logger.info("Fallback: Created new document successfully")
            except Exception as e2:
                logger.error(f"Error creating new document: {e2}")
                raise
        
        if doc is None:
            raise Exception("Failed to create document")
        
        # Add content to document (don't clear template content)
        logger.info("Adding content to document...")
        
        # Debug document structure
        try:
            logger.info(f"Document has {len(doc.paragraphs)} paragraphs")
            logger.info(f"Document has {len(doc.tables)} tables")
            if hasattr(doc, 'element') and hasattr(doc.element, 'body'):
                logger.info(f"Document body has {len(doc.element.body)} elements")
        except Exception as e:
            logger.warning(f"Could not inspect document structure: {e}")
        
        # Parse dateTime to separate Date and Time
        date_time_str = order_data.get('dateTime', '')
        date_value = ''
        time_value = ''
        
        if date_time_str:
            try:
                # Format: "DD/MM/YYYY, HH:MM:SS AM/PM" or "DD/MM/YYYY, HH:MM AM/PM"
                parts = date_time_str.split(',')
                if len(parts) >= 2:
                    date_value = parts[0].strip()  # "DD/MM/YYYY"
                    time_value = parts[1].strip()  # "HH:MM:SS AM/PM" or "HH:MM AM/PM"
            except Exception as e:
                logger.warning(f"Error parsing dateTime '{date_time_str}': {e}. Using full string.")
                date_value = date_time_str
                time_value = ''
        
        if not date_value:
            date_value = 'N/A'
        if not time_value:
            time_value = 'N/A'
        
        # Get order items count
        order_items = order_data.get('orderItems', [])
        total_items = len(order_items)
        
        # Create Indent Details Table (horizontal format)
        try:
            logger.info("Creating indent details table...")
            
            # Add heading paragraph
            logger.info("Adding indent details heading paragraph...")
            details_heading = doc.add_paragraph("Indent Details")
            details_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in details_heading.runs:
                run.bold = True
                run.font.size = Pt(12)
            logger.info("Indent details heading paragraph created successfully")
            
            # Create horizontal table with 9 columns: Date, Time, Factory, Order ID, Given By, Type, Importance, Total Items, Description
            logger.info("Creating horizontal table with 9 columns...")
            try:
                # Create table with 1 header row and 1 data row
                details_table = doc.add_table(rows=2, cols=9)
                logger.info(f"Table created with {len(details_table.rows)} rows and {len(details_table.rows[0].cells)} columns")
            except Exception as table_error:
                logger.warning(f"Error creating table with template, trying fallback approach: {table_error}")
                # Fallback: create a new document without template
                logger.info("Creating fallback document without template...")
                doc = Document()
                details_table = doc.add_table(rows=2, cols=9)
                logger.info("Fallback table created successfully")
            
            # Set table alignment
            logger.info("Setting table alignment...")
            details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            logger.info("Indent details table created successfully")
            
            # Add header row
            header_cells = details_table.rows[0].cells
            headers = ['Date', 'Time', 'Factory', 'Order ID', 'Given By', 'Type', 'Importance', 'Total Items', 'Description']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            # Add data row
            data_cells = details_table.rows[1].cells
            data_values = [
                date_value,
                time_value,
                factory,
                order_id,
                order_data.get('givenBy', 'N/A'),
                order_data.get('type', 'N/A'),
                order_data.get('importance', 'Normal'),
                str(total_items),
                order_data.get('description', 'N/A')
            ]
            
            for i, value in enumerate(data_values):
                data_cells[i].text = str(value)
                data_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in data_cells[i].paragraphs:
                    if i in [0, 1, 2, 3, 6, 7]:  # Date, Time, Factory, Order ID, Importance, Total Items - center aligned
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:  # Given By, Type, Description - left aligned
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            
            logger.info("Indent details table populated successfully")
            
        except Exception as e:
            logger.error(f"Error creating indent details table: {e}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise
        
        # Apply professional formatting
        try:
            format_table_professional(details_table, is_header=True, logger=logger)
        except Exception as e:
            logger.warning(f"Could not apply professional formatting to indent details table: {e}")
            # Continue without formatting
        
        doc.add_paragraph()  # Spacing
        
        # Create Item Details Table
        items_heading = doc.add_paragraph("Item Details")
        items_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in items_heading.runs:
            run.bold = True
            run.font.size = Pt(12)
        
        if order_items:
            # Create table with headers
            items_table = doc.add_table(rows=1, cols=9)
            items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add header row
            header_cells = items_table.rows[0].cells
            headers = ['S.No', 'Category', 'Sub Category', 'Material Name', 'Particulars', 'Quantity', 'UOM', 'Preferred Vendor', 'Place']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            # Add data rows
            for idx, item in enumerate(order_items, 1):
                row = items_table.add_row()
                cells = row.cells
                
                cells[0].text = str(idx)
                cells[1].text = item.get('category', '')
                cells[2].text = item.get('subCategory') or '-'
                cells[3].text = item.get('materialName', '')  # Material Name in column 3
                cells[4].text = item.get('specifications') or '-'  # Particulars in column 4
                cells[5].text = str(item.get('quantity', ''))
                cells[6].text = item.get('uom', '')
                cells[7].text = item.get('partyName') or '-'  # Preferred Vendor in column 7
                cells[8].text = item.get('place') or '-'  # Place in column 8
                
                # Format cells
                for i, cell in enumerate(cells):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        if i == 0:  # S.No
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif i in [5, 6]:  # Quantity and UOM
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # Apply professional formatting
            try:
                format_table_professional(items_table, is_header=True, logger=logger)
            except Exception as e:
                logger.warning(f"Could not apply professional formatting to items table: {e}")
                # Continue without formatting
        
        # Generate filename in format: {FACTORY}_INDT_{Order_ID}_{Day}_{HHMM} {AM/PM}
        # Parse dateTime to extract day, time, and AM/PM
        date_time_str = order_data.get('dateTime', '')
        day = ''
        time_hhmm_ampm = ''
        
        if date_time_str:
            try:
                # Format: "DD/MM/YYYY, HH:MM AM/PM" (e.g., "03/12/2025, 06:10 PM")
                parts = date_time_str.split(',')
                if len(parts) >= 2:
                    # Extract day from date part (DD/MM/YYYY)
                    date_part = parts[0].strip()
                    date_components = date_part.split('/')
                    if len(date_components) >= 1:
                        day = date_components[0].strip()  # Get day (DD)
                    
                    # Extract time and AM/PM from time part (HH:MM AM/PM)
                    time_part = parts[1].strip()  # "06:10 PM"
                    # Remove colon and keep AM/PM
                    time_hhmm_ampm = time_part.replace(':', '')  # "0610 PM"
            except Exception as e:
                logger.warning(f"Error parsing dateTime '{date_time_str}': {e}. Using fallback.")
        
        # Fallback to current date/time if parsing failed
        if not day or not time_hhmm_ampm:
            now = datetime.now()
            day = now.strftime("%d")
            time_hhmm_ampm = now.strftime("%I%M %p")  # "0610 PM" format
        
        # Generate filename: {FACTORY}_INDT_{Order_ID}_{Day}_{HHMM} {AM/PM}
        base_filename = f"INDT_{order_id}_{day}_{time_hhmm_ampm}"
        docx_filename = f"{base_filename}.docx"
        pdf_filename = f"{base_filename}.pdf"
        
        docx_path = os.path.join(output_dir, docx_filename)
        doc.save(docx_path)
        
        logger.info(f"Order document created: {docx_path}")
        
        # Convert to PDF
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        pdf_created = False
        if convert_docx_to_pdf(docx_path, pdf_path):
            logger.info(f"Successfully converted to PDF: {pdf_path}")
            pdf_created = True
            # Use PDF for notifications
            notification_file = pdf_path
        else:
            logger.warning("PDF conversion failed, using DOCX file")
            result["warnings"].append("PDF conversion failed, using DOCX file")
            # Use DOCX for notifications
            notification_file = docx_path
        
        # Upload PDF to Google Drive before notifications (only if PDF was created)
        drive_upload_success = None  # None = not attempted, True = success, False = failed
        if pdf_created and pdf_path:
            try:
                from Utils.drive_utils import upload_store_document_to_drive
                
                # Get date string from order_data
                date_string = order_data.get('dateTime', '')
                if not date_string:
                    # Fallback: use current date
                    date_string = datetime.now().strftime("%d/%m/%Y, %I:%M:%S %p")
                
                logger.info(f"Uploading order PDF to Google Drive (factory: {factory}, date: {date_string})")
                upload_success, file_id, folder_id, upload_error = upload_store_document_to_drive(
                    pdf_path=pdf_path,
                    factory=factory,
                    date_string=date_string,
                    process_type="place_order",
                    file_name=pdf_filename,
                    logger=logger
                )
                
                if upload_success:
                    logger.info(f"Successfully uploaded order PDF to Google Drive. File ID: {file_id}, Folder ID: {folder_id}")
                    drive_upload_success = True
                    result["drive_upload"] = {
                        "success": True,
                        "file_id": file_id,
                        "folder_id": folder_id
                    }
                else:
                    logger.warning(f"Failed to upload order PDF to Google Drive: {upload_error}")
                    result["warnings"].append(f"Google Drive upload failed: {upload_error}")
                    drive_upload_success = False
                    result["drive_upload"] = {
                        "success": False,
                        "error": upload_error
                    }
            except Exception as e:
                logger.error(f"Error during Google Drive upload: {e}")
                result["warnings"].append(f"Error during Google Drive upload: {e}")
                drive_upload_success = False
                result["drive_upload"] = {
                    "success": False,
                    "error": str(e)
                }
        
        # Extract order items for notifications
        order_items = order_data.get('orderItems', [])
        
        # Fetch designation from authority sheet based on givenBy name
        given_by_name = order_data.get('givenBy', '')
        designation = get_designation_from_authority_sheet(factory, given_by_name, logger) if given_by_name else 'N/A'
        
        # Generate HTML table rows for order items
        items_table_rows_html = ''
        if order_items:
            items_table_rows_html = ''.join([
                f'''
                        <tr>
                            <td style="padding: 8px; border: 1px solid #ddd; text-align: center;">{idx}</td>
                            <td style="padding: 8px; border: 1px solid #ddd;">{item.get('category', '')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; {'text-align: center;' if not item.get('subCategory') else ''}">{(item.get('subCategory') or '-')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd;">{item.get('materialName', '')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; {'text-align: center;' if not item.get('specifications') else ''}">{(item.get('specifications') or '-')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; text-align: center;">{item.get('quantity', '')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; text-align: center;">{item.get('uom', '')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; {'text-align: center;' if not item.get('partyName') else ''}">{(item.get('partyName') or '-')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; {'text-align: center;' if not item.get('place') else ''}">{(item.get('place') or '-')}</td>
                        </tr>'''
                for idx, item in enumerate(order_items, 1)
            ])
        else:
            items_table_rows_html = '<tr><td colspan="9" style="padding: 8px; border: 1px solid #ddd; text-align: center;">No items found</td></tr>'
        
        # Prepare notification message
        notification_message = f"""📦 *New Material Indent Raised*

*Date:* {date_value}
*Time:* {time_value}
*Factory:* {factory}
*Order ID:* {order_id}
*Given By:* {order_data.get('givenBy', 'N/A')}
*Type:* {order_data.get('type', 'N/A')}
*Importance:* {order_data.get('importance', 'Normal')}
*Total Items:* {len(order_items)}

*Description:*
{order_data.get('description', 'N/A')}


Please find the detailed order document attached."""
        
        # Send notifications to recipients
        send_email = method in ['email', 'both']
        send_whatsapp = method in ['whatsapp', 'both']
        
        # ============================================
        # EMAIL: Collect all recipients for single email send
        # ============================================
        # Track bulk email result and recipient email status
        bulk_email_success = None  # None = not attempted, True = success, False = failed
        bulk_email_failure_reason = None
        recipients_with_email = set()  # Track which recipients have email addresses
        
        if send_email:
            logger.info("=" * 60)
            logger.info("Collecting all email addresses for single email send")
            logger.info("=" * 60)
            
            # Helper function to clean and split emails
            def clean_and_split_emails(email_str):
                """Clean email string and return list of individual emails"""
                if not email_str or not email_str.strip():
                    return []
                import re
                emails = [e.strip() for e in re.split(r'[\n,]+', email_str) if e.strip()]
                return emails
            
            # Collect all email addresses
            all_to_emails = set()  # Use set to automatically deduplicate
            all_cc_emails = set()
            all_bcc_emails = set()
            recipients_with_no_email = []
            
            for recipient in recipients:
                recipient_name = recipient.get('Name', 'Unknown')
                recipient_to = recipient.get('Email ID - To', '').strip()
                recipient_cc = recipient.get('Email ID - CC', '').strip()
                recipient_bcc = recipient.get('Email ID - BCC', '').strip()
                
                logger.info(f"Processing recipient for email collection: {recipient_name}")
                logger.info(f"  Email ID - To: '{recipient_to}' (present: {bool(recipient_to)})")
                logger.info(f"  Email ID - CC: '{recipient_cc}' (present: {bool(recipient_cc)})")
                logger.info(f"  Email ID - BCC: '{recipient_bcc}' (present: {bool(recipient_bcc)})")
                
                # Collect To emails
                if recipient_to:
                    to_emails = clean_and_split_emails(recipient_to)
                    all_to_emails.update(to_emails)
                    logger.info(f"  Added {len(to_emails)} To email(s)")
                
                # Collect CC emails
                if recipient_cc:
                    cc_emails = clean_and_split_emails(recipient_cc)
                    all_cc_emails.update(cc_emails)
                    logger.info(f"  Added {len(cc_emails)} CC email(s)")
                
                # Collect BCC emails
                if recipient_bcc:
                    bcc_emails = clean_and_split_emails(recipient_bcc)
                    all_bcc_emails.update(bcc_emails)
                    logger.info(f"  Added {len(bcc_emails)} BCC email(s)")
                
                # Track recipients with no email addresses
                if not recipient_to and not recipient_cc and not recipient_bcc:
                    recipients_with_no_email.append(recipient_name)
                    logger.warning(f"  ⚠ No email addresses found for {recipient_name}")
                else:
                    # Track that this recipient has email
                    recipients_with_email.add(recipient_name)
            
            # Convert sets to sorted lists for consistent ordering
            all_to_emails_list = sorted(list(all_to_emails))
            all_cc_emails_list = sorted(list(all_cc_emails))
            all_bcc_emails_list = sorted(list(all_bcc_emails))
            
            logger.info("=" * 60)
            logger.info("Email Collection Summary:")
            logger.info(f"  Total To emails: {len(all_to_emails_list)}")
            logger.info(f"  Total CC emails: {len(all_cc_emails_list)}")
            logger.info(f"  Total BCC emails: {len(all_bcc_emails_list)}")
            logger.info(f"  Recipients with no email: {len(recipients_with_no_email)}")
            if recipients_with_no_email:
                logger.warning(f"  Recipients without email: {', '.join(recipients_with_no_email)}")
            logger.info("=" * 60)
            
            # Handle case where no To addresses exist but CC/BCC exists
            # Gmail API requires a To address, so use sender's email as To
            final_to_emails = all_to_emails_list
            if not all_to_emails_list and (all_cc_emails_list or all_bcc_emails_list):
                final_to_emails = [user_email]
                logger.warning(f"No To addresses found, using sender's email ({user_email}) as To address")
                logger.info(f"CC emails will remain in CC: {len(all_cc_emails_list)} email(s)")
                logger.info(f"BCC emails will remain in BCC: {len(all_bcc_emails_list)} email(s)")
            
            # Check if we have any email addresses to send to
            has_any_email = bool(final_to_emails) or bool(all_cc_emails_list) or bool(all_bcc_emails_list)
            
            if has_any_email:
                # Prepare email content (same for all recipients)
                email_subject = f"New Material Indent Raised - {order_data.get('givenBy', 'N/A')} - {order_id}"
                
                # Parse dateTime to extract separate date and time
                date_time_str = order_data.get('dateTime', '')
                date_value = 'N/A'
                time_value = 'N/A'
                
                if date_time_str:
                    try:
                        # Format: "DD/MM/YYYY, HH:MM AM/PM" (e.g., "03/12/2025, 06:10 PM")
                        parts = date_time_str.split(',')
                        if len(parts) >= 2:
                            date_value = parts[0].strip()  # Extract date part: "DD/MM/YYYY"
                            time_value = parts[1].strip()  # Extract time part: "HH:MM AM/PM"
                    except Exception as e:
                        logger.warning(f"Error parsing dateTime '{date_time_str}': {e}. Using fallback.")
                        # Fallback to current date/time if parsing fails
                        now = datetime.now()
                        date_value = now.strftime("%d/%m/%Y")
                        time_value = now.strftime("%I:%M %p")
                else:
                    # If dateTime is missing, use current date/time
                    now = datetime.now()
                    date_value = now.strftime("%d/%m/%Y")
                    time_value = now.strftime("%I:%M %p")
                
                email_body = f"""
                    <html>
                    <body>
                    <p>Dear Sir/Ma'am,</p>
                    <p>A new indent has been raised for your review and further processing.</p>
                    Please find the details below:
                    
                    <br><p></p>
                    <b style="font-size: 1.1rem;">Indent Details</b>
                    <table style="border-collapse: collapse; margin-top: 0.2rem;">
                        <tr style="text-align: center;">
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Date</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Time</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Factory</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Order ID</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Given By</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Type</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Importance</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Total Items</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Description</strong></td>
                        </tr>
                        <tr>
                            <td style="padding: 8px; border: 1px solid #ddd; text-align: center;">{date_value}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; text-align: center;">{time_value}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; text-align: center;">{factory}</td>
                            <td style="padding: 8px; border: 1px solid #ddd;">{order_id}</td>
                            <td style="padding: 8px; border: 1px solid #ddd;">{order_data.get('givenBy', 'N/A')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd;">{order_data.get('type', 'N/A')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd;">{order_data.get('importance', 'Normal')}</td>
                            <td style="padding: 8px; border: 1px solid #ddd; text-align: center;">{len(order_items)}</td>
                            <td style="padding: 8px; border: 1px solid #ddd;">{order_data.get('description', 'N/A')}</td>
                        </tr>
                    </table>
                    <br><p></p>
                    
                    <b style="font-size: 1.1rem;">Item Details</b>
                    <table style="border-collapse: collapse; margin-top: 0.2rem;">
                        <tr style="text-align: center;">
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>S.No</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Category</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Sub Category</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Material Name</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Particulars</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Quantity</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>UOM</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Preferred Vendor</strong></td>
                            <td style="padding: 8px; border: 1px solid #ddd;"><strong>Place</strong></td>
                        </tr>
                        {items_table_rows_html}
                    </table>
                    
                    <br><p></p>
                    Please find the detailed order document attached.
                    
                    <br><br><p></p>
                    Best Regards,<br>{order_data.get('givenBy', 'N/A')} <br> {designation}
                    </body>
                    </html>
                    """
                
                # Prepare comma-separated email strings
                to_emails_str = ','.join(final_to_emails) if final_to_emails else None
                cc_emails_str = ','.join(all_cc_emails_list) if all_cc_emails_list else None
                bcc_emails_str = ','.join(all_bcc_emails_list) if all_bcc_emails_list else None
                
                logger.info("=" * 60)
                logger.info("Sending single email to all recipients:")
                logger.info(f"  All To recipients ({len(final_to_emails)}): {to_emails_str}")
                logger.info(f"  All CC recipients ({len(all_cc_emails_list)}): {cc_emails_str if cc_emails_str else '(none)'}")
                logger.info(f"  All BCC recipients ({len(all_bcc_emails_list)}): {bcc_emails_str if bcc_emails_str else '(none)'}")
                logger.info("=" * 60)
                
                try:
                    # Send single email with all recipients
                    # Gmail API accepts comma-separated strings for To, CC, and BCC
                    success = send_email_gmail_api(
                        user_email=user_email,
                        recipient_email=to_emails_str,  # Pass all To emails as comma-separated string
                        subject=email_subject,
                        body=email_body,
                        attachment_paths=[notification_file],
                        cc=cc_emails_str,  # Pass all CC emails as comma-separated string
                        bcc=bcc_emails_str  # Pass all BCC emails as comma-separated string
                    )
                    
                    if success is True:
                        # Count all recipients for stats
                        total_email_recipients = len(final_to_emails) + len(all_cc_emails_list) + len(all_bcc_emails_list)
                        result["delivery_stats"]["email_successful"] = total_email_recipients
                        result["delivery_stats"]["successful_deliveries"] += total_email_recipients
                        bulk_email_success = True
                        
                        logger.info("=" * 60)
                        logger.info(f"✓ Single email sent successfully!")
                        logger.info(f"  Total recipients: {total_email_recipients}")
                        logger.info(f"    - To: {len(final_to_emails)}")
                        logger.info(f"    - CC: {len(all_cc_emails_list)}")
                        logger.info(f"    - BCC: {len(all_bcc_emails_list)}")
                        logger.info("=" * 60)
                    else:
                        failure_reason = f"Email error: {success}" if isinstance(success, str) else "Email send failed"
                        bulk_email_success = False
                        bulk_email_failure_reason = failure_reason
                        total_email_recipients = len(final_to_emails) + len(all_cc_emails_list) + len(all_bcc_emails_list)
                        result["delivery_stats"]["failed_deliveries"] += total_email_recipients
                        
                        logger.error("=" * 60)
                        logger.error(f"✗ Failed to send single email: {success}")
                        logger.error(f"  Total recipients affected: {total_email_recipients}")
                        logger.error(f"    - To: {len(final_to_emails)}")
                        logger.error(f"    - CC: {len(all_cc_emails_list)}")
                        logger.error(f"    - BCC: {len(all_bcc_emails_list)}")
                        logger.error("=" * 60)
                        
                except Exception as e:
                    exception_reason = f"Email exception: {str(e)}"
                    bulk_email_success = False
                    bulk_email_failure_reason = exception_reason
                    total_email_recipients = len(final_to_emails) + len(all_cc_emails_list) + len(all_bcc_emails_list)
                    result["delivery_stats"]["failed_deliveries"] += total_email_recipients
                    
                    logger.error("=" * 60)
                    logger.error(f"✗ Error sending single email: {e}")
                    logger.error(f"  Total recipients affected: {total_email_recipients}")
                    logger.error("=" * 60)
            else:
                logger.warning("=" * 60)
                logger.warning("⚠ No email addresses found in any recipients")
                logger.warning(f"  Recipients processed: {len(recipients)}")
                logger.warning(f"  Recipients without email: {len(recipients_with_no_email)}")
                logger.warning("=" * 60)
        
        # ============================================
        # WHATSAPP: Send individual messages (unchanged)
        # ============================================
        for recipient in recipients:
            recipient_name = recipient.get('Name', 'Unknown')
            recipient_phone_raw = recipient.get('Contact No.', '')
            country_code = recipient.get('Country Code', '91')
            
            # Format phone number
            recipient_phone = f"{country_code}{recipient_phone_raw}".replace(' ', '') if recipient_phone_raw else ''
            contact_identifiers = [value for value in [recipient_phone] if value]
            contact_display = " | ".join(contact_identifiers) if contact_identifiers else "N/A"
            recipient_success = False
            recipient_fail_reasons = []
            channel_status = {
                "email": {"status": "not_enabled", "reason": ""},
                "whatsapp": {"status": "not_enabled", "reason": ""}
            }
            
            # Set email status based on bulk email result
            if send_email:
                if recipient_name in recipients_with_email:
                    # Recipient has email - set status based on bulk email result
                    if bulk_email_success is True:
                        channel_status["email"]["status"] = "success"
                    elif bulk_email_success is False:
                        channel_status["email"]["status"] = "failed"
                        channel_status["email"]["reason"] = bulk_email_failure_reason or "Email send failed"
                        recipient_fail_reasons.append(bulk_email_failure_reason or "Email send failed")
                    else:
                        # Bulk email not attempted (shouldn't happen, but handle it)
                        channel_status["email"]["status"] = "not_enabled"
                else:
                    # Recipient has no email address
                    no_email_reason = "Email: No recipient address provided (To, CC, or BCC required)"
                    channel_status["email"]["status"] = "skipped"
                    channel_status["email"]["reason"] = no_email_reason
                    recipient_fail_reasons.append(no_email_reason)
            else:
                channel_status["email"]["status"] = "not_enabled"
            
            # Send WhatsApp
            if send_whatsapp:
                channel_status["whatsapp"]["status"] = "pending"
                if recipient_phone_raw and country_code:
                    try:
                        success = send_whatsapp_message(
                            contact_name=recipient_name,
                            message=notification_message,
                            file_paths=[notification_file],
                            whatsapp_number=recipient_phone,
                            process_name="order_notification"
                        )
                        
                        if success is True:
                            recipient_success = True
                            channel_status["whatsapp"]["status"] = "success"
                            result["delivery_stats"]["whatsapp_successful"] += 1
                            logger.info(f"WhatsApp sent successfully to {recipient_name} ({recipient_phone})")
                        else:
                            failure_reason = f"WhatsApp error: {success}" if isinstance(success, str) else "WhatsApp send failed"
                            recipient_fail_reasons.append(failure_reason)
                            channel_status["whatsapp"]["status"] = "failed"
                            channel_status["whatsapp"]["reason"] = failure_reason
                            logger.error(f"Failed to send WhatsApp to {recipient_name}: {success}")
                            
                    except Exception as e:
                        exception_reason = f"WhatsApp exception: {str(e)}"
                        recipient_fail_reasons.append(exception_reason)
                        channel_status["whatsapp"]["status"] = "failed"
                        channel_status["whatsapp"]["reason"] = exception_reason
                        logger.error(f"Error sending WhatsApp to {recipient_name}: {e}")
                else:
                    whatsapp_skip_reason = "WhatsApp: Missing phone number or country code"
                    recipient_fail_reasons.append(whatsapp_skip_reason)
                    channel_status["whatsapp"]["status"] = "skipped"
                    channel_status["whatsapp"]["reason"] = whatsapp_skip_reason
            else:
                channel_status["whatsapp"]["status"] = "not_enabled"
            
            # Count overall success: recipient is successful if at least one enabled channel succeeded
            enabled_channels = [ch for ch, status in channel_status.items() if status.get("status") != "not_enabled"]
            has_success = any(
                channel_status[ch].get("status") == "success" 
                for ch in enabled_channels
            )
            
            if has_success:
                result["delivery_stats"]["successful_deliveries"] += 1
            else:
                result["delivery_stats"]["failed_deliveries"] += 1

            # Add to failed_contacts if ANY enabled channel failed or was skipped
            # This allows the frontend to show which specific channel failed for each contact
            any_channel_failed_or_skipped = any(
                channel_status[ch].get("status") in ("failed", "skipped")
                for ch in enabled_channels
            ) if enabled_channels else False

            if any_channel_failed_or_skipped:
                reason_text = " | ".join(recipient_fail_reasons) if recipient_fail_reasons else "No delivery methods attempted"
                result["delivery_stats"]["failed_contacts"].append({
                    "name": recipient_name,
                    "contact": contact_display,
                    "reason": reason_text,
                    "channel_status": channel_status
                })
        
        # Delete generated documents after notifications are completed
        # Only delete if Drive upload succeeded
        if pdf_created and pdf_path:
            try:
                from Utils.process_utils import delete_generated_files
                
                # Check if Drive upload was successful
                drive_upload_status = result.get("drive_upload", {}).get("success") if "drive_upload" in result else drive_upload_success
                
                deletion_result = delete_generated_files(
                    file_paths=[pdf_path, docx_path],
                    drive_upload_success=drive_upload_status,
                    logger=logger
                )
                
                if deletion_result['success']:
                    logger.info(f"Deleted order documents after successful Drive upload: {pdf_path} and {docx_path}")
                elif deletion_result['skipped_reason']:
                    logger.info(f"Kept order documents: {deletion_result['skipped_reason']} - {pdf_path}")
                else:
                    if deletion_result['failed_files']:
                        logger.warning(f"Failed to delete some order files: {deletion_result['failed_files']}")
            except Exception as e:
                logger.error(f"Error during order document cleanup: {e}")
        else:
            # If PDF wasn't created, still try to clean up DOCX
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                    logger.info(f"Cleaned up temporary DOCX: {docx_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up DOCX file: {e}")
        
        # Set final status
        successful = result["delivery_stats"]["successful_deliveries"]
        total = result["delivery_stats"]["total_recipients"]
        
        if successful > 0:
            result["success"] = True
            result["message"] = f"Notifications sent successfully to {successful}/{total} recipients"
        else:
            result["message"] = "Failed to send notifications to any recipients"
        
        logger.info(f"Order notification processing completed: {successful}/{total} successful")
        
        return result
        
    except Exception as e:
        logger.error(f"Error in process_order_notification: {e}")
        return {
            "success": False,
            "message": f"Error processing order notification: {str(e)}",
            "errors": [str(e)],
            "delivery_stats": {
                "total_recipients": len(recipients),
                "successful_deliveries": 0,
                "failed_deliveries": len(recipients),
                "email_successful": 0,
                "whatsapp_successful": 0,
                "failed_contacts": []
            }
        }

def process_material_notification(material_data, notification_type, recipients, method, factory, template_path, output_dir, user_email, logger, send_whatsapp_message):
    """
    Process material inward/outward notification by creating a formatted document and sending via email/WhatsApp
    """
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        
        # Initialize result tracking
        result = {
            "success": False,
            "message": "",
            "delivery_stats": {
                "total_recipients": len(recipients),
                "successful_deliveries": 0,
                "failed_deliveries": 0,
                "email_successful": 0,
                "whatsapp_successful": 0,
                "failed_contacts": []
            },
            "errors": [],
            "warnings": []
        }
        
        # Determine notification type details
        is_inward = notification_type == 'material_inward'
        notification_title = "Material Inward" if is_inward else "Material Outward"
        logger.info(f"Processing {notification_title} notification for Factory: {factory}")
        logger.info(f"Recipients count: {len(recipients)}, Method: {method}")
        
        # Create document
        doc = None
        try:
            if template_path and os.path.exists(template_path):
                logger.info(f"Attempting to use template: {template_path}")
                doc = Document(template_path)
                logger.info(f"Template loaded successfully: {template_path}")
            else:
                logger.info("Template not found, creating new document")
                doc = Document()
        except Exception as e:
            logger.warning(f"Error loading template, falling back to new document: {e}")
            try:
                doc = Document()
                logger.info("Fallback: Created new document successfully")
            except Exception as e2:
                logger.error(f"Error creating new document: {e2}")
                raise
        
        if doc is None:
            raise Exception("Failed to create document")
        
        # Add title
        try:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"📦 {notification_title} - {factory}")
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.error(f"Error adding title paragraph: {e}")
            raise
        
        doc.add_paragraph()  # Spacing
        
        # Prepare material details data
        if is_inward:
            details_data = [
                ("Date & Time", material_data.get('dateTime', 'N/A')),
                ("Party Name", material_data.get('partyName', 'N/A')),
                ("Place", material_data.get('place', 'N/A')),
            ]
        else:  # outward
            details_data = [
                ("Date & Time", material_data.get('dateTime', 'N/A')),
                ("Given To", material_data.get('givenTo', 'N/A')),
                ("Description", material_data.get('description', 'N/A')),
            ]
        
        # Create Details Table
        try:
            details_heading = doc.add_paragraph(f"{notification_title} Details")
            details_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in details_heading.runs:
                run.bold = True
                run.font.size = Pt(12)
            
            details_table = doc.add_table(rows=len(details_data), cols=2)
            details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add details
            for i, (label, value) in enumerate(details_data):
                row = details_table.rows[i]
                cells = row.cells
                cells[0].text = label
                cells[1].text = str(value)
            
            # Format cells
            for i, (label, value) in enumerate(details_data):
                row = details_table.rows[i]
                cells = row.cells
                for cell in cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
            
            # Apply professional formatting
            try:
                format_table_professional(details_table, is_header=False, logger=logger)
            except Exception as e:
                logger.warning(f"Could not apply professional formatting to details table: {e}")
            
        except Exception as e:
            logger.error(f"Error creating details table: {e}")
            raise
        
        doc.add_paragraph()  # Spacing
        
        # Create Items Table
        items_heading = doc.add_paragraph(f"{notification_title} Items")
        items_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in items_heading.runs:
            run.bold = True
            run.font.size = Pt(12)
        
        # Get items based on type
        items = material_data.get('inwardItems', []) if is_inward else material_data.get('outwardItems', [])
        
        if items:
            # Create table with headers
            items_table = doc.add_table(rows=1, cols=7)
            items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add header row
            header_cells = items_table.rows[0].cells
            headers = ['S.No', 'Category', 'Sub Category', 'Material Name', 'Specifications', 'Quantity', 'UOM']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            # Add data rows
            for idx, item in enumerate(items, 1):
                row = items_table.add_row()
                cells = row.cells
                
                cells[0].text = str(idx)
                cells[1].text = item.get('category', '')
                cells[2].text = item.get('subCategory') or '-'
                cells[3].text = item.get('materialName', '')
                cells[4].text = item.get('specifications') or '-'
                cells[5].text = str(item.get('quantity', ''))
                cells[6].text = item.get('uom', '')
                
                # Format cells
                for i, cell in enumerate(cells):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        if i == 0:  # S.No
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif i in [5, 6]:  # Quantity and UOM
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # Apply professional formatting
            try:
                format_table_professional(items_table, is_header=True, logger=logger)
            except Exception as e:
                logger.warning(f"Could not apply professional formatting to items table: {e}")
        
        # Add quantity updates if available
        quantity_updates = material_data.get('quantityUpdates', [])
        if quantity_updates:
            doc.add_paragraph()  # Spacing
            qty_heading = doc.add_paragraph("Quantity Updates")
            qty_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in qty_heading.runs:
                run.bold = True
                run.font.size = Pt(12)
            
            qty_table = doc.add_table(rows=1, cols=4)
            qty_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add header row
            qty_header_cells = qty_table.rows[0].cells
            qty_headers = ['Material', 'Previous', 'New', 'Change']
            for i, header_text in enumerate(qty_headers):
                qty_header_cells[i].text = header_text
                qty_header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in qty_header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            # Add data rows
            for update in quantity_updates:
                row = qty_table.add_row()
                cells = row.cells
                
                cells[0].text = update.get('material', '')
                cells[1].text = str(update.get('previous', ''))
                cells[2].text = str(update.get('new', ''))
                change_label = f"+{update.get('added', '')}" if is_inward else f"-{update.get('removed', '')}"
                cells[3].text = change_label
                
                # Format cells
                for i, cell in enumerate(cells):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # Apply professional formatting
            try:
                format_table_professional(qty_table, is_header=True, logger=logger)
            except Exception as e:
                logger.warning(f"Could not apply professional formatting to quantity table: {e}")
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by Bajaj Earths - {factory} Store")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.italic = True
        
        # Save document
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"{notification_type}_{factory}_{timestamp_str}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        doc.save(docx_path)
        
        logger.info(f"{notification_title} document created: {docx_path}")
        
        # Convert to PDF
        pdf_filename = f"{notification_type}_{factory}_{timestamp_str}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        pdf_created = False
        if convert_docx_to_pdf(docx_path, pdf_path):
            logger.info(f"Successfully converted to PDF: {pdf_path}")
            pdf_created = True
            notification_file = pdf_path
        else:
            logger.warning("PDF conversion failed, using DOCX file")
            result["warnings"].append("PDF conversion failed, using DOCX file")
            notification_file = docx_path
        
        # Upload PDF to Google Drive before notifications (only if PDF was created)
        drive_upload_success = None  # None = not attempted, True = success, False = failed
        if pdf_created and pdf_path:
            try:
                from Utils.drive_utils import upload_store_document_to_drive
                
                # Get date string from material_data
                date_string = material_data.get('dateTime', '')
                if not date_string:
                    # Fallback: use current date
                    date_string = datetime.now().strftime("%d/%m/%Y, %I:%M:%S %p")
                
                # Determine process type
                process_type = "material_inward" if is_inward else "material_outward"
                
                logger.info(f"Uploading {notification_type} PDF to Google Drive (factory: {factory}, date: {date_string})")
                upload_success, file_id, folder_id, upload_error = upload_store_document_to_drive(
                    pdf_path=pdf_path,
                    factory=factory,
                    date_string=date_string,
                    process_type=process_type,
                    file_name=pdf_filename,
                    logger=logger
                )
                
                if upload_success:
                    logger.info(f"Successfully uploaded {notification_type} PDF to Google Drive. File ID: {file_id}, Folder ID: {folder_id}")
                    drive_upload_success = True
                    result["drive_upload"] = {
                        "success": True,
                        "file_id": file_id,
                        "folder_id": folder_id
                    }
                else:
                    logger.warning(f"Failed to upload {notification_type} PDF to Google Drive: {upload_error}")
                    result["warnings"].append(f"Google Drive upload failed: {upload_error}")
                    drive_upload_success = False
                    result["drive_upload"] = {
                        "success": False,
                        "error": upload_error
                    }
            except Exception as e:
                logger.error(f"Error during Google Drive upload: {e}")
                result["warnings"].append(f"Error during Google Drive upload: {e}")
                drive_upload_success = False
                result["drive_upload"] = {
                    "success": False,
                    "error": str(e)
                }
        
        # Prepare notification message
        if is_inward:
            notification_message = f"""📥 *Material Inward Notification*

*Factory:* {factory}
*Date & Time:* {material_data.get('dateTime', 'N/A')}
*Party Name:* {material_data.get('partyName', 'N/A')}
*Place:* {material_data.get('place', 'N/A')}

*Total Items:* {len(items)}

Please find the detailed inward document attached."""
        else:
            notification_message = f"""📤 *Material Outward Notification*

*Factory:* {factory}
*Date & Time:* {material_data.get('dateTime', 'N/A')}
*Given To:* {material_data.get('givenTo', 'N/A')}

*Description:*
{material_data.get('description', 'N/A')}

*Total Items:* {len(items)}

Please find the detailed outward document attached."""
        
        # Send notifications to recipients
        send_email = method in ['email', 'both']
        send_whatsapp = method in ['whatsapp', 'both']
        
        for recipient in recipients:
            recipient_name = recipient.get('Name', 'Unknown')
            recipient_email = recipient.get('Email ID - To', '').strip()
            recipient_phone_raw = recipient.get('Contact No.', '').strip()
            country_code = recipient.get('Country Code', '91').strip() or '91'
            
            # Format phone number
            recipient_phone = f"{country_code}{recipient_phone_raw}".replace(' ', '') if recipient_phone_raw else ''
            contact_identifiers = [value for value in [recipient_email, recipient_phone] if value]
            contact_display = " | ".join(contact_identifiers) if contact_identifiers else "N/A"
            recipient_success = False
            recipient_fail_reasons = []
            channel_status = {
                "email": {"status": "not_enabled", "reason": ""},
                "whatsapp": {"status": "not_enabled", "reason": ""}
            }
            
            logger.info(f"Processing recipient: {recipient_name}")
            
            # Send email
            if send_email:
                channel_status["email"]["status"] = "pending"
                if recipient_email:
                    try:
                        email_subject = f"{notification_title} - {factory}"
                        if is_inward:
                            email_body = f"""
                        <html>
                        <body>
                        <h2>Material Inward Notification</h2>
                        <p>Dear <b>{recipient_name}</b>,</p>
                        <p>New material inward has been recorded. Please find the details below:</p>
                        
                        <table style="border-collapse: collapse; margin: 20px 0;">
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Factory:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{factory}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Date & Time:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{material_data.get('dateTime', 'N/A')}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Party Name:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{material_data.get('partyName', 'N/A')}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Place:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{material_data.get('place', 'N/A')}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Total Items:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{len(items)}</td>
                            </tr>
                        </table>
                        
                        <p>Please find the detailed inward document attached.</p>
                        
                        <p>Best regards,<br>Bajaj Earths - {factory} Store</p>
                        </body>
                        </html>
                        """
                        else:
                            email_body = f"""
                        <html>
                        <body>
                        <h2>Material Outward Notification</h2>
                        <p>Dear <b>{recipient_name}</b>,</p>
                        <p>New material outward has been recorded. Please find the details below:</p>
                        
                        <table style="border-collapse: collapse; margin: 20px 0;">
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Factory:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{factory}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Date & Time:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{material_data.get('dateTime', 'N/A')}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Given To:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{material_data.get('givenTo', 'N/A')}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Description:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{material_data.get('description', 'N/A')}</td>
                            </tr>
                            <tr>
                                <td style="padding: 8px; border: 1px solid #ddd;"><strong>Total Items:</strong></td>
                                <td style="padding: 8px; border: 1px solid #ddd;">{len(items)}</td>
                            </tr>
                        </table>
                        
                        <p>Please find the detailed outward document attached.</p>
                        
                        <p>Best regards,<br>Bajaj Earths - {factory} Store</p>
                        </body>
                        </html>
                        """
                        
                        success = send_email_gmail_api(
                            user_email=user_email,
                            recipient_email=recipient_email,
                            subject=email_subject,
                            body=email_body,
                            attachment_paths=[notification_file]
                        )
                        
                        if success is True:
                            recipient_success = True
                            channel_status["email"]["status"] = "success"
                            result["delivery_stats"]["email_successful"] += 1
                            logger.info(f"Email sent successfully to {recipient_name} ({recipient_email})")
                        else:
                            failure_reason = f"Email error: {success}" if isinstance(success, str) else "Email send failed"
                            recipient_fail_reasons.append(failure_reason)
                            channel_status["email"]["status"] = "failed"
                            channel_status["email"]["reason"] = failure_reason
                            logger.error(f"Failed to send email to {recipient_name}: {success}")
                            
                    except Exception as e:
                        exception_reason = f"Email exception: {str(e)}"
                        recipient_fail_reasons.append(exception_reason)
                        channel_status["email"]["status"] = "failed"
                        channel_status["email"]["reason"] = exception_reason
                        logger.error(f"Error sending email to {recipient_name}: {e}")
                else:
                    no_email_reason = "Email: No recipient address provided"
                    recipient_fail_reasons.append(no_email_reason)
                    channel_status["email"]["status"] = "skipped"
                    channel_status["email"]["reason"] = no_email_reason
            else:
                channel_status["email"]["status"] = "not_enabled"
            
            # Send WhatsApp
            if send_whatsapp:
                channel_status["whatsapp"]["status"] = "pending"
                if recipient_phone_raw and country_code:
                    try:
                        success = send_whatsapp_message(
                            contact_name=recipient_name,
                            message=notification_message,
                            file_paths=[notification_file],
                            whatsapp_number=recipient_phone,
                            process_name="material_notification"
                        )
                        
                        if success is True:
                            recipient_success = True
                            channel_status["whatsapp"]["status"] = "success"
                            result["delivery_stats"]["whatsapp_successful"] += 1
                            logger.info(f"WhatsApp sent successfully to {recipient_name} ({recipient_phone})")
                        else:
                            failure_reason = f"WhatsApp error: {success}" if isinstance(success, str) else "WhatsApp send failed"
                            recipient_fail_reasons.append(failure_reason)
                            channel_status["whatsapp"]["status"] = "failed"
                            channel_status["whatsapp"]["reason"] = failure_reason
                            logger.error(f"Failed to send WhatsApp to {recipient_name}: {success}")
                            
                    except Exception as e:
                        exception_reason = f"WhatsApp exception: {str(e)}"
                        recipient_fail_reasons.append(exception_reason)
                        channel_status["whatsapp"]["status"] = "failed"
                        channel_status["whatsapp"]["reason"] = exception_reason
                        logger.error(f"Error sending WhatsApp to {recipient_name}: {e}")
                else:
                    whatsapp_skip_reason = "WhatsApp: Missing phone number or country code"
                    recipient_fail_reasons.append(whatsapp_skip_reason)
                    channel_status["whatsapp"]["status"] = "skipped"
                    channel_status["whatsapp"]["reason"] = whatsapp_skip_reason
            else:
                channel_status["whatsapp"]["status"] = "not_enabled"
            
            # Count overall success: recipient is successful if at least one enabled channel succeeded
            enabled_channels = [ch for ch, status in channel_status.items() if status.get("status") != "not_enabled"]
            has_success = any(
                channel_status[ch].get("status") == "success" 
                for ch in enabled_channels
            )
            
            if has_success:
                result["delivery_stats"]["successful_deliveries"] += 1
            else:
                result["delivery_stats"]["failed_deliveries"] += 1

            # Add to failed_contacts if ANY enabled channel failed or was skipped
            # This allows the frontend to show which specific channel failed for each contact
            any_channel_failed_or_skipped = any(
                channel_status[ch].get("status") in ("failed", "skipped")
                for ch in enabled_channels
            ) if enabled_channels else False

            if any_channel_failed_or_skipped:
                reason_text = " | ".join(recipient_fail_reasons) if recipient_fail_reasons else "No delivery methods attempted"
                result["delivery_stats"]["failed_contacts"].append({
                    "name": recipient_name,
                    "contact": contact_display,
                    "reason": reason_text,
                    "channel_status": channel_status
                })
        
        # Delete generated documents after notifications are completed
        # Only delete if Drive upload succeeded
        if pdf_created and pdf_path:
            try:
                from Utils.process_utils import delete_generated_files
                
                # Check if Drive upload was successful
                drive_upload_status = result.get("drive_upload", {}).get("success") if "drive_upload" in result else drive_upload_success
                
                deletion_result = delete_generated_files(
                    file_paths=[pdf_path, docx_path],
                    drive_upload_success=drive_upload_status,
                    logger=logger
                )
                
                if deletion_result['success']:
                    logger.info(f"Deleted {notification_type} documents after successful Drive upload: {pdf_path} and {docx_path}")
                elif deletion_result['skipped_reason']:
                    logger.info(f"Kept {notification_type} documents: {deletion_result['skipped_reason']} - {pdf_path}")
                else:
                    if deletion_result['failed_files']:
                        logger.warning(f"Failed to delete some {notification_type} files: {deletion_result['failed_files']}")
            except Exception as e:
                logger.error(f"Error during {notification_type} document cleanup: {e}")
        else:
            # If PDF wasn't created, still try to clean up DOCX
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                    logger.info(f"Cleaned up temporary DOCX: {docx_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up DOCX file: {e}")
        
        # Set final status
        successful = result["delivery_stats"]["successful_deliveries"]
        total = result["delivery_stats"]["total_recipients"]
        
        if successful > 0:
            result["success"] = True
            result["message"] = f"Notifications sent successfully to {successful}/{total} recipients"
        else:
            result["message"] = "Failed to send notifications to any recipients"
        
        logger.info(f"{notification_title} notification processing completed: {successful}/{total} successful")
        
        return result
        
    except Exception as e:
        logger.error(f"Error in process_material_notification: {e}")
        return {
            "success": False,
            "message": f"Error processing material notification: {str(e)}",
            "errors": [str(e)],
            "delivery_stats": {
                "total_recipients": len(recipients),
                "successful_deliveries": 0,
                "failed_deliveries": len(recipients),
                "email_successful": 0,
                "whatsapp_successful": 0,
                "failed_contacts": []
            }
        }
