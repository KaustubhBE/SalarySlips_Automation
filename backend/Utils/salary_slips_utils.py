import os
import re
import json
import logging
from docx import Document
from Utils.email_utils import send_email_with_attachment, get_employee_email
from Utils.whatsapp_utils import send_whatsapp_message, get_employee_contact
from Utils.drive_utils import upload_to_google_drive
import shutil
import subprocess
# import platform
# import pythoncom
# from comtypes.client import CreateObject

# Configure logging
logging.basicConfig(level=logging.INFO)

# Load message templates
def load_message_templates():
    try:
        with open('backend/Utils/message.json', 'r') as f:
            return json.load(f)
    except Exception as e:
        logging.error(f"Error loading message templates: {e}")
        return None

# Preprocess headers
def preprocess_headers(headers):
    return [header.replace("\n", " ").strip().strip('"') for header in headers]

def convert_docx_to_pdf(input_path, output_path):
    try:
        process = subprocess.Popen([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path),
            input_path
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        process.wait()

        # pythoncom.CoInitialize()
        # word = CreateObject('Word.Application')
        # word.Visible = False  # Keep Word hidden
        # doc = word.Documents.Open(input_path)
        # doc.SaveAs(output_path, FileFormat=17)  # 17 is PDF format
        # doc.Close()
        # word.Quit()
        # print(f'Converted {input_path} to {output_path}')
        # return True

        return True
    except Exception as e:
        logging.error(f"Error converting DOCX to PDF: {e}")
        return False

# Format file path
def format_file_path(file_path):
    if isinstance(file_path, str):
        return file_path.replace("\\\\", "\\")
    elif isinstance(file_path, list):
        return [f.replace("\\\\", "\\") for f in file_path]
    return file_path

# Clear the Salary_Slips folder
def clear_salary_slips_folder(output_dir):
    try:
        for filename in os.listdir(output_dir):
            file_path = os.path.join(output_dir, filename)
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        logging.info(f"Cleared the contents of the folder: {output_dir}")
    except Exception as e:
        logging.error(f"Error clearing the folder {output_dir}: {e}")

def format_months_list(months_data):
    """Format the list of months for display in messages."""
    if not months_data:
        return ""
    
    if isinstance(months_data, list):
        # Format for WhatsApp message
        return "\n".join([f"   -  {month['month']} {month['year']}" for month in months_data])
    else:
        # Format for email HTML
        return "".join([f"<li>{month['month']} {month['year']}</li>" for month in months_data])

def handle_whatsapp_notification(contact_name, full_month, full_year, whatsapp_number, file_path, is_special=False, months_data=None):
    if whatsapp_number:
        # Load message templates
        templates = load_message_templates()
        if not templates:
            logging.error("Failed to load message templates")
            return False
            
        # Choose appropriate template
        template_key = "special_whatsapp_message" if is_special else "monthly_whatsapp_message"
        
        # Format the message with actual values
        if is_special and months_data:
            message = [line.format(
                contact_name=contact_name,
                months_list=format_months_list(months_data)
            ) for line in templates[template_key]]
        else:
            message = [line.format(
                contact_name=contact_name,
                full_month=full_month,
                full_year=full_year
            ) for line in templates[template_key]]
        
        # Log attempt
        logging.info(f"Attempting to send WhatsApp message to {contact_name} ({whatsapp_number})")
        
        # Send message
        return send_whatsapp_message(contact_name, message, file_path, whatsapp_number)
    return False

def handle_whatsapp_notification(contact_name, full_month, full_year, whatsapp_number, file_path, is_special=False, months_data=None):
    if whatsapp_number:
        # Load message templates
        templates = load_message_templates()
        if not templates:
            logging.error("Failed to load message templates")
            return False
            
        # Choose appropriate template
        template_key = "special_whatsapp_message" if is_special else "monthly_whatsapp_message"
        
        # Format the message with actual values
        if is_special and months_data:
            message = [line.format(
                contact_name=contact_name,
                months_list=format_months_list(months_data)
            ) for line in templates[template_key]]
        else:
            message = [line.format(
                contact_name=contact_name,
                full_month=full_month,
                full_year=full_year
            ) for line in templates[template_key]]
        
        # Log attempt
        logging.info(f"Attempting to send WhatsApp message to {contact_name} ({whatsapp_number})")
        
        # Send message
        return send_whatsapp_message(contact_name, message, file_path, whatsapp_number)
    return False

# Generate and process salary slips for a single employee
def process_salary_slip(template_path, output_dir, employee_data, headers, drive_data, email_employees, contact_employees, month, year, full_month, full_year, send_whatsapp, send_email, is_special=False, months_data=None, collected_pdfs=None):
    logging.info("Starting process_salary_slip function")
    headers = preprocess_headers(headers)
    placeholders = dict(zip(headers, employee_data))

    # Add month and year placeholders to the dictionary
    placeholders["Month"] = full_month
    placeholders["Year"] = full_year

    # Merge data from "Official Details" sheet
    official_details = next((item for item in drive_data if item.get("Employee Code") == placeholders.get("Employee Code") or 
                           item.get("Employee\nCode") == placeholders.get("Employee\nCode")), {})
    placeholders.update(official_details)

    # Calculate components of salary
    try:
        present_salary_str = placeholders.get("Present Salary", "")
        present_salary = float(re.sub(r'[^\d.]', '', present_salary_str))
        if present_salary <= 0:
            raise ValueError("Present Salary must be greater than zero.")
        placeholders["BS"] = str(round(present_salary * 0.40))
        placeholders["HRA"] = str(round(present_salary * 0.20))
        placeholders["SA"] = str(round(present_salary * 0.40))
    except ValueError as e:
        logging.error(f"Invalid Present Salary for {placeholders.get('Name', 'Unknown')}: {e}. Skipping.")
        return

    # Ensure all placeholders are strings
    placeholders = {k: str(v) for k, v in placeholders.items()}

    # Load template and replace placeholders
    try:
        template = Document(template_path)
        for paragraph in template.paragraphs:
            for run in paragraph.runs:
                for placeholder, value in placeholders.items():
                    if f"{{{placeholder}}}" in run.text:
                        run.text = run.text.replace(f"{{{placeholder}}}", value)

        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for placeholder, value in placeholders.items():
                                if f"{{{placeholder}}}" in run.text:
                                    run.text = run.text.replace(f"{{{placeholder}}}", value)

        # Save output files
        employee_name = re.sub(r'[^\w\s]', '', placeholders.get("Name", "Employee"))
        output_docx = os.path.join(output_dir, f"Salary Slip_{employee_name}_{month}{year}.docx")
        template.save(output_docx)
        output_pdf = os.path.join(output_dir, f"Salary Slip_{employee_name}_{month}{year}.pdf")
        template.save(output_pdf)

        if convert_docx_to_pdf(output_docx, output_pdf):
            # Upload to Google Drive
            folder_id = official_details.get("Google Drive ID")
            if folder_id:
                logging.info(f"Found Google Drive ID {folder_id} for employee {employee_name}")
                upload_success = upload_to_google_drive(output_pdf, folder_id, employee_name, month, year)
            else:
                logging.error(f"No Google Drive ID found for employee: {employee_name}")
                logging.error(f"Available keys in drive data: {list(drive_data[0].keys()) if drive_data else 'No drive data'}")

            # Load message templates
            templates = load_message_templates()
            if not templates:
                logging.error("Failed to load message templates")
                return

            # If this is part of a multi-month process, collect the PDF
            if collected_pdfs is not None:
                collected_pdfs.append(output_pdf)
                return

            # Send email if enabled
            if send_email:
                recipient_email = get_employee_email(placeholders.get("Name"), email_employees)
                if recipient_email:
                    email_subject = f"Salary Slip{'s' if is_special else ''} for {full_month} {full_year} - Bajaj Earths Pvt. Ltd."
                    template_key = "special_email_body" if is_special else "monthly_email_body"
                    email_body = templates[template_key].format(
                        employee_name=placeholders.get("Name"),
                        full_month=full_month,
                        full_year=full_year,
                        months_list=format_months_list(months_data) if is_special else ""
                    )
                    logging.info(f"Sending email to {recipient_email}")
                    send_email_with_attachment(recipient_email, email_subject, email_body, output_pdf)
                else:
                    logging.info(f"No email found for {placeholders.get('Name')}.")
            
            # Send WhatsApp message if enabled
            if send_whatsapp:
                contact_name = placeholders.get("Name")
                whatsapp_number = get_employee_contact(contact_name, contact_employees)
                handle_whatsapp_notification(
                    contact_name,
                    full_month,
                    full_year,
                    whatsapp_number,
                    output_pdf,
                    is_special,
                    months_data
                )
    except Exception as e:
        logging.error(f"Error processing salary slip for {placeholders.get('Name', 'Unknown')}: {e}")
    logging.info("Finished process_salary_slip function")

# Generate and process salary slips for multiple employees (batch processing)
def process_salary_slips(template_path, output_dir, employee_data, headers, drive_data, email_employees, contact_employees, month, year, full_month, full_year, send_whatsapp, send_email):
    logging.info("Starting process_salary_slip function")
    headers = preprocess_headers(headers)
    placeholders = dict(zip(headers, employee_data))

    # Add month and year placeholders to the dictionary
    placeholders["Month"] = full_month
    placeholders["Year"] = full_year

    # Merge data from "Official Details" sheet
    official_details = next((item for item in drive_data if item.get("Employee Code") == placeholders.get("Employee Code") or 
                           item.get("Employee\nCode") == placeholders.get("Employee\nCode") or 
                           item.get("Name") == placeholders.get("Name")), {})
    placeholders.update(official_details)

    # Calculate components of salary
    try:
        present_salary_str = placeholders.get("Present Salary", "")
        present_salary = float(re.sub(r'[^\d.]', '', present_salary_str))
        if present_salary <= 0:
            raise ValueError("Present Salary must be greater than zero.")
        placeholders["BS"] = str(round(present_salary * 0.40))
        placeholders["HRA"] = str(round(present_salary * 0.20))
        placeholders["SA"] = str(round(present_salary * 0.40))
    except ValueError as e:
        logging.error(f"Invalid Present Salary for {placeholders.get('Name', 'Unknown')}: {e}. Skipping.")
        return

    # Ensure all placeholders are strings
    placeholders = {k: str(v) for k, v in placeholders.items()}

    # Load template and replace placeholders
    try:
        template = Document(template_path)
        for paragraph in template.paragraphs:
            for run in paragraph.runs:
                for placeholder, value in placeholders.items():
                    if f"{{{placeholder}}}" in run.text:
                        run.text = run.text.replace(f"{{{placeholder}}}", value)

        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for placeholder, value in placeholders.items():
                                if f"{{{placeholder}}}" in run.text:
                                    run.text = run.text.replace(f"{{{placeholder}}}", value)

        # Save output files
        employee_name = re.sub(r'[^\w\s]', '', placeholders.get("Name", "Employee"))
        output_docx = os.path.join(output_dir, f"Salary Slip_{employee_name}_{month}{year}.docx")
        template.save(output_docx)
        output_pdf = os.path.join(output_dir, f"Salary Slip_{employee_name}_{month}{year}.pdf")
        template.save(output_pdf)

        if convert_docx_to_pdf(output_docx, output_pdf):
            # Upload to Google Drive
            folder_id = official_details.get("Google Drive ID")
            if folder_id:
                logging.info(f"Found Google Drive ID {folder_id} for employee {employee_name}")
                upload_success = upload_to_google_drive(output_pdf, folder_id, employee_name, month, year)
            else:
                logging.error(f"No Google Drive ID found for employee: {employee_name}")
                logging.error(f"Available keys in drive data: {list(drive_data[0].keys()) if drive_data else 'No drive data'}")

            # Load message templates
            templates = load_message_templates()
            if not templates:
                logging.error("Failed to load message templates")
                return

            # Send email if enabled
            if send_email:
                recipient_email = get_employee_email(placeholders.get("Name"), email_employees)
                if recipient_email:
                    email_subject = f"Salary Slip for {full_month} {full_year} - Bajaj Earths Pvt. Ltd."
                    email_body = templates["monthly_email_body"].format(
                        employee_name=placeholders.get("Name"),
                        full_month=full_month,
                        full_year=full_year
                    )
                    logging.info(f"Sending email to {recipient_email}")
                    send_email_with_attachment(recipient_email, email_subject, email_body, output_pdf)
                else:
                    logging.info(f"No email found for {placeholders.get('Name')}.")
            
            # Send WhatsApp message if enabled
            if send_whatsapp:
                contact_name = placeholders.get("Name")
                whatsapp_number = get_employee_contact(contact_name, contact_employees)
                handle_whatsapp_notification(
                    contact_name,
                    full_month,
                    full_year,
                    whatsapp_number,
                    output_pdf
                )
    except Exception as e:
        logging.error(f"Error processing salary slip for {placeholders.get('Name', 'Unknown')}: {e}")
    logging.info("Finished process_salary_slip function")